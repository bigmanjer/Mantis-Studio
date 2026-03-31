#!/usr/bin/env python3
"""
MANTIS Studio - Main Entry Point

This is the primary Streamlit application entry point for MANTIS Studio.
It implements a state-driven (not page-driven) navigation system.

Run with:
    streamlit run app/main.py

Or with utility flags:
    python -m app.main --selftest    # Run self-tests
    python -m app.main --repair      # Repair project files

Architecture:
    - State-based navigation via st.session_state.page
    - Views are available in app/views/ directory
    - Business logic in app/services/ directory
    - All navigation happens within this single entry point
    - NO Streamlit multipage routing (no /pages directory)

Directory Structure:
    app/
    +-- state.py              # Session state schema + defaults
    +-- router.py             # Central navigation logic
    +-- layout/               # UI layout components (sidebar, header, styles)
    +-- views/                # UI screens (dashboard, projects, editor, etc.)
    +-- components/           # Reusable UI blocks
    +-- services/             # Business logic (projects, storage, auth, AI, export)
    +-- utils/                # Utilities (versioning, helpers)
"""

import datetime
import difflib
import json
import logging
import os
import random
import re
import shutil
import sys
import time
import uuid
import webbrowser
from collections.abc import Generator
from dataclasses import asdict, dataclass, field, fields
from pathlib import Path
from string import Template
from typing import Any, Callable, Dict, List, Optional

import requests

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from app.utils.navigation import get_nav_config
from app.utils.branding_assets import read_asset_bytes, resolve_asset_path

# NOTE: Streamlit-dependent utilities are imported inside _run_ui() so
# `python -m app.main --selftest` can run without Streamlit installed.


# ===== v45 BRANDING (SAFE, ORIGINAL TEMPLATE) =====
import base64
# ==================================================

SELFTEST_MODE = "--selftest" in sys.argv
REPAIR_MODE = "--repair" in sys.argv


# ============================================================
# 1) CONFIG
# ============================================================

def get_app_version() -> str:
    """Return the user-visible app version.

    Priority:
      1) Explicit env var (useful for CI/CD overrides)
      2) VERSION.txt next to this file (repo-controlled)
      3) A safe fallback string
    """
    env_version = os.getenv("MANTIS_APP_VERSION")
    if env_version:
        return env_version

    try:
        version_path = Path(__file__).resolve().parents[1] / "VERSION.txt"
        if version_path.exists():
            raw = version_path.read_text(encoding="utf-8").strip()
            if raw:
                return raw
    except Exception:
        # Never block app start on version metadata.
        pass

    return "47.0"


def _safe_int_env(env_var: str, default: int) -> int:
    raw = os.getenv(env_var, "")
    if not raw:
        return default
    try:
        return int(raw)
    except (ValueError, TypeError):
        return default


def _safe_float_env(env_var: str, default: float) -> float:
    raw = os.getenv(env_var, "")
    if not raw:
        return default
    try:
        return float(raw)
    except (ValueError, TypeError):
        return default


def _safe_bool_env(env_var: str, default: bool) -> bool:
    raw = os.getenv(env_var, "")
    if not raw:
        return default
    return str(raw).strip().lower() in {"1", "true", "yes", "on"}


class AppConfig:
    APP_NAME = "MANTIS Studio"
    VERSION = get_app_version()
    PROJECTS_DIR = os.getenv("MANTIS_PROJECTS_DIR", "projects")
    BACKUPS_DIR = os.path.join(PROJECTS_DIR, ".backups")

    GROQ_API_URL = os.getenv("GROQ_API_URL", "https://api.groq.com/openai/v1")
    GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
    GROQ_TIMEOUT = _safe_int_env("GROQ_TIMEOUT", 300)
    DEFAULT_MODEL = os.getenv("GROQ_MODEL", "llama3-8b-8192")
    OPENAI_API_URL = os.getenv("OPENAI_API_URL", "https://api.openai.com/v1")
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
    OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    CONFIG_PATH = os.getenv(
        "MANTIS_CONFIG_PATH",
        os.path.join(PROJECTS_DIR, ".mantis_config.json"),
    )
    MAX_PROMPT_CHARS = 16000
    SUMMARY_CONTEXT_CHARS = 4000
    MAX_UPLOAD_MB = _safe_int_env("MANTIS_MAX_UPLOAD_MB", 10)
    SAVE_LOCK_TIMEOUT = _safe_int_env("MANTIS_SAVE_LOCK_TIMEOUT", 5)
    SAVE_LOCK_RETRY_SLEEP = _safe_float_env("MANTIS_SAVE_LOCK_RETRY_SLEEP", 0.1)
    WORLD_BIBLE_CONFIDENCE = _safe_float_env("MANTIS_WORLD_BIBLE_CONFIDENCE", 0.83)
    WORLD_MEMORY_SOFT_CONFIDENCE = _safe_float_env("MANTIS_WORLD_MEMORY_SOFT_CONFIDENCE", 0.92)
    WORLD_MEMORY_HARD_CONFIDENCE = _safe_float_env("MANTIS_WORLD_MEMORY_HARD_CONFIDENCE", 0.97)
    WORLD_MEMORY_AUTO_SOFT = _safe_bool_env("MANTIS_WORLD_MEMORY_AUTO_SOFT", True)
    WORLD_MEMORY_AUTO_HARD = _safe_bool_env("MANTIS_WORLD_MEMORY_AUTO_HARD", False)
    
    # Documentation URLs
    GETTING_STARTED_URL = "https://github.com/bigmanjer/Mantis-Studio/blob/main/GETTING_STARTED.md"


os.makedirs(AppConfig.PROJECTS_DIR, exist_ok=True)
os.makedirs(AppConfig.BACKUPS_DIR, exist_ok=True)


logging.basicConfig(
    level=logging.DEBUG if os.getenv("MANTIS_DEBUG") else logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s:%(lineno)d - %(message)s"
)
logger = logging.getLogger("MANTIS")
logger.info("=" * 60)
logger.info("MANTIS Studio Starting...")
logger.info("=" * 60)


def load_app_config() -> Dict[str, str]:
    try:
        with open(AppConfig.CONFIG_PATH, "r", encoding="utf-8") as fh:
            data = json.load(fh)
            if isinstance(data, dict):
                return data
    except FileNotFoundError:
        return {}
    except Exception:
        logger.warning("Failed to load app config", exc_info=True)
    return {}


def save_app_config(data: Dict[str, str]) -> None:
    try:
        os.makedirs(os.path.dirname(AppConfig.CONFIG_PATH) or ".", exist_ok=True)
        tmp = AppConfig.CONFIG_PATH + ".tmp"
        with open(tmp, "w", encoding="utf-8") as fh:
            json.dump(data, fh, indent=2)
        os.replace(tmp, AppConfig.CONFIG_PATH)
    except Exception:
        logger.warning("Failed to save app config", exc_info=True)



def _load_json_with_encoding_fallback(path: str) -> Dict[str, Any]:
    """Load JSON with UTF-8 first and tolerant fallback encodings."""
    encodings = ("utf-8", "cp1252", "latin-1")
    last_decode_error: Optional[UnicodeDecodeError] = None
    for encoding in encodings:
        try:
            with open(path, "r", encoding=encoding) as fh:
                return json.load(fh)
        except UnicodeDecodeError as exc:
            last_decode_error = exc
            continue
        except json.JSONDecodeError:
            raise
    if last_decode_error is not None:
        raise last_decode_error
    raise ValueError(f"Cannot decode JSON file ({path})")


def _read_text_with_encoding_fallback(path: Path) -> str:
    """Read text with UTF-8 first and safe Windows fallback encodings."""
    encodings = ("utf-8", "cp1252", "latin-1")
    last_decode_error: Optional[UnicodeDecodeError] = None
    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            last_decode_error = exc
            continue
    if last_decode_error is not None:
        raise last_decode_error
    raise ValueError(f"Cannot decode text file ({path})")
def _truncate_prompt(prompt: str, limit: int) -> str:
    if not prompt:
        return prompt
    if len(prompt) <= limit:
        return prompt
    logger.warning("Prompt length %s exceeded limit %s; truncating", len(prompt), limit)
    return prompt[:limit]


def sanitize_ai_input(text: str, max_length: int = 0) -> str:
    """Sanitize user-provided text before sending to AI APIs.

    Strips leading/trailing whitespace and removes null bytes.
    When *max_length* is positive the text is truncated to that limit.
    """
    if not text:
        return ""
    cleaned = text.strip().replace("\x00", "")
    if max_length > 0 and len(cleaned) > max_length:
        cleaned = cleaned[:max_length]
    return cleaned


def _acquire_lock(lock_path: str, timeout: int, retry_sleep: float) -> bool:
    start = time.time()
    while True:
        try:
            fd = os.open(lock_path, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
            with os.fdopen(fd, "w", encoding="utf-8") as lock_file:
                lock_file.write(str(os.getpid()))
            return True
        except FileExistsError:
            try:
                age = time.time() - os.path.getmtime(lock_path)
            except OSError:
                age = 0
            if age > max(timeout * 5, 5):
                try:
                    os.remove(lock_path)
                except OSError:
                    pass
            if time.time() - start >= timeout:
                return False
            time.sleep(retry_sleep)


def _release_lock(lock_path: str) -> None:
    try:
        os.remove(lock_path)
    except OSError:
        pass


def _get_streamlit():
    if "streamlit" not in sys.modules:
        return None
    try:
        import streamlit as st

        return st
    except Exception:
        return None


def _normalize_provider(provider: str) -> str:
    return (provider or "groq").strip().lower()


def _provider_label(provider: str) -> str:
    return "OpenAI" if _normalize_provider(provider) == "openai" else "Groq"


def _ensure_session_keys(st) -> Dict[str, str]:
    keys = st.session_state.setdefault("ai_session_keys", {})
    for provider in ("openai", "groq"):
        keys.setdefault(provider, "")
    st.session_state["ai_session_keys"] = keys
    return keys


def _validate_api_key_format(key: str) -> bool:
    """Basic sanity check for API key format.

    Returns ``True`` when the key looks plausible (only printable ASCII,
    reasonable length).  This is not authoritative  the real validation
    happens server-side  but it catches obvious mistakes like pasting
    HTML or binary content into the key field.
    """
    if not key:
        return False
    if len(key) < 8 or len(key) > 256:
        return False
    # Only printable ASCII characters (no control chars or non-ASCII)
    return all(32 <= ord(ch) < 127 for ch in key)


def set_session_key(provider: str, key: str) -> None:
    st = _get_streamlit()
    if not st:
        return
    provider = _normalize_provider(provider)
    cleaned = (key or "").strip()
    if cleaned and not _validate_api_key_format(cleaned):
        logger.warning("Rejected API key for %s: invalid format", provider)
        return
    keys = _ensure_session_keys(st)
    keys[provider] = cleaned
    st.session_state["ai_session_keys"] = keys
    # Persist to config file so the key survives a page refresh
    config = load_app_config()
    config[f"{provider}_api_key"] = cleaned
    save_app_config(config)


def clear_session_key(provider: str) -> None:
    st = _get_streamlit()
    if not st:
        return
    provider = _normalize_provider(provider)
    keys = _ensure_session_keys(st)
    keys[provider] = ""
    st.session_state["ai_session_keys"] = keys
    # Remove from config file as well
    config = load_app_config()
    config.pop(f"{provider}_api_key", None)
    save_app_config(config)


def _get_secret_key(st, provider: str) -> str:
    if not st or not hasattr(st, "secrets"):
        return ""
    try:
        secrets = st.secrets
        provider = _normalize_provider(provider)
        if provider == "openai":
            return str(
                secrets.get("openai_api_key")
                or secrets.get("OPENAI_API_KEY")
                or (secrets.get("openai") or {}).get("api_key")
                or ""
            ).strip()
        return str(
            secrets.get("groq_api_key")
            or secrets.get("GROQ_API_KEY")
            or (secrets.get("groq") or {}).get("api_key")
            or ""
        ).strip()
    except (KeyError, FileNotFoundError):
        return ""


def _get_server_default_key(provider: str) -> str:
    st = _get_streamlit()
    key = _get_secret_key(st, provider)
    if key:
        return key
    provider = _normalize_provider(provider)
    if provider == "openai":
        return (AppConfig.OPENAI_API_KEY or "").strip()
    return (AppConfig.GROQ_API_KEY or "").strip()


def get_effective_key(provider: str, user_id: Optional[str] = None) -> tuple[str, str]:
    provider = _normalize_provider(provider)
    st = _get_streamlit()
    if st:
        session_key = _ensure_session_keys(st).get(provider, "")
        if session_key:
            return session_key, "session"
    # Check config file for a previously saved key
    config = load_app_config()
    saved_key = (config.get(f"{provider}_api_key") or "").strip()
    if saved_key:
        return saved_key, "saved"
    default_key = _get_server_default_key(provider)
    if default_key:
        return default_key, "default"
    return "", "none"


def get_active_provider() -> str:
    st = _get_streamlit()
    if st:
        return _normalize_provider(st.session_state.get("ai_provider", "groq"))
    return "groq"


def _get_provider_base_url(provider: str) -> str:
    provider = _normalize_provider(provider)
    st = _get_streamlit()
    if provider == "openai":
        if st:
            return st.session_state.get("openai_base_url", AppConfig.OPENAI_API_URL)
        return AppConfig.OPENAI_API_URL
    if st:
        return st.session_state.get("groq_base_url", AppConfig.GROQ_API_URL)
    return AppConfig.GROQ_API_URL


# ============================================================
# 2) MODELS
# ============================================================

@dataclass
class Entity:
    id: str
    name: str
    category: str
    description: str = ""
    aliases: List[str] = field(default_factory=list)
    tags: str = ""
    created_at: float = field(default_factory=time.time)

    def merge(self, new_desc: str):
        """Smart merge that avoids exact duplicates and formats as bullets."""
        if not new_desc:
            return
        new_desc = new_desc.strip()
        if not new_desc or new_desc in (self.description or ""):
            return

        # Normalize text to check for duplicates (basic)
        curr = (self.description or "").lower()
        current_sentences = set(re.split(r"[.!?\n]", curr))
        new_sentences = [s.strip() for s in re.split(r"[.!?\n]", new_desc) if s.strip()]

        to_add = []
        for s in new_sentences:
            if s.lower() not in current_sentences:
                to_add.append(s)

        if to_add:
            if self.description and not self.description.startswith("-"):
                self.description = f"- {self.description.strip()}"
            elif not self.description:
                self.description = ""

            for item in to_add:
                if self.description:
                    self.description += f"\n- {item}"
                else:
                    self.description = f"- {item}"


@dataclass
class Chapter:
    id: str
    index: int
    title: str = "Untitled Chapter"
    content: str = ""
    summary: str = ""
    word_count: int = 0
    target_words: int = 1000
    created_at: float = field(default_factory=time.time)
    modified_at: float = field(default_factory=time.time)
    history: List[Dict[str, Any]] = field(default_factory=list)

    def update_content(self, new_text: str, source: str = "manual"):
        if self.content == new_text:
            return
        self.history.append(
            {"timestamp": time.time(), "source": source, "previous_text": self.content}
        )
        if len(self.history) > 10:
            self.history.pop(0)
        self.content = new_text
        self.word_count = len((new_text or "").split())
        self.modified_at = time.time()

    def restore_revision(self, text: str):
        self.content = text
        self.word_count = len((text or "").split())
        self.modified_at = time.time()


@dataclass
class Project:
    id: str
    title: str
    author: str = ""
    genre: str = ""
    outline: str = ""
    memory: str = ""
    memory_hard: str = ""
    memory_soft: str = ""
    author_note: str = ""
    style_guide: str = ""
    default_word_count: int = 1000
    world_db: Dict[str, Entity] = field(default_factory=dict)
    chapters: Dict[str, Chapter] = field(default_factory=dict)
    created_at: float = field(default_factory=time.time)
    last_modified: float = field(default_factory=time.time)
    filepath: Optional[str] = None
    storage_dir: Optional[str] = None

    @classmethod
    def create(
        cls,
        title: str,
        author: str = "",
        genre: str = "",
        storage_dir: Optional[str] = None,
    ) -> "Project":
        storage_dir = storage_dir or AppConfig.PROJECTS_DIR
        os.makedirs(storage_dir, exist_ok=True)
        final_title = (title or "").strip()
        if not final_title:
            base = "Untitled Project"
            final_title = base
            counter = 1
            if os.path.exists(storage_dir):
                while True:
                    found = False
                    for f in os.listdir(storage_dir):
                        if final_title.replace(" ", "_") in f:
                            found = True
                            break
                    if not found:
                        break
                    final_title = f"{base} ({counter})"
                    counter += 1
        return cls(
            id=str(uuid.uuid4()),
            title=final_title,
            author=author.strip(),
            genre=genre.strip(),
            storage_dir=storage_dir,
        )

    @staticmethod
    def _normalize_category(category: str) -> str:
        normalized = (category or "").strip().lower()
        mapping = {
            "character": "Character",
            "characters": "Character",
            "location": "Location",
            "locations": "Location",
            "faction": "Faction",
            "factions": "Faction",
            "lore": "Lore",
            "rule": "Lore",
            "rules": "Lore",
            "item": "Item",
            "items": "Item",
        }
        return mapping.get(normalized, "Lore")

    @staticmethod
    def _normalize_entity_name(name: str) -> str:
        base = (name or "").strip().lower()
        if not base:
            return ""
        base = re.sub(r"[\"'`]", "", base)
        base = re.sub(r"\((.*?)\)", r" \1 ", base)
        base = re.sub(r"\b(mr|mrs|ms|dr|prof|sir|madam)\.?\b", "", base)
        base = re.sub(r"[^\w\s/-]", " ", base)
        base = re.sub(r"\s+", " ", base).strip()
        abbrev_map = {
            "st": "saint",
            "mt": "mount",
            "ft": "fort",
        }
        tokens = [abbrev_map.get(token, token) for token in base.split()]
        return " ".join(tokens).strip()

    @classmethod
    def _entity_aliases(cls, name: str) -> List[str]:
        normalized = cls._normalize_entity_name(name)
        if not normalized:
            return []
        parts = re.split(r"\s*/\s*|\s+or\s+|\s+\|\s+", normalized)
        aliases = []
        for part in parts:
            part = part.strip()
            if part:
                aliases.append(part)
        if normalized not in aliases:
            aliases.append(normalized)
        return aliases

    @staticmethod
    def _token_set(name: str) -> set[str]:
        return {token for token in re.split(r"\s+", name) if token}

    @classmethod
    def _initials_match(cls, left_norm: str, right_norm: str) -> bool:
        left_tokens = left_norm.split()
        right_tokens = right_norm.split()
        if not left_tokens or not right_tokens:
            return False
        if left_tokens[-1] != right_tokens[-1]:
            return False
        left_initials = {t[0] for t in left_tokens[:-1] if len(t) == 1}
        right_initials = {t[0] for t in right_tokens[:-1] if len(t) == 1}
        if left_initials and right_tokens[0] and right_tokens[0][0] in left_initials:
            return True
        if right_initials and left_tokens[0] and left_tokens[0][0] in right_initials:
            return True
        return False

    @classmethod
    def _names_match(cls, left: str, right: str) -> bool:
        left_norm = cls._normalize_entity_name(left)
        right_norm = cls._normalize_entity_name(right)
        if not left_norm or not right_norm:
            return False
        if left_norm == right_norm:
            return True

        left_tokens = cls._token_set(left_norm)
        right_tokens = cls._token_set(right_norm)
        if left_tokens & right_tokens:
            if left_norm in right_norm or right_norm in left_norm:
                return True
            if len(left_tokens) > 1 and len(right_tokens) > 1:
                if left_tokens.intersection(right_tokens):
                    return True

        if cls._initials_match(left_norm, right_norm):
            return True

        similarity = difflib.SequenceMatcher(None, left_norm, right_norm).ratio()
        return similarity >= 0.85

    @classmethod
    def _merge_aliases(cls, entity: Entity, incoming: List[str], primary: str) -> None:
        normalized_existing = {
            cls._normalize_entity_name(entity.name)
        }
        for alias in entity.aliases:
            normalized_existing.add(cls._normalize_entity_name(alias))
        for alias in incoming:
            clean = (alias or "").strip()
            if not clean:
                continue
            normalized = cls._normalize_entity_name(clean)
            if not normalized or normalized in normalized_existing:
                continue
            if cls._normalize_entity_name(primary) == normalized:
                continue
            entity.aliases.append(clean)
            normalized_existing.add(normalized)

    def _build_match_aliases(self, name: str, aliases: Optional[List[str]] = None) -> List[str]:
        incoming_match_aliases = self._entity_aliases(name)
        for alias in (aliases or []):
            incoming_match_aliases.extend(self._entity_aliases(alias))
        return incoming_match_aliases

    def find_entity_match(
        self,
        name: str,
        category: str,
        aliases: Optional[List[str]] = None,
    ) -> Optional[Entity]:
        clean_name = (name or "").strip()
        if not clean_name:
            return None

        normalized_category = self._normalize_category(category)
        incoming_match_aliases = self._build_match_aliases(clean_name, aliases)

        for ent in self.world_db.values():
            if self._normalize_category(ent.category) != normalized_category:
                continue
            candidates = [ent.name] + (ent.aliases or [])
            matched = any(
                self._names_match(candidate, incoming)
                for candidate in candidates
                for incoming in incoming_match_aliases
            )
            if matched:
                return ent
        return None

    def upsert_entity(
        self,
        name: str,
        category: str,
        desc: str = "",
        aliases: Optional[List[str]] = None,
        allow_merge: bool = True,
        allow_alias: bool = True,
    ) -> tuple[Optional[Entity], str]:
        clean_name = (name or "").strip()
        if not clean_name:
            return None, "skipped"

        normalized_category = self._normalize_category(category)
        incoming_aliases = [a for a in (aliases or []) if (a or "").strip()]
        incoming_aliases = list(dict.fromkeys([a.strip() for a in incoming_aliases if a.strip()]))
        incoming_match_aliases = self._build_match_aliases(clean_name, incoming_aliases)

        for ent in self.world_db.values():
            if self._normalize_category(ent.category) != normalized_category:
                continue

            candidates = [ent.name] + (ent.aliases or [])
            matched = any(
                self._names_match(candidate, incoming)
                for candidate in candidates
                for incoming in incoming_match_aliases
            )
            if matched:
                if allow_alias:
                    self._merge_aliases(ent, [clean_name] + incoming_aliases, ent.name)
                if allow_merge:
                    ent.merge(desc)
                self.last_modified = time.time()
                return ent, "matched"

        e = Entity(
            id=str(uuid.uuid4()),
            name=clean_name,
            category=normalized_category,
            description=(desc or "").strip(),
            aliases=[],
        )
        if allow_alias:
            self._merge_aliases(e, incoming_aliases, clean_name)
        self.world_db[e.id] = e
        self.last_modified = time.time()
        return e, "created"

    def add_entity(self, name: str, category: str, desc: str = "") -> Optional[Entity]:
        """Smart Add with Fuzzy Matching."""
        ent, _ = self.upsert_entity(name, category, desc, allow_merge=True, allow_alias=True)
        return ent

    def delete_entity(self, eid: str):
        if eid in self.world_db:
            del self.world_db[eid]
            self.last_modified = time.time()

    def delete_chapter(self, cid: str):
        if cid in self.chapters:
            del self.chapters[cid]
            for new_idx, ch in enumerate(self.get_ordered_chapters(), start=1):
                ch.index = new_idx
            self.last_modified = time.time()

    def add_chapter(self, title: str = "Untitled", content: str = "") -> Chapter:
        existing = [c.index for c in self.chapters.values()]
        index = (max(existing) + 1) if existing else 1

        if (title == "Untitled" or title.startswith("Chapter")) and self.outline:
            pattern = re.compile(rf"Chapter {index}[:\s]+(.*?)(?=\n|$)", re.IGNORECASE)
            match = pattern.search(self.outline)
            if match:
                raw_text = match.group(1).strip()
                split_text = re.split(r" [-:] ", raw_text, 1)
                final_title = sanitize_chapter_title(split_text[0].strip())
                if len(final_title) > 2:
                    title = final_title

        c = Chapter(
            id=str(uuid.uuid4()),
            index=index,
            title=sanitize_chapter_title(title),
            content=content or "",
            target_words=self.default_word_count,
        )
        c.word_count = len((content or "").split())
        self.chapters[c.id] = c
        self.last_modified = time.time()
        return c

    def get_ordered_chapters(self) -> List[Chapter]:
        return sorted(self.chapters.values(), key=lambda c: c.index)

    def get_total_word_count(self) -> int:
        return sum(c.word_count for c in self.chapters.values())

    def import_text_file(self, full_text: str) -> int:
        split_pattern = r"(?i)\n\s*(?:chapter|part)\s+(?:\d+|[a-z]+)[:.]?\s*.*?(?=\n)"
        parts = re.split(split_pattern, full_text or "")
        headers = re.findall(split_pattern, full_text or "")
        if len(parts) < 2:
            self.add_chapter("Imported Text", full_text or "")
            return 1

        count = 0
        if parts[0].strip():
            self.add_chapter("Prologue", parts[0].strip())
            count += 1
        for i, content in enumerate(parts[1:]):
            if not content.strip():
                continue
            title = headers[i].strip() if i < len(headers) else f"Chapter {i+1}"
            title = sanitize_chapter_title(title)
            self.add_chapter(title, content.strip())
            count += 1
        return count

    def to_dict(self) -> Dict[str, Any]:
        d = asdict(self)
        d.pop("filepath", None)
        d.pop("storage_dir", None)
        d["world_db"] = {k: asdict(v) for k, v in self.world_db.items()}
        d["chapters"] = {k: asdict(v) for k, v in self.chapters.items()}
        return d

    def save(self) -> str:
        safe_title = re.sub(r'[<>:"/\\|?*]', "_", self.title)[:60]
        filename = f"{self.id}_{safe_title.replace(' ', '_')}.json"
        storage_dir = self.storage_dir or AppConfig.PROJECTS_DIR
        try:
            os.makedirs(storage_dir, exist_ok=True)
        except OSError:
            logger.error("Cannot create storage directory %s", storage_dir, exc_info=True)
            return ""
        path = os.path.join(storage_dir, filename)
        lock_path = path + ".lock"
        tmp = path + ".tmp"
        last_error: Optional[Exception] = None
        if not _acquire_lock(
            lock_path,
            AppConfig.SAVE_LOCK_TIMEOUT,
            AppConfig.SAVE_LOCK_RETRY_SLEEP,
        ):
            logger.error("Save lock timeout for %s", path)
            return ""
        try:
            saved = False
            for attempt in range(1, 4):
                try:
                    if os.path.exists(path):
                        os.makedirs(AppConfig.BACKUPS_DIR, exist_ok=True)
                        stamp = time.strftime("%Y%m%d_%H%M%S")
                        backup_name = f"{stamp}__{os.path.basename(path)}"
                        backup_path = os.path.join(AppConfig.BACKUPS_DIR, backup_name)
                        try:
                            shutil.copy2(path, backup_path)
                        except Exception:
                            logger.warning("Backup failed for %s", path, exc_info=True)
                    with open(tmp, "w", encoding="utf-8") as f:
                        json.dump(self.to_dict(), f, indent=2, ensure_ascii=False)
                    os.replace(tmp, path)
                    saved = True
                    break
                except Exception as exc:
                    last_error = exc
                    logger.warning("Save attempt %s failed for %s", attempt, path, exc_info=True)
                    time.sleep(0.1)
            else:
                logger.error("Save failed after retries for %s: %s", path, last_error)
        finally:
            _release_lock(lock_path)
            # Clean up stale temp file on failure
            if not saved:
                try:
                    os.remove(tmp)
                except OSError:
                    pass

        if not saved:
            return ""

        self.filepath = path
        self.storage_dir = storage_dir
        self.last_modified = time.time()
        return path

    @classmethod
    def load(cls, path: str) -> "Project":
        try:
            try:
                data = _load_json_with_encoding_fallback(path)
            except json.JSONDecodeError as exc:
                logger.error("Corrupt project file %s: %s", path, exc)
                raise ValueError(
                    f"Cannot load project: file is not valid JSON ({path})"
                ) from exc
            except UnicodeDecodeError as exc:
                logger.error("Project file has unsupported encoding %s: %s", path, exc)
                raise ValueError(
                    f"Cannot load project: unsupported file encoding ({path})"
                ) from exc
        except FileNotFoundError as exc:
            logger.error("Project file not found: %s", path, exc_info=True)
            raise ValueError(f"Cannot load project: file not found ({path})") from exc
        if not isinstance(data, dict):
            raise ValueError(f"Cannot load project: expected JSON object ({path})")

        proj = cls(
            id=data.get("id") or str(uuid.uuid4()),
            title=data.get("title") or "Untitled Project",
        )
        proj.author = data.get("author", "")
        proj.genre = data.get("genre", "")
        proj.outline = data.get("outline", "")
        proj.memory = data.get("memory", data.get("memory_notes", ""))
        proj.memory_hard = data.get("memory_hard", "")
        proj.memory_soft = data.get("memory_soft", "")
        proj.author_note = data.get("author_note", data.get("authors_note", ""))
        proj.style_guide = data.get("style_guide", data.get("style_note", ""))
        proj.default_word_count = data.get("default_word_count", 1000)
        proj.created_at = data.get("created_at", time.time())
        proj.last_modified = data.get("last_modified", time.time())

        ent_fields = {f.name for f in fields(Entity)}
        world_db_data = data.get("world_db") or data.get("characters") or {}
        if isinstance(world_db_data, list):
            world_db_data = {
                (item.get("id") or str(uuid.uuid4())): item
                for item in world_db_data
                if isinstance(item, dict)
            }
        if not isinstance(world_db_data, dict):
            world_db_data = {}
        for k, v in world_db_data.items():
            if "category" not in v:
                v["category"] = "Character"
            clean = {key: val for key, val in v.items() if key in ent_fields}
            clean["id"] = clean.get("id") or k or str(uuid.uuid4())
            clean["name"] = clean.get("name") or "Unknown"
            clean["category"] = Project._normalize_category(clean.get("category", "Lore"))
            if "aliases" in clean and not isinstance(clean["aliases"], list):
                clean["aliases"] = [str(clean["aliases"])]
            try:
                proj.world_db[clean["id"]] = Entity(**clean)
            except TypeError:
                logger.warning("Skipping malformed entity entry: %s", k, exc_info=True)

        chap_fields = {f.name for f in fields(Chapter)}
        chapter_data = data.get("chapters") or {}
        if isinstance(chapter_data, list):
            chapter_data = {
                (item.get("id") or str(uuid.uuid4())): item
                for item in chapter_data
                if isinstance(item, dict)
            }
        if not isinstance(chapter_data, dict):
            chapter_data = {}
        fallback_index = 1
        for k, v in chapter_data.items():
            clean = {key: val for key, val in v.items() if key in chap_fields}
            clean["id"] = clean.get("id") or k or str(uuid.uuid4())
            try:
                clean_index = int(clean.get("index") or fallback_index)
            except (TypeError, ValueError):
                clean_index = fallback_index
            clean["index"] = clean_index
            clean["title"] = clean.get("title") or f"Chapter {clean_index}"
            fallback_index = max(fallback_index, clean_index + 1)
            try:
                proj.chapters[clean["id"]] = Chapter(**clean)
            except TypeError:
                logger.warning("Skipping malformed chapter entry: %s", k, exc_info=True)

        proj.storage_dir = os.path.dirname(path)
        proj.filepath = path
        return proj

    @classmethod
    def delete_file(cls, filepath: str) -> tuple[bool, str]:
        if not filepath:
            return False, "Missing project file path."
        if not os.path.exists(filepath):
            return True, "Project file already removed."

        last_error: Optional[Exception] = None
        for _ in range(3):
            try:
                # Clear read-only bit if needed (common on Windows copies).
                os.chmod(filepath, 0o666)
            except OSError:
                pass
            try:
                os.remove(filepath)
                lock_path = filepath + ".lock"
                tmp_path = filepath + ".tmp"
                for extra in (lock_path, tmp_path):
                    try:
                        if os.path.exists(extra):
                            os.remove(extra)
                    except OSError:
                        logger.debug("Could not remove sidecar file during delete: %s", extra, exc_info=True)
                return True, "Deleted."
            except OSError as exc:
                last_error = exc
                time.sleep(0.1)

        logger.warning("Failed to delete project file: %s", filepath, exc_info=True)
        return False, str(last_error) if last_error else "Unknown delete error."


# ============================================================
# 3) AI ENGINE
# ============================================================

class AIEngine:
    def __init__(
        self,
        timeout: int = AppConfig.GROQ_TIMEOUT,
        base_url: Optional[str] = None,
        provider: Optional[str] = None,
        api_key: Optional[str] = None,
    ):
        self.provider = _normalize_provider(provider or get_active_provider())
        self.base_url = (base_url or _get_provider_base_url(self.provider)).rstrip("/")
        self.timeout = timeout
        self.session = requests.Session()
        self.api_key = (api_key or "").strip()

    def _resolve_api_key(self) -> str:
        if self.api_key:
            return self.api_key
        key, _ = get_effective_key(self.provider)
        return key

    def probe_models(self, api_key: str) -> List[str]:
        headers = {"Authorization": f"Bearer {api_key}"} if api_key else {}
        try:
            r = self.session.get(f"{self.base_url}/models", headers=headers, timeout=5)
            r.raise_for_status()
            data = r.json()
            return [m.get("id") for m in data.get("data", []) if m.get("id")]
        except Exception:
            logger.warning("Model probe failed for %s", self.base_url, exc_info=True)
            return []

    def generate_stream(self, prompt: str, model: str) -> Generator[str, None, None]:
        if "streamlit" in sys.modules:
            import streamlit as st

            results = st.session_state.get("coherence_results", [])
            if len(results) > 2:
                yield (
                    "ERROR: Canon violation detected.\n"
                    "Resolve Hard Canon conflicts before generating AI content."
                )
                return
            project = st.session_state.get("project")
            hard_rules = ""
            if project:
                hard_rules = (project.memory_hard or project.memory or "").strip()
            if hard_rules:
                prompt = (
                    "HARD CANON RULES (NON-NEGOTIABLE):\n"
                    f"{hard_rules}\n\n"
                    f"{prompt}"
                )
        if not model:
            yield f"{_provider_label(self.provider)} model not configured."
            return

        api_key = self._resolve_api_key()
        if not api_key:
            yield f"{_provider_label(self.provider)} API key not configured."
            return

        prompt = sanitize_ai_input(prompt)
        prompt = _truncate_prompt(prompt, AppConfig.MAX_PROMPT_CHARS)
        headers = {"Content-Type": "application/json"}
        headers["Authorization"] = f"Bearer {api_key}"

        def _groq_non_stream() -> Generator[str, None, None]:
            payload = {
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
            }
            try:
                r = self.session.post(
                    f"{self.base_url}/chat/completions",
                    json=payload,
                    headers=headers,
                    timeout=self.timeout,
                )
                r.raise_for_status()
                data = r.json()
                choice = (data.get("choices") or [{}])[0]
                content = (choice.get("message") or {}).get("content") or choice.get("text") or ""
                if content:
                    yield content
                else:
                    yield f"{_provider_label(self.provider)} response empty."
            except requests.exceptions.Timeout:
                logger.warning("%s request timed out after %ss", _provider_label(self.provider), self.timeout)
                yield f"{_provider_label(self.provider)} request timed out. Try again or increase timeout."
            except requests.exceptions.HTTPError as exc:
                status = getattr(exc.response, "status_code", None)
                logger.warning("%s HTTP %s error", _provider_label(self.provider), status, exc_info=True)
                if status == 401:
                    yield f"{_provider_label(self.provider)} API key is invalid or expired."
                elif status == 429:
                    yield f"{_provider_label(self.provider)} rate limit exceeded. Please wait and try again."
                else:
                    yield f"{_provider_label(self.provider)} generation failed (HTTP {status})."
            except Exception:
                logger.warning("%s non-stream generation failed", _provider_label(self.provider), exc_info=True)
                yield f"{_provider_label(self.provider)} generation failed."

        payload = {
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "stream": True,
        }
        try:
            with self.session.post(
                f"{self.base_url}/chat/completions",
                json=payload,
                headers=headers,
                stream=True,
                timeout=self.timeout,
            ) as r:
                r.raise_for_status()
                yielded = False
                for raw in r.iter_lines():
                    if not raw:
                        continue
                    line = raw.decode("utf-8").strip()
                    if not line.startswith("data:"):
                        continue
                    data = line.replace("data:", "", 1).strip()
                    if data == "[DONE]":
                        break
                    try:
                        chunk = json.loads(data)
                    except json.JSONDecodeError:
                        continue
                    delta = chunk.get("choices", [{}])[0].get("delta", {})
                    content = delta.get("content")
                    if content:
                        yielded = True
                        yield content
                if not yielded:
                    yield from _groq_non_stream()
        except Exception:
            logger.warning(
                "%s streaming generation failed, retrying non-stream",
                _provider_label(self.provider),
                exc_info=True,
            )
            yield from _groq_non_stream()

    def generate(self, prompt: str, model: str) -> Dict[str, str]:
        chunks: List[str] = []
        for chunk in self.generate_stream(prompt, model):
            chunks.append(chunk)
        return {"text": "".join(chunks)}

    def generate_json(self, prompt: str, model: str) -> Optional[List[Dict[str, Any]]]:
        res = self.generate(f"Return ONLY valid JSON. List of objects. No markdown.\n\n{prompt}", model)
        txt = (res.get("text", "") or "").strip()
        if not txt:
            return None
        txt = re.sub(r"```json\s*", "", txt)
        txt = re.sub(r"```\s*", "", txt).strip()
        # Try to extract JSON array or object even if surrounded by text
        if not txt.startswith(("[", "{")):
            bracket = txt.find("[")
            brace = txt.find("{")
            if bracket >= 0 and (brace < 0 or bracket < brace):
                txt = txt[bracket:]
            elif brace >= 0:
                txt = txt[brace:]
        try:
            d = json.loads(txt)
            if isinstance(d, list):
                return d
            if isinstance(d, dict):
                return [d]
            return None
        except (json.JSONDecodeError, ValueError, TypeError):
            logger.warning("JSON parse failed for AI response (length=%d)", len(txt))
            return None


class AnalysisEngine:
    @staticmethod
    def extract_entities(text: str, model: str) -> list:
        if not text or len(text) < 50:
            return []
        prompt = (
            "Analyze text. Identify Characters, Locations, Factions, and important Lore.\n"
            "Return JSON List: [{'name': 'X', 'category': 'Character|Location|Faction|Lore', "
            "'description': 'Z', 'aliases': ['Alt Name'], 'confidence': 0.0}]\n"
            f"TEXT:\n{text[:6000]}"
        )
        return AIEngine().generate_json(prompt, model) or []

    @staticmethod
    def generate_title(outline: str, genre: str, model: str) -> str:
        prompt = (
            f"GENRE: {genre}\n"
            f"OUTLINE: {outline[:2000]}\n"
            "TASK: Create a creative, catchy title for this story.\n"
            "RULES: Output ONLY the title. No quotes. No prefixes like 'Title:'."
        )
        raw = (AIEngine().generate(prompt, model).get("text", "Untitled") or "").strip()
        clean = re.sub(r"(?i)^(here is a title|sure|suggested title|title)[:\s-]*", "", raw).strip()
        clean = clean.replace('"', "").replace("'", "").strip()
        return clean.split("\n")[0] if clean else "Untitled"


    @staticmethod
    def detect_genre(outline: str, model: str) -> str:
        """
        Attempts to infer a concise genre label from the outline text.
        Uses AI if available, otherwise falls back to a lightweight heuristic.
        """
        outline = (outline or "").strip()
        if not outline:
            return ""

        # Heuristic fallback first (fast, offline-safe)
        low = outline.lower()
        heuristic_map = [
            ("space", "Science Fiction"),
            ("spaceship", "Science Fiction"),
            ("alien", "Science Fiction"),
            ("cyber", "Cyberpunk"),
            ("hacker", "Cyberpunk"),
            ("detective", "Mystery/Thriller"),
            ("murder", "Mystery/Thriller"),
            ("serial killer", "Mystery/Thriller"),
            ("dragon", "Fantasy"),
            ("magic", "Fantasy"),
            ("kingdom", "Fantasy"),
            ("vampire", "Horror"),
            ("zombie", "Horror"),
            ("haunted", "Horror"),
            ("romance", "Romance"),
            ("love triangle", "Romance"),
            ("coming of age", "Coming-of-Age"),
            ("war", "Historical/War"),
            ("wwii", "Historical/War"),
            ("post-apocalyptic", "Post-Apocalyptic"),
            ("apocalypse", "Post-Apocalyptic"),
        ]
        for needle, genre in heuristic_map:
            if needle in low:
                return genre

        # AI-based inference
        prompt = (
            "You are a book editor. Infer the most fitting genre from the outline.\n"
            "Rules:\n"
            "- Output ONLY a short genre label (2-4 words).\n"
            "- No quotes, no prefix, no punctuation at the end.\n\n"
            f"OUTLINE:\n{outline[:2500]}"
        )
        try:
            raw = (AIEngine().generate(prompt, model).get("text", "") or "").strip()
            raw = re.sub(r'(?i)^(genre)[:\s-]*', '', raw).strip()
            raw = raw.replace('"', "").replace("'", "").strip()
            genre = raw.splitlines()[0].strip()
            # Keep it clean / short
            genre = re.sub(r"[^\w\s/&-]", "", genre).strip()
            if len(genre) > 40:
                genre = genre[:40].strip()
            return genre
        except Exception:
            logger.warning("Genre detection failed", exc_info=True)
            return ""

    @staticmethod
    def enrich_entity(name: str, category: str, desc: str, model: str) -> str:
        prompt = (
            f"NAME: {name}\nCATEGORY: {category}\nCURRENT INFO: {desc}\n"
            "TASK: Expand on this entry with creative details, backstory, or physical description.\n"
            "RULES: Keep it concise (under 50 words). Bullet points are okay."
        )
        return AIEngine().generate(prompt, model).get("text", "") or ""

    @staticmethod
    def coherence_check(
        memory: str,
        author_note: str,
        style_guide: str,
        outline: str,
        world_bible: str,
        chapters: List[Dict[str, Any]],
        model: str,
    ) -> List[Dict[str, Any]]:
        chapter_payload = chapters or [
            {
                "chapter_index": 0,
                "summary": "No chapter summaries yet.",
                "excerpt": (outline or world_bible or memory or "")[:1200],
            }
        ]
        prompt = (
            "You are a developmental editor checking continuity and canon coherence.\n"
            "Identify contradictions, timeline issues, or inconsistent details.\n"
            "Return ONLY JSON (no markdown) as a list of objects with keys:\n"
            "- chapter_index (number)\n"
            "- issue (short description)\n"
            "- target_excerpt (exact text to replace, must exist in the chapter)\n"
            "- suggested_rewrite (replacement text to fix coherence)\n"
            "- confidence (low|medium|high)\n\n"
            f"PROJECT MEMORY:\n{(memory or '')[:2000]}\n\n"
            f"AUTHOR NOTE:\n{(author_note or '')[:1200]}\n\n"
            f"STYLE GUIDE:\n{(style_guide or '')[:1200]}\n\n"
            f"WORLD BIBLE:\n{(world_bible or '')[:2400]}\n\n"
            f"OUTLINE:\n{(outline or '')[:2000]}\n\n"
            "CHAPTERS (summaries + excerpts):\n"
            f"{json.dumps(chapter_payload, ensure_ascii=False)}"
        )
        return AIEngine().generate_json(prompt, model) or []


class StoryEngine:
    @staticmethod
    def summarize(text: str, model: str) -> str:
        if not text:
            return ""
        prompt = f"Summarize this scene concisely (under 100 words):\n\n{text[:4000]}"
        return AIEngine().generate(prompt, model).get("text", "") or ""

    @staticmethod
    def generate_chapter_prompt(project: Project, chapter_index: int, target_words: int) -> str:
        hard_rules = (project.memory_hard or project.memory or "").strip()[:1500]
        hard_block = f"HARD CANON RULES (STRICT):\n{hard_rules}\n\n" if hard_rules else ""
        soft_guidelines = (project.memory_soft or "").strip()[:1500]
        soft_block = f"SOFT GUIDELINES (STYLE/CONTEXT):\n{soft_guidelines}\n\n" if soft_guidelines else ""
        memory_context = (project.memory or "").strip()[:1500]
        memory_block = f"PROJECT MEMORY:\n{memory_context}\n\n" if memory_context else ""
        author_note = (project.author_note or "").strip()[:1200]
        author_block = f"AUTHOR NOTE:\n{author_note}\n\n" if author_note else ""
        style_guide = (project.style_guide or "").strip()[:1200]
        style_block = f"STYLE GUIDE:\n{style_guide}\n\n" if style_guide else ""
        outline_context = (project.outline or "")[:3000]
        match = re.search(rf"(?i)Chapter {chapter_index}[:\s]+(.*?)(?=\n|$)", project.outline or "")
        if match:
            specific_beat = match.group(1).strip()
            outline_context += f"\n\nCURRENT CHAPTER OBJECTIVE: {specific_beat}"

        prev_chaps = [c for c in project.get_ordered_chapters() if c.index < chapter_index]
        story_so_far = ""
        prev_text = ""

        if prev_chaps:
            story_so_far = "PREVIOUS EVENTS:\n"
            for c in prev_chaps[-5:]:
                summ = c.summary if c.summary else "No summary."
                story_so_far += f"Ch {c.index}: {summ}\n"
            prev_text = f"\nIMMEDIATELY PRECEDING SCENE:\n{(prev_chaps[-1].content or '')[-1500:]}\n"

        prompt = (
            f"TITLE: {project.title}\nGENRE: {project.genre}\n"
            f"{hard_block}"
            f"{soft_block}"
            f"{memory_block}"
            f"{author_block}"
            f"{style_block}"
            f"OUTLINE CONTEXT:\n{outline_context}\n\n"
            f"{story_so_far}"
            f"{prev_text}\n"
            f"TASK: Write Chapter {chapter_index}.\n"
            f"LENGTH GOAL: {target_words} words.\n"
            "INSTRUCTIONS:\n"
            "1. Continue directly from the preceding scene (if any).\n"
            "2. Match the writing style and tone of the previous text.\n"
            "3. Focus on the 'Current Chapter Objective' if provided.\n"
            "4. Do not output the chapter title. Just the story prose."
        )
        return prompt

    @staticmethod
    def reverse_engineer_outline(project: Project, model: str) -> str:
        chaps = project.get_ordered_chapters()
        if not chaps:
            return "No content."
        txt = ""
        for c in chaps:
            txt += f"Ch {c.index} ({c.title}): {c.summary or (c.content or '')[:300]}\n"
        prompt = f"Create a structured outline based on these chapters:\n\n{txt}"
        return AIEngine().generate(prompt, model).get("text", "") or ""


# ============================================================
# 4) UTILS
# ============================================================

REWRITE_PRESETS = {
    "Custom": "Use your own custom instructions",
    "Fix Grammar": "Correct grammar and spelling only.",
    "Improve Flow": "Fix awkward phrasing and sentence rhythm.",
    "Sensory Expansion": "Show, Don't Tell - Add smells, sounds, and textures.",
    "Deepen POV": "Add more internal monologue and emotional reaction.",
    "Make Witty": "Add humor, banter, and irony.",
    "Make Darker": "Grim, moody, noir-style descriptions.",
    "Expand": "Write more details, double the length.",
}

def rewrite_prompt(text: str, preset: str, custom: str = "") -> str:
    text = sanitize_ai_input(text, AppConfig.MAX_PROMPT_CHARS)
    custom = sanitize_ai_input(custom, 2000)
    instr = custom if preset == "Custom" else REWRITE_PRESETS.get(preset, "")
    return f"Act as an editor.\nTASK: {preset}\nDETAILS: {instr}\n\nINPUT TEXT:\n{text}"

def project_to_markdown(project: Project) -> str:
    md = [f"# {project.title}", f"**Genre:** {project.genre}\n"]
    if project.outline:
        md.append("## Outline")
        md.append(project.outline + "\n")
    if project.memory:
        md.append("## Memory")
        md.append(project.memory + "\n")
    if project.author_note:
        md.append("## Author Note")
        md.append(project.author_note + "\n")
    if project.style_guide:
        md.append("## Style Guide")
        md.append(project.style_guide + "\n")
    md.append("## World Bible")
    for c in project.world_db.values():
        md.append(f"- **{c.name}** ({c.category}): {c.description}")
    md.append("\n## Chapters")
    for c in project.get_ordered_chapters():
        md.append(f"### {c.index}. {c.title}")
        md.append((c.content or "") + "\n")
    return "\n".join(md)


def project_to_text(project: Project) -> str:
    lines = [project.title, f"Genre: {project.genre}", ""]
    if project.outline:
        lines.extend(["OUTLINE", project.outline, ""])
    if project.world_db:
        lines.append("WORLD BIBLE")
        for e in sorted(project.world_db.values(), key=lambda x: (x.category, x.name.lower())):
            lines.append(f"- {e.name} ({e.category}): {e.description}")
        lines.append("")
    for c in project.get_ordered_chapters():
        lines.append(f"CHAPTER {c.index}: {c.title}")
        lines.append(c.content or "")
        lines.append("")
    return "\n".join(lines).strip()


def project_to_json(project: Project) -> str:
    payload = project.to_dict()
    return json.dumps(payload, ensure_ascii=False, indent=2)


def project_to_docx_bytes(project: Project) -> bytes:
    from io import BytesIO
    import docx  # type: ignore

    doc = docx.Document()
    doc.add_heading(project.title, level=1)
    if project.genre:
        doc.add_paragraph(f"Genre: {project.genre}")
    if project.outline:
        doc.add_heading("Outline", level=2)
        doc.add_paragraph(project.outline)
    if project.world_db:
        doc.add_heading("World Bible", level=2)
        for e in sorted(project.world_db.values(), key=lambda x: (x.category, x.name.lower())):
            doc.add_paragraph(f"{e.name} ({e.category}): {e.description}")
    if project.chapters:
        doc.add_heading("Chapters", level=2)
        for c in project.get_ordered_chapters():
            doc.add_heading(f"Chapter {c.index}: {c.title}", level=3)
            doc.add_paragraph(c.content or "")
    buff = BytesIO()
    doc.save(buff)
    return buff.getvalue()


def sanitize_chapter_title(title: str) -> str:
    raw = (title or "").strip()
    if not raw:
        return ""
    clean = raw.replace(""", "").replace(""", "").replace("'", "")
    clean = clean.replace('"', "").replace("'", "")
    clean = re.sub(r"^\*+|\*+$", "", clean).strip()
    clean = re.sub(r"\s{2,}", " ", clean).strip()
    return clean


# ============================================================
# 5) STREAMLIT UI (Appearance + First-time Onboarding)
# ============================================================



def _run_ui():
    import streamlit as st
    import streamlit.components.v1 as components
    from contextlib import contextmanager
    from pathlib import Path
    from app.components.ui import action_card, card, primary_button, section_header, stat_tile
    from app.ui.components import card_block, cta_tile, empty_state, header_bar, render_tag_list
    from app.ui.ui_layout import render_card, render_metric, render_section_header
    from app.layout.layout import render_footer
    from app.utils.keys import ui_key
    # Import enhanced UI feedback components
    from app.ui.feedback import (
        loading_state,
        step_indicator,
        progress_bar_with_message,
        feedback_message,
    )
    from app.ui.navigation import help_tooltip, quick_action_card

    widget_counters: Dict[tuple, int] = {}
    key_prefix_stack: List[str] = []

    ASSETS_DIR = Path(__file__).resolve().parents[1] / "assets"

    @st.cache_data(show_spinner=False)
    def load_asset_bytes(filename: str) -> Optional[bytes]:
        try:
            return read_asset_bytes(ASSETS_DIR, filename)
        except Exception:
            logging.getLogger("MANTIS").warning(
                "Failed to load asset %s",
                filename,
                exc_info=True,
            )
            return None

    def asset_base64(filename: str) -> str:
        payload = load_asset_bytes(filename)
        if not payload:
            return ""
        return base64.b64encode(payload).decode("utf-8")

    def load_first_asset_bytes(*filenames: str) -> Optional[bytes]:
        for filename in filenames:
            payload = load_asset_bytes(filename)
            if payload:
                return payload
        return None

    def resolve_first_asset_path(*filenames: str) -> Optional[Path]:
        for filename in filenames:
            path = resolve_asset_path(ASSETS_DIR, filename)
            if path and path.exists():
                return path
        return None

    def _current_prefix() -> str:
        return "__".join(key_prefix_stack) or "global"

    def _slugify(value: str) -> str:
        slug = re.sub(r"[^a-z0-9]+", "_", (value or "").lower()).strip("_")
        return slug[:40] or "widget"

    def _auto_key(widget_type: str, label: Optional[str], key: Optional[str]) -> str:
        if key:
            return key
        prefix = _current_prefix()
        slug = _slugify(label or widget_type)
        counter_key = (prefix, widget_type, slug)
        widget_counters[counter_key] = widget_counters.get(counter_key, 0) + 1
        index = widget_counters[counter_key]
        return ui_key(prefix, widget_type, slug, str(index))

    def init_state(key: str, default: Any) -> None:
        if key not in st.session_state:
            st.session_state[key] = default

    def queue_widget_update(key: str, value: Any) -> None:
        pending = st.session_state.setdefault("_pending_widget_updates", {})
        pending[key] = value

    def apply_pending_widget_updates() -> None:
        pending = st.session_state.pop("_pending_widget_updates", None)
        if pending:
            for pending_key, pending_value in pending.items():
                st.session_state[pending_key] = pending_value

    def normalize_theme_name(value: Any) -> str:
        return "Light" if str(value or "").strip().lower() == "light" else "Dark"

    def _record_action(action_label: Optional[str], action_key: Optional[str]) -> None:
        label = (action_label or action_key or "action").strip()
        st.session_state["last_action"] = label
        st.session_state["last_action_ts"] = time.time()

    @contextmanager
    def key_scope(prefix: str) -> Generator[None, None, None]:
        key_prefix_stack.append(prefix)
        try:
            yield
        finally:
            key_prefix_stack.pop()

    def _wrap_widget(widget_fn, widget_type: str):
        def _wrapped(label=None, *args, **kwargs):
            resolved_key = _auto_key(widget_type, label, kwargs.get("key"))
            kwargs["key"] = resolved_key
            result = widget_fn(label, *args, **kwargs)
            if widget_type in {"button", "form_submit_button"} and result:
                _record_action(label, resolved_key)
            return result

        _wrapped._mantis_wrapped = True
        return _wrapped

    def _wrap_widget_no_label(widget_fn, widget_type: str):
        def _wrapped(*args, **kwargs):
            resolved_key = _auto_key(widget_type, None, kwargs.get("key"))
            kwargs["key"] = resolved_key
            result = widget_fn(*args, **kwargs)
            if widget_type in {"button", "form_submit_button"} and result:
                _record_action(None, resolved_key)
            return result

        _wrapped._mantis_wrapped = True
        return _wrapped

    def _maybe_wrap(name: str, widget_type: str, has_label: bool = True) -> None:
        if not hasattr(st, name):
            return
        current = getattr(st, name)
        # Skip if already wrapped to avoid stacking wrappers on repeated
        # Streamlit reruns, which would eventually cause a RecursionError.
        if getattr(current, "_mantis_wrapped", False):
            return
        wrapped = _wrap_widget(current, widget_type) if has_label else _wrap_widget_no_label(current, widget_type)
        setattr(st, name, wrapped)

    _maybe_wrap("button", "button")
    _maybe_wrap("text_input", "text_input")
    _maybe_wrap("text_area", "text_area")
    _maybe_wrap("selectbox", "selectbox")
    _maybe_wrap("radio", "radio")
    _maybe_wrap("checkbox", "checkbox")
    _maybe_wrap("number_input", "number_input")
    _maybe_wrap("slider", "slider")
    _maybe_wrap("multiselect", "multiselect")
    _maybe_wrap("file_uploader", "file_uploader")
    _maybe_wrap("download_button", "download_button")
    _maybe_wrap("form_submit_button", "form_submit_button")
    _maybe_wrap("toggle", "toggle")
    _maybe_wrap("feedback", "feedback")
    _maybe_wrap("date_input", "date_input")
    _maybe_wrap("time_input", "time_input")
    _maybe_wrap("color_picker", "color_picker")
    _maybe_wrap("camera_input", "camera_input")
    _maybe_wrap("audio_input", "audio_input")
    _maybe_wrap("chat_input", "chat_input", has_label=False)
    # NOTE: st.page_link does not take a "label" as its first positional argument.
    # Wrapping it with the generic label-first wrapper can break routing.
    # If we need auto-keys for page_link later, add a signature-aware wrapper.
    # _maybe_wrap("page_link", "page_link")

    def get_canon_health() -> tuple[str, str]:
        results = st.session_state.get("coherence_results", [])
        issue_count = len(results)
        if issue_count == 0:
            return "OK", "Canon Stable"
        if issue_count <= 2:
            return "WARN", "Minor Canon Drift"
        return "RISK", "High Canon Risk"

    def detect_hard_canon_violation(project: Project, chapter_index: int, new_text: str) -> List[Dict[str, Any]]:
        hard_rules = (project.memory_hard or project.memory or "").strip()
        if not hard_rules or not (new_text or "").strip():
            return []
        active_provider = _normalize_provider(st.session_state.get("ai_provider", "groq"))
        active_key, _ = get_effective_key(active_provider, st.session_state.get("user_id"))
        if not active_key or not get_ai_model():
            return []
        compiled_world_bible = "\n".join(
            f"{e.name} ({e.category}): {e.description}"
            for e in project.world_db.values()
        )
        chapters_payload = [
            {
                "chapter_index": chapter_index,
                "summary": "",
                "excerpt": (new_text or "")[:800],
            }
        ]
        results = AnalysisEngine.coherence_check(
            memory=hard_rules,
            author_note="",
            style_guide="",
            outline=project.outline or "",
            world_bible=compiled_world_bible,
            chapters=chapters_payload,
            model=get_ai_model(),
        )
        return results or []

    def update_locked_chapters() -> None:
        """Refresh the set of chapter indices that are locked due to canon violations."""
        results = st.session_state.get("coherence_results", [])
        locked = set()
        for issue in results:
            try:
                idx = int(issue.get("chapter_index", -1))
                if idx >= 0:
                    locked.add(idx)
            except (TypeError, ValueError):
                pass
        st.session_state["locked_chapters"] = locked

    LEGAL_DIR = Path(__file__).resolve().parents[1] / "legal"

    def _read_legal_file(filename: str, fallback: str) -> str:
        path = LEGAL_DIR / filename
        if path.exists():
            return _read_text_with_encoding_fallback(path)
        return fallback

    def render_privacy():
        content = _read_legal_file("privacy.md", "## Privacy Policy\n\nLocal-only storage. No analytics.")
        st.markdown(content)
        if st.button("Back to All Policies", key=ui_key("privacy", "back")):
            st.session_state.page = "legal"
            st.rerun()

    def render_terms():
        content = _read_legal_file("terms.md", "## Terms of Service\n\nProvided as-is for creative use.")
        st.markdown(content)
        if st.button("Back to All Policies", key=ui_key("terms", "back")):
            st.session_state.page = "legal"
            st.rerun()

    def render_copyright():
        content = _read_legal_file("copyright.md", "## Copyright\n\n MANTIS Studio")
        st.markdown(content)
        if st.button("Back to All Policies", key=ui_key("copyright", "back")):
            st.session_state.page = "legal"
            st.rerun()

    def render_cookie():
        content = _read_legal_file("cookie.md", "## Cookie Policy\n\nEssential cookies only.")
        st.markdown(content)
        if st.button("Back to All Policies", key=ui_key("cookie", "back")):
            st.session_state.page = "legal"
            st.rerun()

    def render_help():
        content = _read_legal_file("help.md", "## Help\n\nVisit our GitHub for support.")
        st.markdown(content)
        if st.button("Back to Dashboard", key=ui_key("help", "back")):
            st.session_state.page = "home"
            st.rerun()

    logger.info("Starting UI initialization...")
    logger.debug(f"Assets directory: {ASSETS_DIR}")
    
    icon_path = resolve_first_asset_path(
        "branding/mantis_favicon.png",
        "mantis_favicon.png",
        "branding/mantis_emblem.png",
        "mantis_emblem.png",
    )
    page_icon = str(icon_path) if icon_path and icon_path.exists() else "M"
    logger.debug(f"Icon path exists: {bool(icon_path and icon_path.exists())}, using: {page_icon}")
    
    try:
        st.set_page_config(page_title=AppConfig.APP_NAME, page_icon=page_icon, layout="wide")
        logger.info("Page config set successfully")
    except Exception as e:
        logger.error(f"Failed to set page config: {e}", exc_info=True)
        raise
    
    # Theme injection is handled by the enhanced design system below.
    
    # Use centralized session initialization module
    from app.session_init import initialize_session_state
    try:
        initialize_session_state(st)
        logger.info("Session state initialized successfully via session_init module")
    except Exception as e:
        logger.error(f"Failed to initialize session state: {e}", exc_info=True)
        # Set absolute minimum state to prevent total failure
        st.session_state.setdefault("page", "home")
        st.session_state.setdefault("initialized", True)
        st.warning("Failed to initialize application state. Some features may not work correctly.")

    apply_pending_widget_updates()

    # Apply enhanced theme system from design system
    try:
        from app.ui.enhanced_theme import inject_enhanced_theme
        theme = "Light" if str(st.session_state.get("ui_theme", "Dark")).strip().lower() == "light" else "Dark"
        st.session_state.ui_theme = theme
        inject_enhanced_theme(theme, page_key=st.session_state.get("page", "home"))
        logger.info(f"Enhanced theme '{theme}' injected successfully")
    except Exception as e:
        logger.error(f"Failed to inject enhanced theme: {e}", exc_info=True)
        st.error("Theme failed to load. Please restart the app.")

    # Apply additional header/card styles that are specific to main.py layouts
    st.html(
        """
    <style>
    .mantis-header {{
        display:flex;
        align-items:center;
        justify-content: space-between;
        gap:14px;
        padding:18px 24px;
        border-radius:22px;
        background: var(--mantis-header-gradient);
        border: 1px solid var(--mantis-primary-border);
        margin-top: 18px;
        margin-bottom: 18px;
        box-shadow: var(--mantis-shadow-strong);
    }}
    .mantis-header-left {{
        display:flex;
        align-items:center;
        gap:14px;
    }}
    .mantis-header-right {{
        display:flex;
        gap:10px;
        align-items:center;
    }}
    .mantis-header-logo {
        width:236px;
        height:82px;
        border-radius:16px;
        background: var(--mantis-header-logo-bg);
        background: linear-gradient(135deg, color-mix(in srgb, var(--mantis-header-logo-bg) 82%, transparent), transparent);
        border: 1px solid var(--mantis-primary-border);
        border: 1px solid color-mix(in srgb, var(--mantis-accent) 28%, transparent);
        display:flex;
        align-items:center;
        justify-content:center;
        overflow:hidden;
        box-shadow:
            inset 0 0 0 1px rgba(0,0,0,0.04),
            0 8px 22px var(--mantis-accent-glow),
            0 8px 22px color-mix(in srgb, var(--mantis-accent-glow) 60%, transparent);
    }
    .mantis-header-logo img {
        width:100%;
        max-height:62px;
        object-fit:contain;
        padding:0;
        border-radius:0;
        filter: drop-shadow(0 0 10px color-mix(in srgb, var(--mantis-accent-glow) 75%, transparent));
    }
    .mantis-header-title {{
        font-size:22px;
        font-weight:800;
        color: var(--mantis-text);
        letter-spacing: 0.01em;
    }}
    .mantis-header-sub {{
        color: var(--mantis-header-sub);
        font-size:12px;
    }}
    .mantis-header-meta {{
        display:flex;
        flex-direction:column;
        align-items:flex-end;
        gap:4px;
        font-size:12px;
        color: var(--mantis-muted);
    }}
    </style>
    """
    )
    
    st.html(
        """
    <style>
    section[data-testid="stSidebar"] details summary {
        color: var(--mantis-text) !important;
        font-weight: 600;
        background: var(--mantis-surface-alt);
        border: 1px solid var(--mantis-card-border);
        border-radius: 12px;
        padding: 6px 10px;
    }
    section[data-testid="stSidebar"] details[open] {
        margin-bottom: 8px;
    }
    div[data-testid="stSidebarNav"] { display: none; }
    div[data-testid="stToast"] { border-radius: 14px !important; }

    /* --- CARD POLISH --- */
    div[data-testid="stContainer"] {
        background: var(--mantis-card-bg);
        border-radius: 20px !important;
        padding: 22px !important;
        border: 1px solid var(--mantis-card-border) !important;
        box-shadow: var(--mantis-shadow-strong);
        margin-bottom: 18px;
    }
    .mantis-pill {
        display:inline-flex;
        align-items:center;
        gap:6px;
        padding:6px 10px;
        border-radius:999px;
        font-size:12px;
        background: var(--mantis-accent-soft);
        border: 1px solid var(--mantis-accent-glow);
        color: var(--mantis-text);
        letter-spacing: 0.02em;
    }
    .mantis-hero-caption {
        font-size:12px;
        color: var(--mantis-muted);
        margin-top:4px;
    }
    div[data-testid="stContainer"] h3 {
        margin-top: 0;
        margin-bottom: 12px;
        color: var(--mantis-text);
    }
    .mantis-hero {
        background: var(--mantis-surface);
        border-radius: 22px;
        padding: 22px 24px;
        border: 1px solid var(--mantis-card-border);
        box-shadow: var(--mantis-shadow-strong);
    }
    .mantis-hero-title {
        font-size: 28px;
        font-weight: 700;
        margin-bottom: 6px;
    }
    .mantis-hero-sub {
        color: var(--mantis-muted);
        font-size: 14px;
    }
    .mantis-page-header {
        display: flex;
        justify-content: space-between;
        align-items: flex-start;
        gap: 12px;
    }
    .mantis-page-title {
        font-size: 26px;
        font-weight: 700;
        margin: 0;
        color: var(--mantis-text);
    }
    .mantis-page-sub {
        color: var(--mantis-muted);
        margin-top: 4px;
        font-size: 14px;
    }
    .mantis-tag {
        display:inline-flex;
        align-items:center;
        gap:6px;
        padding:6px 12px;
        border-radius:999px;
        font-size:12px;
        font-weight:600;
        background: var(--mantis-accent-soft);
        color: var(--mantis-text);
        border: 1px solid var(--mantis-accent-glow);
    }
    .mantis-kpi-grid {
        display:grid;
        grid-template-columns: repeat(auto-fit, minmax(140px, 1fr));
        gap:12px;
        margin-top: 14px;
    }
    .mantis-kpi-card {
        padding: 12px 14px;
        border-radius: 16px;
        background: var(--mantis-surface-alt);
        border: 1px solid var(--mantis-card-border);
    }
    .mantis-kpi-label {
        font-size:11px;
        text-transform: uppercase;
        letter-spacing: 0.08em;
        color: var(--mantis-muted);
        margin-bottom: 6px;
    }
    .mantis-kpi-value {
        font-size:20px;
        font-weight:700;
    }
    .mantis-section-title {
        font-size:20px;
        font-weight:700;
        margin-bottom: 6px;
    }
    .mantis-section-header {
        display:flex;
        align-items:flex-end;
        justify-content:space-between;
        gap:16px;
        margin: 6px 0 14px;
    }
    .mantis-section-caption {
        color: var(--mantis-muted);
        font-size: 13px;
    }
    .mantis-soft {
        background: var(--mantis-surface-alt);
        border-radius: 16px;
        padding: 14px;
        border: 1px solid var(--mantis-card-border);
    }
    .mantis-stat-tile {
        display:flex;
        flex-direction:column;
        gap:6px;
        padding: 12px 14px;
        border-radius: 16px;
        background: var(--mantis-surface-alt);
        border: 1px solid var(--mantis-card-border);
    }
    .mantis-stat-icon {
        width: 30px;
        height: 30px;
        border-radius: 10px;
        display:flex;
        align-items:center;
        justify-content:center;
        background: var(--mantis-accent-soft);
        border: 1px solid var(--mantis-accent-glow);
        font-size: 14px;
    }
    .mantis-stat-label {
        font-size: 11px;
        text-transform: uppercase;
        letter-spacing: 0.08em;
        color: var(--mantis-muted);
    }
    .mantis-stat-value {
        font-size: 20px;
        font-weight: 700;
        color: var(--mantis-text);
    }
    .mantis-stat-help {
        font-size: 12px;
        color: var(--mantis-muted);
    }

    /* --- SIDEBAR POLISH --- */
    section[data-testid="stSidebar"] {
        background: var(--mantis-sidebar-bg);
        border-right: 1px solid var(--mantis-sidebar-border);
    }
    section[data-testid="stSidebar"] h3 {
        color: var(--mantis-sidebar-title);
        font-weight: 700;
        letter-spacing: 0.04em;
        text-transform: uppercase;
        font-size: 12px;
    }
    section[data-testid="stSidebar"] .stMarkdown,
    section[data-testid="stSidebar"] p,
    section[data-testid="stSidebar"] label {
        color: var(--mantis-text);
    }
    section[data-testid="stSidebar"] .stSelectbox label {
        font-size: 11px;
        text-transform: uppercase;
        letter-spacing: 0.08em;
        color: var(--mantis-muted);
        font-weight: 700;
    }
    section[data-testid="stSidebar"] .stSelectbox [data-baseweb="select"] {
        border-radius: 12px;
        border: 1px solid var(--mantis-card-border);
        background: var(--mantis-surface-alt);
    }
    .mantis-nav-section {
        font-size: 11px;
        text-transform: uppercase;
        letter-spacing: 0.08em;
        color: var(--mantis-muted);
        font-weight: 700;
        margin: 4px 2px 6px;
    }
    section[data-testid="stSidebar"] .stButton > button {
        border-radius: 12px !important;
    }


    /* --- SIDEBAR BRAND --- */
    .mantis-sidebar-brand{
        display:flex;
        flex-direction: column;
        align-items:center;
        padding:8px 8px 6px;
        margin: 2px 4px 8px;
        border-radius: 16px;
        background: var(--mantis-sidebar-brand-bg);
        border: 1px solid var(--mantis-sidebar-brand-border);
        box-shadow: var(--mantis-shadow-button);
    }
    .mantis-sidebar-brand--modern {
        gap: 2px;
    }
    .mantis-sidebar-logo {
        width: 100%;
        max-width: 148px;
        margin: 0 auto 6px;
        padding: 8px;
        border-radius: 14px;
        background: var(--mantis-sidebar-logo-bg);
        background: linear-gradient(145deg, color-mix(in srgb, var(--mantis-sidebar-logo-bg) 92%, transparent), transparent);
        border: 1px solid var(--mantis-sidebar-brand-border);
        border: 1px solid color-mix(in srgb, var(--mantis-accent) 20%, transparent);
    }
    .mantis-sidebar-logo img {
        width: 100%;
        display: block;
        object-fit: contain;
        filter: drop-shadow(0 0 8px color-mix(in srgb, var(--mantis-accent-glow) 80%, transparent));
    }
    .mantis-sidebar-brand--modern [data-testid="stImage"] {
        width: 100%;
        margin: 0 0 4px;
    }
    .mantis-sidebar-brand--modern [data-testid="stImage"] img{
        width: 100%;
        max-height: 84px;
        display:block;
        margin: 0 auto;
        object-fit: contain;
        filter: drop-shadow(0 0 8px rgba(34,197,94,0.26));
        transition: transform 180ms ease, filter 180ms ease;
    }
    .mantis-sidebar-brand--modern [data-testid="stImage"] img:hover {
        transform: scale(1.02);
        filter: drop-shadow(0 0 10px rgba(34,197,94,0.36));
    }
    .mantis-avatar {
        height:44px;
        width:44px;
        border-radius:50%;
        background: var(--mantis-surface-alt);
        color: var(--mantis-text);
        display:flex;
        align-items:center;
        justify-content:center;
        font-weight:700;
        font-size:0.9rem;
        border: 1px solid var(--mantis-card-border);
    }
    .mantis-logo-fallback {
        font-size: 0.95rem;
        font-weight: 700;
        color: var(--mantis-sidebar-title);
    }
    .mantis-sidebar-meta {
        width: 100%;
        text-align: center;
        margin-top: 2px;
        margin-bottom: 4px;
    }
    .mantis-sidebar-version{
        font-size:11px;
        letter-spacing: 0.04em;
        text-transform: uppercase;
        color: var(--mantis-muted);
        opacity: 0.9;
    }
    .mantis-sidebar-divider {
        width: 100%;
        height: 1px;
        margin-top: 2px;
        background: linear-gradient(90deg, rgba(34,197,94,0), rgba(34,197,94,0.35), rgba(34,197,94,0));
    }
    @media (max-width: 1080px) {
        .mantis-sidebar-brand--modern [data-testid="stImage"] img{
            max-height: 64px;
        }
    }
    .mantis-banner img {
        max-height: 180px;
        object-fit: contain;
    }

    /* --- NAV RADIO STYLE --- */
    div[role="radiogroup"] > label {
        background: var(--mantis-surface-alt);
        padding: 10px 12px;
        border-radius: 12px;
        border: 1px solid var(--mantis-card-border);
        margin-bottom: 8px;
    }
    div[role="radiogroup"] > label span {
        color: var(--mantis-text);
    }
    div[role="radiogroup"] > label:has(input:checked) {
        border-color: var(--mantis-accent);
        box-shadow: 0 0 0 1px var(--mantis-accent-glow);
    }
    </style>
    """
    )

    st.html(
        """
        <style>
        :root {
            --mantis-bg-primary: var(--mantis-bg);
            --mantis-bg-secondary: var(--mantis-surface);
            --mantis-accent-enterprise: #23f7c0;
            --mantis-accent-teal: #4be5ff;
            --mantis-success: var(--mantis-success, #28f29c);
            --mantis-warning: var(--mantis-warning, #ffb347);
            --mantis-danger: var(--mantis-danger, #ff6b6b);
        }
        [data-testid="stAppViewContainer"] { background: var(--mantis-bg-primary); }
        :root { --mantis-top-offset: 1.25rem; }
        section.main { padding-top: var(--mantis-top-offset) !important; overflow: visible !important; }
        div[data-testid="stAppViewContainer"] { overflow: visible !important; }
        div[data-testid="stAppViewContainer"] > .main { overflow: visible !important; }
        div[data-testid="stMainBlockContainer"] { padding-top: 0.2rem !important; }
        .stApp { overflow: visible !important; }
        .block-container { max-width: 1440px; padding-top: 0.85rem !important; overflow: visible !important; }
        .block-container > div:first-child { margin-top: 0 !important; }
        .block-container > div[data-testid="stVerticalBlock"]:first-child { margin-top: 0.65rem !important; }
        section[data-testid="stSidebar"] .block-container {
            padding-top: 0 !important;
            padding-left: 0.35rem !important;
            padding-right: 0.35rem !important;
            padding-bottom: 0.6rem !important;
        }
        section[data-testid="stSidebar"] div[data-testid="stSidebarUserContent"] {
            padding-top: 0 !important;
            padding-left: 0.6rem !important;
            padding-right: 0.6rem !important;
        }
        section[data-testid="stSidebar"] > div:first-child {
            padding-top: 0 !important;
        }
        /* Prevent first-page headings from clipping under the top bar */
        h1, h2, h3,
        .stMarkdown h1, .stMarkdown h2, .stMarkdown h3 {
            line-height: 1.25 !important;
            padding-top: 0.12em !important;
            overflow: visible !important;
        }
        .block-container > div:first-child h1,
        .block-container > div:first-child h2,
        .block-container > div:first-child h3 {
            margin-top: 0.25rem !important;
        }
        header[data-testid="stHeader"],
        div[data-testid="stToolbar"],
        div[data-testid="stDecoration"],
        div[data-testid="stStatusWidget"] {
            display: none !important;
            height: 0 !important;
            min-height: 0 !important;
            visibility: hidden !important;
        }
        button[data-testid="collapsedControl"] {
            display: none !important;
            visibility: hidden !important;
            opacity: 0 !important;
            pointer-events: none !important;
        }
        section[data-testid="stSidebar"] {
            transform: translateX(0) !important;
            margin-left: 0 !important;
            min-width: 280px !important;
        }
        .mantis-topnav-brand { display:flex; flex-direction:column; justify-content:center; }
        .mantis-topnav-title { font-size:1.1rem; letter-spacing:0.08em; font-weight:800; }
        .mantis-topnav-sub { font-size:0.78rem; color:var(--mantis-muted); }
        .mantis-status-badge { margin-top:2px; display:inline-flex; gap:8px; align-items:center; padding:8px 12px; border-radius:999px; font-size:0.8rem; border:1px solid var(--mantis-card-border); }
        .mantis-status-badge .dot { width:8px; height:8px; border-radius:50%; display:inline-block; }
        .mantis-status-badge.ready .dot { background: var(--mantis-success); box-shadow: 0 0 12px var(--mantis-success); }
        .mantis-status-badge.processing .dot { background: var(--mantis-warning); animation: mantisPulse 1.2s infinite; }
        .mantis-status-badge.limited .dot { background: var(--mantis-danger); }
        .mantis-enterprise-hero { display:flex; justify-content:space-between; align-items:flex-start; gap:20px; padding: 20px 22px; border-radius:16px; background: var(--mantis-surface); border:1px solid var(--mantis-card-border); }
        .mantis-enterprise-hero h1 { margin:0; font-size:2rem; letter-spacing:-0.02em; }
        .mantis-enterprise-hero p { margin:6px 0 0 0; color:var(--mantis-muted); }
        .mantis-status-pill { border:1px solid var(--mantis-primary-border); background:var(--mantis-accent-soft); color:var(--mantis-text); border-radius:999px; padding:8px 12px; font-size:0.82rem; display:inline-flex; align-items:center; gap:8px; }
        .mantis-status-pill span { width:8px; height:8px; border-radius:50%; background:var(--mantis-success); box-shadow:0 0 10px var(--mantis-success); }
        .mantis-metric-card { border-radius:14px; padding:16px; border:1px solid var(--mantis-card-border); background:var(--mantis-bg-secondary); min-height:104px; transition:transform .16s ease, border-color .16s ease; }
        .mantis-metric-card:hover { transform: translateY(-2px); border-color: var(--mantis-accent-teal); }
        .mantis-metric-label { font-size:0.75rem; text-transform:uppercase; letter-spacing:0.08em; color:var(--mantis-muted); }
        .mantis-metric-value { margin-top:8px; font-size:1.6rem; font-weight:750; }
        .mantis-activity-item { display:grid; grid-template-columns: 120px 120px 1fr; gap:12px; padding:10px 0; border-bottom:1px solid var(--mantis-card-border); }
        .mantis-activity-time { color:var(--mantis-muted); font-size:0.78rem; }
        .mantis-activity-status { font-size:0.78rem; color:var(--mantis-text); }
        .mantis-activity-text { font-size:0.9rem; }
        .mantis-project-chip { border:1px solid var(--mantis-card-border); border-radius:12px; padding:10px 12px; background:var(--mantis-surface-alt); }
        .mantis-project-chip__title { font-size:0.88rem; font-weight:650; }
        .mantis-project-chip__meta { font-size:0.74rem; color:var(--mantis-muted); margin-top:4px; }
        @keyframes mantisPulse { 0%{opacity:1;} 50%{opacity:.3;} 100%{opacity:1;} }
        </style>
        """
    )

    logger.info("Initializing session state...")
    
    # Load app configuration for session state initialization
    config_data = load_app_config()
    logger.debug(f"Loaded app config: {list(config_data.keys())}")
    
    init_state("user_id", None)
    init_state("projects_dir", None)
    init_state("project", None)
    init_state("page", "home")
    logger.debug(f"Default page set to: {st.session_state.page}")
    init_state("auto_save", True)
    init_state("ghost_text", "")
    init_state("pending_improvement_text", "")
    init_state("pending_improvement_meta", {})
    init_state("chapter_text_prev", {})
    init_state("chapter_drafts", [])
    init_state("chapters", [])
    init_state("chapters_project_id", None)
    init_state("active_chapter_id", None)
    init_state("editor_improve__copy_buffer", "")
    init_state("first_run", True)
    init_state("is_premium", True)
    init_state("pending_action", None)
    st.session_state.setdefault("ai_keys", {})
    logger.info("Session state initialization complete")

    def _resolve_api_key(provider: str, default_value: str) -> str:
        session_key = (st.session_state.get("ai_keys") or {}).get(provider, "")
        if session_key:
            return session_key
        config_key = config_data.get(f"{provider}_api_key", "")
        if config_key:
            return config_key
        return default_value or ""
    init_state(
        "openai_base_url",
        config_data.get("openai_base_url", AppConfig.OPENAI_API_URL),
    )
    init_state("ai_provider", config_data.get("ai_provider", "groq"))
    init_state("ai_session_keys", {
        "openai": config_data.get("openai_api_key", ""),
        "groq": config_data.get("groq_api_key", ""),
    })
    init_state("openai_key_input", "")
    init_state("openai_model", config_data.get("openai_model", AppConfig.OPENAI_MODEL))
    init_state("openai_model_list", [])
    init_state("openai_model_tests", {})
    init_state("ui_theme", config_data.get("ui_theme", "Dark"))
    init_state("world_bible_confidence_threshold", float(config_data.get("world_bible_confidence_threshold", AppConfig.WORLD_BIBLE_CONFIDENCE)))
    init_state("groq_base_url", config_data.get("groq_base_url", AppConfig.GROQ_API_URL))
    init_state("groq_key_input", "")
    init_state("groq_model", config_data.get("groq_model", AppConfig.DEFAULT_MODEL))
    init_state("groq_model_list", [])
    init_state("groq_model_tests", {})
    init_state("_force_nav", False)
    init_state("_scroll_top_nonce", 0)
    init_state("_last_scroll_top_nonce", -1)
    init_state("memory_auto_soft_enabled", bool(config_data.get("memory_auto_soft_enabled", AppConfig.WORLD_MEMORY_AUTO_SOFT)))
    init_state("memory_auto_soft_threshold", float(config_data.get("memory_auto_soft_threshold", AppConfig.WORLD_MEMORY_SOFT_CONFIDENCE)))
    init_state("memory_auto_hard_enabled", bool(config_data.get("memory_auto_hard_enabled", AppConfig.WORLD_MEMORY_AUTO_HARD)))
    init_state("memory_auto_hard_threshold", float(config_data.get("memory_auto_hard_threshold", AppConfig.WORLD_MEMORY_HARD_CONFIDENCE)))

    AppConfig.GROQ_API_URL = st.session_state.groq_base_url
    AppConfig.DEFAULT_MODEL = st.session_state.groq_model
    AppConfig.OPENAI_API_URL = st.session_state.openai_base_url
    AppConfig.OPENAI_MODEL = st.session_state.openai_model

    def debug_enabled() -> bool:
        try:
            secrets = st.secrets
        except Exception:
            secrets = {}
        if isinstance(secrets, dict):
            return bool(secrets.get("DEBUG")) or bool(st.session_state.get("debug"))
        try:
            return bool(secrets["DEBUG"]) or bool(st.session_state.get("debug"))
        except Exception:
            return bool(st.session_state.get("debug"))

    def open_legal_page() -> None:
        st.session_state.page = "legal"
        st.rerun()

    def persist_project(project: "Project", *, action: str = "save") -> bool:
        """Save project to local storage."""
        path = project.save()
        if not path:
            logger.error("persist_project failed for '%s' (action=%s)", project.title, action)
            try:
                st.toast("Save failed - check file permissions and disk space.", icon="\u26A0")
            except Exception as e:
                st.error(f"Save failed for '{project.title}': {e}")
            return False
        # Remember the last active project so it can be restored after refresh
        config = load_app_config()
        config["last_project_path"] = path
        save_app_config(config)
        return True

    def check_and_show_whats_new() -> None:
        """Check if app version has changed and show What's New notification."""
        config = load_app_config()
        last_seen_version = config.get("last_seen_version", "")
        current_version = AppConfig.VERSION

        # Check if this is a new version
        if last_seen_version and last_seen_version != current_version:
            # Show What's New banner
            with st.container(border=True):
                st.markdown("### What's New in Mantis Studio")
                st.markdown(f"""
                **Version {current_version} is now available!** (Updated from {last_seen_version})
                
                **What Changed:**
                - **NEW**: You can now see what changed between versions!
                  - This notification system was added to address user feedback
                  - Version updates are now visible and transparent
                - **Improved**: Complete changelog documentation for all recent versions
                - **Enhanced**: Better version tracking and update notifications
                
                **Why This Matters:**
                - You previously reported: "merged 4 times now with no changed from users point of view"
                - This fix ensures you always know when the app updates and what's new
                - All future updates will show clear release notes
                
                **See the [full changelog](https://github.com/bigmanjer/Mantis-Studio/blob/main/docs/CHANGELOG.md) for complete details**
                """)
                
                col1, col2 = st.columns([1, 1])
                with col1:
                    if st.button("Got it, thanks!", use_container_width=True, type="primary"):
                        # Update last seen version
                        config["last_seen_version"] = current_version
                        save_app_config(config)
                        st.rerun()
                with col2:
                    if st.button("View Full Changelog", use_container_width=True):
                        webbrowser.open("https://github.com/bigmanjer/Mantis-Studio/blob/main/docs/CHANGELOG.md")
                        config["last_seen_version"] = current_version
                        save_app_config(config)
                        st.rerun()
        elif not last_seen_version:
            # First time user - set version silently
            config["last_seen_version"] = current_version
            save_app_config(config)

    def render_welcome_banner(context: str) -> None:
        """Show helpful context-aware guidance for first-time users."""
        # Show What's New banner if version changed
        if context == "home":
            check_and_show_whats_new()

        if context == "home" and st.session_state.get("first_run", True):
            provider, active_key, _ = get_active_key_status()
            with st.container(border=True):
                st.markdown("### Welcome to MANTIS Studio")
                st.caption("Plan your story. Write your draft. Keep canon consistent.")
                col_a, col_b = st.columns(2)
                with col_a:
                    st.markdown("**Start here**")
                    st.markdown(
                        f"""
                        - Create your first project from **Projects**
                        - Build an outline before drafting chapters
                        - Follow the [Getting Started Guide]({AppConfig.GETTING_STARTED_URL})
                        - Projects auto-save locally to your workspace
                        """
                    )
                with col_b:
                    st.markdown("**Core workflow**")
                    st.markdown(
                        """
                        1. Outline your narrative structure  
                        2. Draft chapters in the editor  
                        3. Use World Bible to track canon  
                        4. Export when the manuscript is ready
                        """
                    )
                if not active_key:
                    st.info(
                        f"AI is optional to start. Connect {_provider_label(provider)} in AI Settings when you're ready for generation and summaries."
                    )
                else:
                    st.success("AI is connected. Generation, summaries, and entity scans are ready.")
                col1, col2, col3 = st.columns(3)
                with col1:
                    if st.button("Create my first project", use_container_width=True, type="primary", key="welcome_create_first"):
                        st.session_state.first_run = False
                        st.session_state.page = "projects"
                        st.rerun()
                with col2:
                    if st.button("Open AI settings", use_container_width=True, key="welcome_open_ai"):
                        st.session_state.first_run = False
                        st.session_state.page = "ai"
                        st.rerun()
                with col3:
                    if st.button("Continue to dashboard", use_container_width=True, key="welcome_dismiss"):
                        st.session_state.first_run = False
                        st.rerun()
        elif st.session_state.get("first_run", True) and context != "home":
            # Intentionally no tip banner on non-home pages so the page header
            # card remains the first visible element.
            return

    def render_app_footer() -> None:
        render_footer(AppConfig.VERSION)

    def _render_header_bar(recent_projects: List[Dict[str, Any]]) -> None:
        logo_bytes = load_first_asset_bytes(
            "branding/mantis_emblem.png",
            "mantis_emblem.png",
            "branding/mantis_wordmark.png",
            "mantis_wordmark.png",
        )
        left_col, mid_col, right_col = st.columns([1.6, 2.2, 1.2], vertical_alignment="center")
        with left_col:
            if logo_bytes:
                st.image(logo_bytes, width=42)
            st.markdown("### MANTIS Studio")
            st.caption("Modular AI Narrative Text Intelligence System")

        with mid_col:
            options = ["No project selected"]
            project_map = {"No project selected": None}
            for entry in recent_projects:
                meta = entry.get("meta", {})
                title = meta.get("title") or os.path.basename(entry.get("path", "Untitled"))
                label = f"{title}"
                options.append(label)
                project_map[label] = entry.get("path")
            current_label = "No project selected"
            if st.session_state.project:
                current_label = st.session_state.project.title or current_label
            selection = st.selectbox(
                "Current project",
                options,
                index=options.index(current_label) if current_label in options else 0,
                help="Quickly jump between recent projects.",
            )
            selected_path = project_map.get(selection)
            if selected_path and (not st.session_state.project or st.session_state.project.filepath != selected_path):
                loaded = load_project_safe(selected_path, context="project selector")
                if loaded:
                    st.session_state.project = loaded
                    st.session_state.page = "chapters"
                    st.rerun()

        with right_col:
            tags = []
            if st.session_state.project:
                tags.append("Project active")
            render_tag_list(tags)

    # Local-first: use the default projects directory
    if not st.session_state.get("projects_dir"):
        st.session_state.projects_dir = AppConfig.PROJECTS_DIR
        logger.info(f"Projects directory set to: {AppConfig.PROJECTS_DIR}")

    # Restore the last active project from config if none is loaded
    if not st.session_state.get("project"):
        last_path = config_data.get("last_project_path", "")
        if last_path and os.path.isfile(last_path):
            try:
                st.session_state.project = Project.load(last_path)
                logger.info(f"Restored last project: {st.session_state.project.title}")
            except Exception as e:
                logger.warning("Failed to restore last project: %s - %s", last_path, e, exc_info=True)
    
    # Startup diagnostics - run once per session
    if not st.session_state.get("_startup_diagnostics_run"):
        st.session_state._startup_diagnostics_run = True
        logger.info("=" * 60)
        logger.info("STARTUP DIAGNOSTICS")
        logger.info("=" * 60)
        logger.info(f"App Version: {AppConfig.VERSION}")
        logger.info(f"Python Version: {sys.version.split()[0]}")
        logger.info(f"Streamlit Version: {st.__version__}")
        logger.info(f"Projects Directory: {AppConfig.PROJECTS_DIR}")
        logger.info(f"Projects Dir Exists: {os.path.exists(AppConfig.PROJECTS_DIR)}")
        logger.info(f"Config Path: {AppConfig.CONFIG_PATH}")
        logger.info(f"Config Exists: {os.path.exists(AppConfig.CONFIG_PATH)}")
        logger.info(f"Assets Directory: {ASSETS_DIR}")
        logger.info(f"Assets Dir Exists: {ASSETS_DIR.exists()}")
        logger.info(f"Current Page: {st.session_state.page}")
        logger.info(f"Project Loaded: {st.session_state.project is not None}")
        if st.session_state.project:
            logger.info(f"  - Project Title: {st.session_state.project.title}")
            logger.info(f"  - Project Path: {st.session_state.project.filepath}")
        logger.info(f"Session State Keys: {len(st.session_state.keys())} total")
        logger.info(f"Debug Mode: {debug_enabled()}")
        logger.info("=" * 60)
        logger.info("STARTUP DIAGNOSTICS COMPLETE - App initialized successfully")
        logger.info("=" * 60)

    # Reliable navigation rerun (avoids Streamlit edge cases when returning early)
    if st.session_state.get("_force_nav"):
        st.session_state._force_nav = False
        st.rerun()

    def get_ai_model() -> str:
        provider = _normalize_provider(st.session_state.get("ai_provider", "groq"))
        if provider == "openai":
            return st.session_state.get("openai_model", AppConfig.OPENAI_MODEL)
        return st.session_state.get("groq_model", AppConfig.DEFAULT_MODEL)

    def get_active_key_status() -> tuple[str, str, str]:
        provider = _normalize_provider(st.session_state.get("ai_provider", "groq"))
        key, source = get_effective_key(provider, st.session_state.get("user_id"))
        return provider, key, source

    def save_p():
        if st.session_state.project and st.session_state.auto_save:
            persist_project(st.session_state.project)

    def get_active_projects_dir() -> Optional[str]:
        """Return the active projects directory (local-first, single user)."""
        return st.session_state.get("projects_dir") or AppConfig.PROJECTS_DIR

    def resume_pending_action() -> None:
        pending = st.session_state.get("pending_action")
        if not pending:
            return
        action = pending.get("action")
        payload = pending.get("payload") or {}
        return_to = pending.get("return_to") or "home"
        st.session_state.pending_action = None

        if action == "create_project":
            p = Project.create(
                payload.get("title") or "Untitled Project",
                author=payload.get("author", ""),
                genre=payload.get("genre", ""),
                storage_dir=get_active_projects_dir() or AppConfig.PROJECTS_DIR,
            )
            p.save()
            st.session_state.project = p
        elif action == "import_project":
            text = payload.get("text") or ""
            p = Project.create("Imported Project", storage_dir=get_active_projects_dir() or AppConfig.PROJECTS_DIR)
            p.import_text_file(text)
            active_provider = _normalize_provider(st.session_state.get("ai_provider", "groq"))
            active_key, _ = get_effective_key(active_provider, st.session_state.get("user_id"))
            if active_key and get_ai_model():
                p.outline = StoryEngine.reverse_engineer_outline(p, get_ai_model())
            p.save()
            st.session_state.project = p
        elif action == "save_project":
            if st.session_state.project:
                st.session_state.project.save()
                st.toast("Project saved.")
        elif action == "save_outline":
            if st.session_state.project:
                st.session_state.project.save()
                st.toast("Outline saved.")
        elif action == "save_chapter":
            if st.session_state.project:
                st.session_state.project.save()
                st.toast("Chapter saved.")
        elif action == "export_project":
            if st.session_state.project:
                st.session_state.project.save()
                st.session_state.export_project_path = st.session_state.project.filepath
        elif action == "save_app_settings":
            save_app_settings()
        elif action == "delete_project":
            delete_path = payload.get("path")
            if delete_path:
                deleted, _msg = Project.delete_file(delete_path)
                if deleted:
                    _bump_projects_refresh()
                else:
                    st.toast(f"Delete failed: {_msg}")

        st.session_state.page = return_to
        st.rerun()

    if st.session_state.get("pending_action"):
        resume_pending_action()


    def load_project_safe(path: str, context: str = "project") -> Optional["Project"]:
        try:
            return Project.load(path)
        except Exception:
            logger.warning("Failed to load %s: %s", context, path, exc_info=True)
            st.error("We couldn't open that project file. It may be missing or corrupted.")
            return None

    def show_api_key_notice(action: str) -> None:
        """Display a notice that an API key is required for the specified action.
        
        Args:
            action: Description of what the user is trying to do (e.g., "scan entities", "generate outlines")
        """
        provider, _, _ = get_active_key_status()
        st.info(f"Add a {_provider_label(provider)} API key in AI Settings to {action}.")

    def get_ai_button_help(
        cooldown_seconds: int,
        has_api_key: bool,
        action_description: str,
    ) -> str:
        """Generate contextual help text for AI-powered buttons.
        
        Args:
            cooldown_seconds: Remaining cooldown time in seconds (0 if no cooldown)
            has_api_key: Whether the user has configured an API key
            action_description: What the button does (e.g., "generate a chapter outline")
        
        Returns:
            Appropriate help text explaining the button's state
        """
        if cooldown_seconds > 0:
            return f"Available in {cooldown_seconds} seconds to prevent API overuse"
        elif not has_api_key:
            provider, _, _ = get_active_key_status()
            return f"Configure {_provider_label(provider)} API key in AI Settings to use this feature"
        else:
            return f"AI will {action_description}"

    def render_page_header(
        title: str,
        subtitle: str,
        primary_label: Optional[str] = None,
        primary_action: Optional[Callable[[], None]] = None,
        primary_disabled: bool = False,
        primary_help: Optional[str] = None,
        secondary_label: Optional[str] = None,
        secondary_action: Optional[Callable[[], None]] = None,
        secondary_disabled: bool = False,
        secondary_help: Optional[str] = None,
        tag: Optional[str] = None,
        key_prefix: str = "page_header",
    ) -> None:
        header_bar(title, subtitle, pill=tag)
        action_cols = st.columns(2)
        with action_cols[0]:
            if primary_label and primary_action:
                if st.button(
                    primary_label,
                    type="primary",
                    use_container_width=True,
                    key=f"{key_prefix}__primary",
                    disabled=primary_disabled,
                    help=primary_help,
                ):
                    primary_action()
        with action_cols[1]:
            if secondary_label and secondary_action:
                if st.button(
                    secondary_label,
                    use_container_width=True,
                    key=f"{key_prefix}__secondary",
                    disabled=secondary_disabled,
                    help=secondary_help,
                ):
                    secondary_action()

    @st.cache_data(show_spinner=False)
    def _cached_models(base_url: str, api_key: str) -> List[str]:
        return AIEngine(base_url=base_url).probe_models(api_key)

    def refresh_models():
        groq_key, _ = get_effective_key("groq", st.session_state.get("user_id"))
        st.session_state.groq_model_list = _cached_models(
            st.session_state.groq_base_url,
            groq_key,
        ) or []

    def refresh_openai_models():
        openai_key, _ = get_effective_key("openai", st.session_state.get("user_id"))
        st.session_state.openai_model_list = _cached_models(
            st.session_state.openai_base_url,
            openai_key,
        ) or []

    def save_app_settings():
        # Merge with existing config to preserve saved API keys
        data = load_app_config()
        data.update({
            "ai_provider": st.session_state.ai_provider,
            "groq_base_url": st.session_state.groq_base_url,
            "groq_model": st.session_state.groq_model,
            "openai_base_url": st.session_state.openai_base_url,
            "openai_model": st.session_state.openai_model,
            "ui_theme": st.session_state.ui_theme,
            "daily_word_goal": int(st.session_state.daily_word_goal),
            "weekly_sessions_goal": int(st.session_state.weekly_sessions_goal),
            "focus_minutes": int(st.session_state.focus_minutes),
            "activity_log": list(st.session_state.activity_log),
        })
        save_app_config(data)
        st.toast("Settings saved.")

    def _today_str() -> str:
        return datetime.date.today().isoformat()

    def _parse_day(day: str) -> Optional[datetime.date]:
        try:
            return datetime.date.fromisoformat(day)
        except ValueError:
            return None

    def _log_activity():
        today = _today_str()
        log = set(st.session_state.activity_log)
        log.add(today)
        st.session_state.activity_log = sorted(log)
        save_app_settings()

    def _weekly_activity_count() -> int:
        today = datetime.date.today()
        cutoff = today - datetime.timedelta(days=6)
        count = 0
        for day in st.session_state.activity_log:
            parsed = _parse_day(day)
            if parsed and parsed >= cutoff:
                count += 1
        return count

    def _activity_streak() -> int:
        if not st.session_state.activity_log:
            return 0
        days = sorted(
            {d for d in (_parse_day(day) for day in st.session_state.activity_log) if d}
        )
        if not days:
            return 0
        streak = 0
        cursor = datetime.date.today()
        day_set = set(days)
        while cursor in day_set:
            streak += 1
            cursor -= datetime.timedelta(days=1)
        return streak

    def _activity_series() -> List[Dict[str, Any]]:
        today = datetime.date.today()
        labels = []
        counts = []
        log_set = set(st.session_state.activity_log)
        for offset in range(6, -1, -1):
            day = today - datetime.timedelta(days=offset)
            labels.append(day.strftime("%a"))
            counts.append(1 if day.isoformat() in log_set else 0)
        return [{"day": label, "sessions": count} for label, count in zip(labels, counts)]

    def _cooldown_remaining(action_key: str, cooldown_s: int = 8) -> int:
        cooldowns = st.session_state.setdefault("action_cooldowns", {})
        last_ts = cooldowns.get(action_key, 0)
        remaining = max(0, cooldown_s - (time.time() - last_ts))
        return int(round(remaining))

    def _mark_action(action_key: str) -> None:
        cooldowns = st.session_state.setdefault("action_cooldowns", {})
        cooldowns[action_key] = time.time()

    def _load_recent_projects(active_dir: Optional[str], refresh_token: int) -> List[Dict[str, Any]]:
        del refresh_token
        if not active_dir or not os.path.exists(active_dir):
            return []
        files = sorted(
            [
                f
                for f in os.listdir(active_dir)
                if f.endswith(".json")
                and not f.startswith(".")
                and f.lower() != ".mantis_config.json"
            ],
            key=lambda x: os.path.getmtime(os.path.join(active_dir, x)),
            reverse=True,
        )
        projects = []
        for filename in files:
            full_path = os.path.join(active_dir, filename)
            try:
                with open(full_path, "r", encoding="utf-8") as fh:
                    meta = json.load(fh)
                if not isinstance(meta, dict):
                    continue
                # Guard against non-project JSON files living in the same folder.
                if not (meta.get("id") or meta.get("title") or meta.get("chapters") or meta.get("outline")):
                    continue
                projects.append({"path": full_path, "meta": meta})
            except Exception:
                logger.warning("Failed to load project metadata: %s", full_path, exc_info=True)
        return projects

    def _bump_projects_refresh() -> None:
        st.session_state.projects_refresh_token += 1

    def _project_snapshot(meta: Dict[str, Any]) -> Dict[str, Any]:
        chapters = meta.get("chapters", {}) or {}
        chapter_list = list(chapters.values())
        word_count = sum(int(c.get("word_count") or 0) for c in chapter_list)
        return {
            "title": meta.get("title") or "Untitled Project",
            "genre": meta.get("genre") or "General Fiction",
            "chapters": len(chapter_list),
            "words": word_count,
            "modified_at": meta.get("last_modified") or meta.get("modified_at"),
        }

    FALLBACK_PROJECT_GENRES = ["Fantasy", "Adventure"]
    MAX_DRAFT_EXCERPT_LENGTH = 600
    AI_ERROR_MARKERS = (
        "not configured",
        "generation failed",
        "response empty",
        "error",
    )

    def _format_genre_list(genres: List[str]) -> str:
        return "  ".join(genres)

    def _parse_genre_list(raw: str) -> List[str]:
        cleaned = re.sub(r"(?i)^(genres?)[:\s-]*", "", (raw or "")).strip()
        if not cleaned:
            return []
        cleaned = cleaned.replace("|", ",")
        parts = re.split(r"[,\n/]+", cleaned)
        return [part.strip() for part in parts if part.strip()]

    def _ai_generation_available() -> bool:
        provider = _normalize_provider(st.session_state.get("ai_provider", "groq"))
        key, _ = get_effective_key(provider, st.session_state.get("user_id"))
        model = get_ai_model()
        if not key or not model:
            return False
        base_url = _get_provider_base_url(provider)
        try:
            if provider == "openai":
                ok, _ = test_openai_model(base_url, key, model)
            else:
                ok, _ = test_groq_model(base_url, key, model)
            return ok
        except Exception:
            logger.warning(
                "AI availability check failed for provider=%s model=%s",
                provider,
                model,
                exc_info=True,
            )
            return False

    def _build_project_context(
        context: str,
        title: str,
        genre: str,
        author: str,
        draft_excerpt: str,
    ) -> str:
        parts = []
        if context:
            parts.append(context.strip())
        if title:
            parts.append(f"Title idea: {title}")
        if genre:
            parts.append(f"Genre hints: {genre}")
        if author:
            parts.append(f"Author: {author}")
        safe_excerpt = draft_excerpt.strip()
        if safe_excerpt:
            parts.append(f"Draft excerpt: {safe_excerpt[:MAX_DRAFT_EXCERPT_LENGTH]}")
        return "\n".join(parts) if parts else "A new creative writing project."

    def _ai_response_has_error(raw: str) -> bool:
        low = (raw or "").lower()
        return any(marker in low for marker in AI_ERROR_MARKERS)

    def _generate_project_title(context: str, genre_hint: str, model: str, provider: str) -> str:
        prompt = (
            "You are initializing a new creative writing project.\n"
            f"PROJECT CONTEXT:\n{context}\n\n"
            f"GENRE HINTS: {genre_hint or 'Open'}\n"
            "TASK: Generate a creative, relevant project title.\n"
            "RULES: Return ONLY the title. No quotes. No prefixes."
        )
        response = AIEngine(provider=provider).generate(prompt, model)
        raw = (response.get("text", "") or "").strip()
        if _ai_response_has_error(raw):
            return ""
        clean = re.sub(r"(?i)^(title|project title|suggested title)[:\s-]*", "", raw).strip()
        clean = clean.replace('"', "").replace("'", "").strip()
        return clean.splitlines()[0].strip() if clean else ""

    def _generate_project_genres(context: str, title: str, model: str, provider: str) -> List[str]:
        prompt = (
            "You are initializing a new creative writing project.\n"
            f"PROJECT CONTEXT:\n{context}\n\n"
            f"TITLE: {title or 'New Project'}\n"
            "TASK: Suggest 2-4 genres that fit this project.\n"
            "RULES: Return ONLY a comma-separated list of genres."
        )
        response = AIEngine(provider=provider).generate(prompt, model)
        raw = (response.get("text", "") or "").strip()
        if _ai_response_has_error(raw):
            return []
        raw = raw.replace('"', "").replace("'", "").strip()
        return _parse_genre_list(raw)

    def _resolve_project_seed(
        title: str,
        genre: str,
        author: str,
        context: str = "",
        draft_excerpt: str = "",
    ) -> tuple[str, str, List[str], bool]:
        final_title = (title or "").strip()
        final_genre = (genre or "").strip()
        ai_used = False

        needs_title = not final_title
        needs_genre = not final_genre
        ai_available = _ai_generation_available()
        if needs_title or needs_genre:
            if ai_available:
                provider = _normalize_provider(st.session_state.get("ai_provider", "groq"))
                model = get_ai_model()
                context_text = _build_project_context(
                    context,
                    final_title,
                    final_genre,
                    author,
                    draft_excerpt,
                )
                if needs_title:
                    generated_title = _generate_project_title(context_text, final_genre, model, provider)
                    if generated_title:
                        final_title = generated_title
                        ai_used = True
                if needs_genre:
                    generated_genres = _generate_project_genres(context_text, final_title, model, provider)
                    if generated_genres:
                        final_genre = _format_genre_list(generated_genres)
                        ai_used = True

        if not final_title:
            final_title = _random_project_title()
        if not final_genre and not ai_available:
            final_genre = _random_project_genres()

        genre_list = _parse_genre_list(final_genre)
        return final_title, final_genre, genre_list, ai_used

    def _random_project_title() -> str:
        adjectives = [
            "Ashen",
            "Verdant",
            "Crimson",
            "Celestial",
            "Obsidian",
            "Luminous",
            "Forgotten",
            "Gilded",
            "Veiled",
            "Hollow",
            "Radiant",
            "Stormbound",
            "Ivory",
            "Eclipsed",
            "Thorned",
            "Mythic",
        ]
        nouns = [
            "Crown",
            "Archive",
            "Sanctum",
            "Labyrinth",
            "Harbor",
            "Citadel",
            "Oath",
            "Chronicle",
            "Constellation",
            "Axiom",
            "Ember",
            "Signal",
            "Throne",
            "Vesper",
            "Pulse",
            "Emissary",
        ]
        suffixes = [
            "of Hollowlight",
            "of the Verdant Sea",
            "of the Last Meridian",
            "of Emberglass",
            "of the Drowned Sky",
            "of Ironfall",
            "of the Crystal District",
            "of Midnight Bloom",
            "of the Sunken Choir",
            "of Starward",
        ]
        return f"The {random.choice(adjectives)} {random.choice(nouns)} {random.choice(suffixes)}"

    def _random_project_genres() -> str:
        genres = [
            "Solarpunk",
            "Mythic Fantasy",
            "Cosmic Horror",
            "Techno-Thriller",
            "Romantic Suspense",
            "Gaslamp Adventure",
            "Urban Fantasy",
            "Dark Academia",
            "Political Intrigue",
            "Epic Fantasy",
            "Noir Mystery",
            "Found Family",
            "Post-Apocalyptic",
            "Speculative Romance",
            "Spy Drama",
            "Weird Western",
        ]
        genre_count = min(4, len(genres))
        picks = random.sample(genres, k=genre_count)
        return "  ".join(picks)

    def test_groq_connection(base_url: str, api_key: str) -> bool:
        headers = {}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        try:
            r = requests.get(
                f"{base_url.rstrip('/')}/models",
                headers=headers,
                timeout=5,
            )
            r.raise_for_status()
            return True
        except Exception:
            return False

    def test_openai_connection(base_url: str, api_key: str) -> bool:
        headers = {}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        try:
            r = requests.get(
                f"{base_url.rstrip('/')}/models",
                headers=headers,
                timeout=5,
            )
            r.raise_for_status()
            return True
        except Exception:
            return False

    def fetch_groq_models(base_url: str, api_key: str) -> tuple[List[str], str]:
        headers = {}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        try:
            r = requests.get(
                f"{base_url.rstrip('/')}/models",
                headers=headers,
                timeout=10,
            )
            r.raise_for_status()
            data = r.json()
            models = [m.get("id") for m in data.get("data", []) if m.get("id")]
            return models, ""
        except Exception as exc:
            return [], str(exc)

    def fetch_openai_models(base_url: str, api_key: str) -> tuple[List[str], str]:
        headers = {}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        try:
            r = requests.get(
                f"{base_url.rstrip('/')}/models",
                headers=headers,
                timeout=10,
            )
            r.raise_for_status()
            data = r.json()
            models = [m.get("id") for m in data.get("data", []) if m.get("id")]
            return models, ""
        except Exception as exc:
            return [], str(exc)

    def test_groq_model(base_url: str, api_key: str, model: str) -> tuple[bool, str]:
        headers = {"Content-Type": "application/json"}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        payload = {
            "model": model,
            "messages": [{"role": "user", "content": "ping"}],
            "max_tokens": 1,
            "temperature": 0,
        }
        try:
            r = requests.post(
                f"{base_url.rstrip('/')}/chat/completions",
                headers=headers,
                json=payload,
                timeout=15,
            )
            r.raise_for_status()
            return True, ""
        except Exception as exc:
            return False, str(exc)

    def test_openai_model(base_url: str, api_key: str, model: str) -> tuple[bool, str]:
        headers = {"Content-Type": "application/json"}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
        payload = {
            "model": model,
            "messages": [{"role": "user", "content": "ping"}],
            "max_tokens": 1,
            "temperature": 0,
        }
        try:
            r = requests.post(
                f"{base_url.rstrip('/')}/chat/completions",
                headers=headers,
                json=payload,
                timeout=15,
            )
            r.raise_for_status()
            return True, ""
        except Exception as exc:
            return False, str(exc)

    def _queue_world_bible_suggestion(item: Dict[str, Any]) -> None:
        from app.services.world_bible_merge import classify_suggestion
        queue = st.session_state.setdefault("world_bible_review", [])
        # Classify against the live project when available
        p = st.session_state.get("project")
        if p is not None:
            item = classify_suggestion(p, item)
            if item.get("type") == "duplicate":
                return
        aliases = item.get("novel_aliases") or item.get("aliases") or []
        if not isinstance(aliases, list):
            aliases = [str(aliases)]
        normalized_aliases = []
        for alias in aliases:
            clean = str(alias).strip()
            if not clean:
                continue
            normalized_aliases.append(Project._normalize_entity_name(clean))
        alias_key = ",".join(sorted({a for a in normalized_aliases if a}))
        key = (
            f"{Project._normalize_category(item.get('category'))}|"
            f"{Project._normalize_entity_name(item.get('name'))}|"
            f"{alias_key}|"
            f"{(item.get('description') or '').strip().lower()}|"
            f"{item.get('type', 'new')}"
        )
        existing_keys = {q.get("_key") for q in queue}
        if key in existing_keys:
            return
        item["_key"] = key
        queue.append(item)

    def _append_unique_memory_line(existing: str, line: str) -> tuple[str, bool]:
        line = (line or "").strip()
        if not line:
            return existing, False
        current = (existing or "").strip()
        lower_current = current.lower()
        if line.lower() in lower_current:
            return existing, False
        updated = f"{current}\n{line}".strip() if current else line
        return updated, True

    def extract_entities_ui(text: str, label: str, show_feedback: bool = True):
        """Scan text for entities and update the World Bible.

        Uses Project.upsert_entity fuzzy matching so that abbreviations, nicknames,
        or slightly different spellings merge into existing entries instead of
        creating duplicates.
        """
        provider, active_key, _ = get_active_key_status()
        if not active_key:
            show_api_key_notice("scan entities")
            return
        last_scan = st.session_state.get("last_entity_scan")
        if last_scan and (time.time() - last_scan) < 15:
            st.warning("Please wait a few seconds before running another entity scan.")
            return
        model = get_ai_model()
        raw_text = text or ""
        p = st.session_state.project

        try:
            ents = AnalysisEngine.extract_entities(raw_text, model)
        except Exception:
            logger.warning("Entity extraction failed for %s", label, exc_info=True)
            st.error("Entity scan failed. Please check your AI settings and try again.")
            return
        total_detected = len(ents)
        added = 0
        matched = 0
        flagged = 0
        suggested = 0
        memory_soft_added = 0
        memory_hard_added = 0
        world_threshold = float(
            st.session_state.get("world_bible_confidence_threshold", AppConfig.WORLD_BIBLE_CONFIDENCE)
        )
        world_threshold = max(0.0, min(1.0, world_threshold))
        soft_enabled = bool(st.session_state.get("memory_auto_soft_enabled", AppConfig.WORLD_MEMORY_AUTO_SOFT))
        hard_enabled = bool(st.session_state.get("memory_auto_hard_enabled", AppConfig.WORLD_MEMORY_AUTO_HARD))
        soft_threshold = float(st.session_state.get("memory_auto_soft_threshold", AppConfig.WORLD_MEMORY_SOFT_CONFIDENCE))
        hard_threshold = float(st.session_state.get("memory_auto_hard_threshold", AppConfig.WORLD_MEMORY_HARD_CONFIDENCE))
        soft_threshold = max(0.0, min(1.0, soft_threshold))
        hard_threshold = max(0.0, min(1.0, hard_threshold))

        for e in ents:
            name = (e.get("name") or "").strip()
            category = (e.get("category") or "").strip()
            if not name or not category:
                continue

            desc = (e.get("description") or "").strip()
            aliases = e.get("aliases") or []
            try:
                confidence = float(e.get("confidence")) if e.get("confidence") is not None else 0.0
            except (TypeError, ValueError):
                confidence = 0.0
            # Normalize confidence into 0..1 range.
            # Some model outputs return percentages (e.g., 82) instead of decimals (0.82).
            if confidence > 1.0:
                confidence = confidence / 100.0
            confidence = max(0.0, min(1.0, confidence))

            if confidence < world_threshold:
                existing = p.find_entity_match(
                    name,
                    category,
                    aliases=aliases if isinstance(aliases, list) else None,
                )
                if existing and desc:
                    _queue_world_bible_suggestion(
                        {
                            "type": "update",
                            "entity_id": existing.id,
                            "name": existing.name,
                            "category": existing.category,
                            "description": desc,
                            "aliases": aliases,
                            "confidence": confidence,
                            "source": label,
                        }
                    )
                else:
                    _queue_world_bible_suggestion(
                        {
                            "type": "new",
                            "name": name,
                            "category": Project._normalize_category(category),
                            "description": desc,
                            "aliases": aliases,
                            "confidence": confidence,
                            "source": label,
                        }
                    )
                flagged += 1
                continue

            ent, status = p.upsert_entity(
                name,
                category,
                desc,
                aliases=aliases if isinstance(aliases, list) else None,
                allow_merge=False,
                allow_alias=True,
            )
            if ent is None:
                continue

            if status == "created":
                added += 1
            else:
                matched += 1
                if desc:
                    _queue_world_bible_suggestion(
                        {
                            "type": "update",
                            "entity_id": ent.id,
                            "name": ent.name,
                            "category": ent.category,
                            "description": desc,
                            "aliases": aliases,
                            "confidence": confidence,
                            "source": label,
                        }
                    )
                    suggested += 1

            # High-confidence facts can be promoted into memory automatically.
            memory_line = ""
            if desc:
                memory_line = f"- {ent.name} ({ent.category}): {desc}"
            if memory_line and soft_enabled and confidence >= soft_threshold:
                updated_soft, changed = _append_unique_memory_line(p.memory_soft, memory_line)
                if changed:
                    p.memory_soft = updated_soft
                    memory_soft_added += 1
            if memory_line and hard_enabled and confidence >= hard_threshold:
                updated_hard, changed = _append_unique_memory_line(p.memory_hard, memory_line)
                if changed:
                    p.memory_hard = updated_hard
                    memory_hard_added += 1

        persist_project(p)
        st.session_state["last_entity_scan"] = time.time()

        if added > 0 or matched > 0 or flagged > 0 or memory_soft_added > 0 or memory_hard_added > 0:
            summary = [f"{added} new", f"{matched} existing", f"{flagged} flagged"]
            if suggested:
                summary.append(f"{suggested} suggested updates")
            if memory_soft_added:
                summary.append(f"{memory_soft_added} soft-memory")
            if memory_hard_added:
                summary.append(f"{memory_hard_added} hard-memory")
            if show_feedback:
                st.toast(f"World Bible updated ({', '.join(summary)})", icon="\u2705")
            return {
                "status": "updated",
                "added": added,
                "matched": matched,
                "flagged": flagged,
                "suggested": suggested,
                "memory_soft_added": memory_soft_added,
                "memory_hard_added": memory_hard_added,
                "total_detected": total_detected,
            }
        else:
            if total_detected > 0:
                if show_feedback:
                    st.toast("Detected entities, but they all matched existing entries.", icon="\u2139")
                return {
                    "status": "matched_only",
                    "added": 0,
                    "matched": matched,
                    "flagged": flagged,
                    "suggested": suggested,
                    "memory_soft_added": 0,
                    "memory_hard_added": 0,
                    "total_detected": total_detected,
                }
            else:
                if show_feedback:
                    st.toast("No entities detected in this text.", icon="\U0001F50E")
                return {
                    "status": "none",
                    "added": 0,
                    "matched": 0,
                    "flagged": 0,
                    "suggested": 0,
                    "memory_soft_added": 0,
                    "memory_hard_added": 0,
                    "total_detected": 0,
                }

    query_raw = st.query_params.get("page")
    query = query_raw[0] if isinstance(query_raw, list) and query_raw else query_raw
    valid_query_pages = {
        "home", "projects", "outline", "chapters", "world", "memory",
        "insights", "ai", "workspace", "privacy", "terms", "copyright",
        "cookie", "legal", "help",
    }
    if query in valid_query_pages:
        # Guard against redirect loops when browser/query state lags behind reruns.
        handled_query = st.session_state.get("_handled_query_page")
        current_page = st.session_state.get("page")
        st.query_params.clear()
        if query != current_page:
            st.session_state.page = query
        if handled_query != query or query != current_page:
            st.session_state["_handled_query_page"] = query
            st.rerun()
            return

    def render_ai_settings():
        # Add helpful guidance
        help_tooltip(
            text="MANTIS Studio supports multiple AI providers. Configure at least one provider to unlock AI-assisted features like outline generation, chapter writing, and world-building.",
            title="Getting Started with AI"
        )
        
        def refresh_all_models() -> None:
            st.cache_data.clear()
            refresh_models()
            refresh_openai_models()
            st.toast("Model list refreshed")

        def save_settings_action() -> None:
            save_app_settings()

        header_bar(
            "AI Settings",
            "Set up your AI providers and API keys. Start with one provider to begin writing.",
            pill="AI Settings",
        )
        
        if st.session_state.pop("ai_settings__flash", False):
            st.success("AI Settings opened. Configure a provider below to get started.")

        # Connection Status Overview
        groq_key, groq_source = get_effective_key("groq", st.session_state.get("user_id"))
        openai_key, openai_source = get_effective_key("openai", st.session_state.get("user_id"))
        
        with card_block("Connection Status"):
            st.caption("Current status of your AI provider connections")
            status_cols = st.columns(3)
            with status_cols[0]:
                st.metric("Groq", "Connected" if groq_key else "Not configured")
            with status_cols[1]:
                st.metric("OpenAI", "Connected" if openai_key else "Not configured")
            with status_cols[2]:
                current_model = get_ai_model() or "None"
                st.metric("Active Model", current_model)
        if not (groq_key or openai_key):
            st.info(
                "You can still create projects, build your world bible, write outlines, and draft chapters without AI. "
                "Connect a provider to enable generation, summaries, and assisted analysis."
            )

        # Provider Selection
        st.markdown("### Choose Your AI Provider")
        st.caption("Select which provider to use for AI-powered features like outlining, brainstorming, and text generation.")
        
        col1, col2 = st.columns(2)
        with col1:
            st.info("**Groq** - Fast & Free\n\nGreat for getting started quickly with no cost.")
        with col2:
            st.info("**OpenAI** - Advanced Models\n\nIncludes GPT-4 and other powerful models (paid).")
        
        provider_choice = st.radio(
            "Which provider would you like to use?",
            ["Groq", "OpenAI"],
            horizontal=True,
            index=0 if st.session_state.ai_provider == "groq" else 1,
            key="ai_provider_choice",
        )
        provider_value = "groq" if provider_choice == "Groq" else "openai"
        if provider_value != st.session_state.ai_provider:
            st.session_state.ai_provider = provider_value
            save_app_settings()
            st.rerun()

        # Provider Configuration Tabs
        st.markdown("### Configure Providers")
        tabs = st.tabs(["Groq", "OpenAI"])

        def _status_label(source: str) -> str:
            return {
                "session": "Using session key",
                "saved": "Using saved key",
                "default": "Using server default",
                "none": "No key set",
            }.get(source, "No key set")

        def _render_key_controls(provider: str) -> None:
            provider_label = _provider_label(provider)
            key_input_state = f"{provider}_key_input"
            init_state(key_input_state, "")
            
            st.markdown("#### API Key")
            _, source = get_effective_key(provider)
            key_status = _status_label(source)
            if source in ["session", "saved"]:
                st.success(f"{key_status}")
            else:
                st.info(f"{key_status}")
            
            key_value = st.text_input(
                f"Enter your {provider_label} API key",
                value=st.session_state.get(key_input_state, ""),
                type="password",
                placeholder="sk-...",
                help=f"Paste your API key here. It will only be stored for this session unless saved.",
                key=key_input_state,
                label_visibility="collapsed",
            )
            
            key_cols = st.columns(3)
            with key_cols[0]:
                if st.button(f"Apply Key", use_container_width=True, key=f"{provider}_session_key", type="primary"):
                    if key_value.strip():
                        set_session_key(provider, key_value)
                        queue_widget_update(key_input_state, "")
                        # Auto-fetch models after applying key
                        base_url = st.session_state.get(f"{provider}_base_url", "")
                        fetch_fn = fetch_openai_models if provider == "openai" else fetch_groq_models
                        models, _err = fetch_fn(base_url, key_value.strip())
                        if models:
                            st.session_state[f"{provider}_model_list"] = models
                            st.session_state[f"{provider}_model_tests"] = {}
                            st.toast(f"{provider_label} key activated  loaded {len(models)} models.")
                        else:
                            st.toast(f"{provider_label} key activated for this session.")
                        st.session_state[f"{provider}_connection_tested"] = False
                        st.session_state[f"{provider}_probe_attempted"] = False
                        st.toast("Use Check Connection to verify the key.")
                        st.rerun()
                    else:
                        st.warning("Please enter an API key first.")
            with key_cols[1]:
                if st.button("Check Connection", use_container_width=True, key=f"{provider}_check_connection_btn"):
                    candidate_key = key_value.strip()
                    if candidate_key:
                        set_session_key(provider, candidate_key)
                        queue_widget_update(key_input_state, "")
                    else:
                        candidate_key, _ = get_effective_key(provider, st.session_state.get("user_id"))
                    if not candidate_key:
                        st.warning("Enter or apply an API key first.")
                    else:
                        base_url = st.session_state.get(f"{provider}_base_url", "")
                        test_fn = test_openai_connection if provider == "openai" else test_groq_connection
                        fetch_fn = fetch_openai_models if provider == "openai" else fetch_groq_models
                        with st.spinner(f"Testing {provider_label} connection..."):
                            connection_ok = bool(test_fn(base_url, candidate_key))
                        st.session_state[f"{provider}_connection_tested"] = connection_ok
                        st.session_state[f"{provider}_probe_attempted"] = True
                        if connection_ok:
                            models, _err = fetch_fn(base_url, candidate_key)
                            if models:
                                st.session_state[f"{provider}_model_list"] = models
                            st.toast(f"{provider_label} connection verified.", icon="\u2705")
                        else:
                            st.toast(
                                f"{provider_label} connection failed. Check your key and base URL.",
                                icon="\u26A0",
                            )
                        st.rerun()
            with key_cols[2]:
                if st.button(f"Clear Key", use_container_width=True, key=f"{provider}_clear_session"):
                    clear_session_key(provider)
                    st.toast(f"{provider_label} key cleared.")
                    st.rerun()

        with tabs[1]:
            with card_block("OpenAI Configuration"):
                st.caption("OpenAI provides advanced models like GPT-4. Requires a paid API account.")
                
                _render_key_controls("openai")
                
                openai_key, _ = get_effective_key("openai", st.session_state.get("user_id"))
                if not openai_key:
                    st.info("**Getting Started:** [Create an OpenAI account](https://platform.openai.com/api-keys) to get your API key.")
                
                st.markdown("#### Model Selection")
                if st.session_state.openai_model_list:
                    models = st.session_state.openai_model_list
                    idx = 0
                    if st.session_state.openai_model in models:
                        idx = models.index(st.session_state.openai_model)
                    openai_model = st.selectbox(
                        "Choose a model",
                        models,
                        index=idx,
                        help="Select which OpenAI model to use for AI features",
                        label_visibility="collapsed",
                        key="openai_model_select",
                    )
                else:
                    openai_model = st.text_input(
                        "Model name",
                        value=st.session_state.openai_model,
                        help="Enter the model name manually, or apply your key to auto-fetch available models",
                        label_visibility="collapsed",
                        key="openai_model_manual_input",
                    )

                if openai_model != st.session_state.openai_model:
                    st.session_state.openai_model = openai_model
                    AppConfig.OPENAI_MODEL = openai_model
                    save_app_settings()
                    if openai_key:
                        st.session_state.openai_connection_tested = test_openai_connection(
                            st.session_state.openai_base_url,
                            openai_key,
                        )
                        if not st.session_state.openai_connection_tested:
                            st.toast("OpenAI connection test failed. Check your API key and base URL.", icon="\u26A0")
                    st.toast(f"OpenAI model set to {openai_model}")
                    st.rerun()
                
                with st.expander("Advanced Settings", expanded=False):
                    openai_url = st.text_input(
                        "API Base URL",
                        value=st.session_state.openai_base_url,
                        help="Change this only if using a custom OpenAI-compatible endpoint"
                    )
                    if openai_url != st.session_state.openai_base_url:
                        st.session_state.openai_base_url = openai_url
                        AppConfig.OPENAI_API_URL = openai_url

                if st.session_state.openai_model_list:
                    with st.expander("Advanced: Test All Models", expanded=False):
                        st.warning("This will test each model individually. It may take time and consume API credits.")
                        openai_test_all_cooldown = _cooldown_remaining("openai_test_models", 20)
                        if st.button(
                            "Run Full Model Test",
                            use_container_width=True,
                            disabled=bool(openai_test_all_cooldown) or not openai_key,
                            key="openai_test_all_models_btn"
                        ):
                            _mark_action("openai_test_models")
                            results = {}
                            total = len(st.session_state.openai_model_list)
                            progress = st.progress(0)
                            for i, model_name in enumerate(st.session_state.openai_model_list, start=1):
                                ok, error_message = test_openai_model(
                                    st.session_state.openai_base_url,
                                    openai_key,
                                    model_name,
                                )
                                results[model_name] = "" if ok else error_message
                                progress.progress(i / total)
                            st.session_state.openai_model_tests = results
                            failures = [m for m, err in results.items() if err]
                            if failures:
                                st.warning(f"{len(failures)} models failed. Check details below.")
                            else:
                                st.success("All models working correctly!")

                        if st.session_state.openai_model_tests:
                            st.markdown("**Test Results:**")
                            for model_name, error_message in sorted(
                                st.session_state.openai_model_tests.items()
                            ):
                                if error_message:
                                    st.error(f"{model_name}: {error_message}")
                                else:
                                    st.success(f"{model_name}")

        with tabs[0]:
            with card_block("Groq Configuration"):
                st.caption("Groq offers fast, free AI models. Great for getting started quickly.")
                
                _render_key_controls("groq")
                
                groq_key, _ = get_effective_key("groq", st.session_state.get("user_id"))
                if not groq_key:
                    st.info("**Getting Started:** [Get a free Groq API key](https://console.groq.com/keys) to begin.")
                
                st.markdown("#### Model Selection")
                if st.session_state.groq_model_list:
                    models = st.session_state.groq_model_list
                    idx = 0
                    if st.session_state.groq_model in models:
                        idx = models.index(st.session_state.groq_model)
                    groq_model = st.selectbox(
                        "Choose a model",
                        models,
                        index=idx,
                        help="Select which Groq model to use for AI features",
                        label_visibility="collapsed",
                        key="groq_model_select",
                    )
                else:
                    groq_model = st.text_input(
                        "Model name",
                        value=st.session_state.groq_model,
                        help="Enter the model name manually, or apply your key to auto-fetch available models",
                        label_visibility="collapsed",
                        key="groq_model_manual_input",
                    )

                if groq_model != st.session_state.groq_model:
                    st.session_state.groq_model = groq_model
                    AppConfig.DEFAULT_MODEL = groq_model
                    save_app_settings()
                    if groq_key:
                        st.session_state.groq_connection_tested = test_groq_connection(
                            st.session_state.groq_base_url,
                            groq_key,
                        )
                        if not st.session_state.groq_connection_tested:
                            st.toast("Groq connection test failed. Check your API key and base URL.", icon="\u26A0")
                    st.toast(f"Groq model set to {groq_model}")
                    st.rerun()
                
                with st.expander("Advanced Settings", expanded=False):
                    groq_url = st.text_input(
                        "API Base URL",
                        value=st.session_state.groq_base_url,
                        help="Change this only if using a custom Groq-compatible endpoint"
                    )
                    if groq_url != st.session_state.groq_base_url:
                        st.session_state.groq_base_url = groq_url
                        AppConfig.GROQ_API_URL = groq_url

                if st.session_state.groq_model_list:
                    with st.expander("Advanced: Test All Models", expanded=False):
                        st.warning("This will test each model individually and may take some time.")
                        groq_test_all_cooldown = _cooldown_remaining("groq_test_models", 20)
                        if st.button(
                            "Run Full Model Test",
                            use_container_width=True,
                            disabled=bool(groq_test_all_cooldown) or not groq_key,
                            key="groq_test_all_models_btn"
                        ):
                            _mark_action("groq_test_models")
                            results = {}
                            total = len(st.session_state.groq_model_list)
                            progress = st.progress(0)
                            for i, model_name in enumerate(st.session_state.groq_model_list, start=1):
                                ok, error_message = test_groq_model(
                                    st.session_state.groq_base_url,
                                    groq_key,
                                    model_name,
                                )
                                results[model_name] = "" if ok else error_message
                                progress.progress(i / total)
                            st.session_state.groq_model_tests = results
                            failures = [m for m, err in results.items() if err]
                            if failures:
                                st.warning(f"{len(failures)} models failed. Check details below.")
                            else:
                                st.success("All models working correctly!")

                        if st.session_state.groq_model_tests:
                            st.markdown("**Test Results:**")
                            for model_name, error_message in sorted(
                                st.session_state.groq_model_tests.items()
                            ):
                                if error_message:
                                    st.error(f"{model_name}: {error_message}")
                                else:
                                    st.success(f"{model_name}")

        # Save Settings Action
        st.markdown("---")
        save_cols = st.columns([2, 1])
        with save_cols[0]:
            st.markdown("### Save Your Configuration")
            st.caption("Save your provider settings and API keys for future sessions.")
        with save_cols[1]:
            if st.button(
                "Save Settings",
                type="primary",
                use_container_width=True,
                help="Save your AI provider configuration"
            ):
                save_settings_action()

    def render_legal_redirect():
        render_page_header(
            "All Policies",
            "Review policies, IP guidance, and acceptable use.",
            tag="Policies",
            key_prefix="legal_header",
        )

        legal_pages = [
            ("Terms of Service", "terms", "terms.md"),
            ("Privacy Policy", "privacy", "privacy.md"),
            ("Copyright", "copyright", "copyright.md"),
            ("Cookie Policy", "cookie", "cookie.md"),
            ("Brand & IP Clarity", "brand_ip", "brand_ip_clarity.md"),
            ("Trademark Path", "trademark", "trademark_path.md"),
            ("Contact", "contact_legal", "contact.md"),
        ]

        cols = st.columns(2)
        for idx, (label, key_suffix, filename) in enumerate(legal_pages):
            with cols[idx % 2]:
                path = LEGAL_DIR / filename
                if path.exists():
                    with st.expander(label):
                        st.markdown(_read_text_with_encoding_fallback(path))
                else:
                    st.caption(f"{label}  not available")

    def render_workspace_settings():
        render_page_header(
            "Workspace Settings",
            "Studio preferences, writing defaults, and canon confidence thresholds.",
            tag="Workspace",
            key_prefix="workspace_header",
        )
        with st.container(border=True):
            st.markdown("### Studio Preferences")
            col1, col2 = st.columns(2)
            with col1:
                auto_save = st.toggle(
                    "Auto-save changes",
                    value=bool(st.session_state.get("auto_save", True)),
                    key="workspace_auto_save",
                )
                daily_goal = st.number_input(
                    "Daily word goal",
                    min_value=100,
                    max_value=10000,
                    step=100,
                    value=int(st.session_state.get("daily_word_goal", 1000)),
                    key="workspace_daily_goal",
                )
            with col2:
                weekly_goal = st.number_input(
                    "Weekly sessions goal",
                    min_value=1,
                    max_value=14,
                    step=1,
                    value=int(st.session_state.get("weekly_sessions_goal", 4)),
                    key="workspace_weekly_goal",
                )
                focus_minutes = st.number_input(
                    "Focus timer (minutes)",
                    min_value=5,
                    max_value=180,
                    step=5,
                    value=int(st.session_state.get("focus_minutes", 25)),
                    key="workspace_focus_minutes",
                )
            st.session_state.auto_save = bool(auto_save)
            st.session_state.daily_word_goal = int(daily_goal)
            st.session_state.weekly_sessions_goal = int(weekly_goal)
            st.session_state.focus_minutes = int(focus_minutes)

        with st.container(border=True):
            st.markdown("### Canon Confidence Rules")
            st.caption("Set confidence gates for automatic World Bible and Memory updates.")
            w_col1, w_col2 = st.columns(2)
            with w_col1:
                wb_threshold = st.slider(
                    "World Bible auto-apply threshold",
                    min_value=0.50,
                    max_value=0.99,
                    value=float(st.session_state.get("world_bible_confidence_threshold", AppConfig.WORLD_BIBLE_CONFIDENCE)),
                    step=0.01,
                    key="workspace_world_bible_threshold",
                )
                st.session_state["world_bible_confidence_threshold"] = float(wb_threshold)
            with w_col2:
                st.session_state["memory_auto_soft_enabled"] = st.toggle(
                    "Auto-add to Soft Guidelines",
                    value=bool(st.session_state.get("memory_auto_soft_enabled", AppConfig.WORLD_MEMORY_AUTO_SOFT)),
                    key="workspace_memory_soft_enabled",
                )
                st.session_state["memory_auto_soft_threshold"] = st.slider(
                    "Soft memory threshold",
                    min_value=0.75,
                    max_value=0.99,
                    value=float(st.session_state.get("memory_auto_soft_threshold", AppConfig.WORLD_MEMORY_SOFT_CONFIDENCE)),
                    step=0.01,
                    key="workspace_memory_soft_threshold",
                )
                st.session_state["memory_auto_hard_enabled"] = st.toggle(
                    "Auto-add to Hard Canon Rules",
                    value=bool(st.session_state.get("memory_auto_hard_enabled", AppConfig.WORLD_MEMORY_AUTO_HARD)),
                    key="workspace_memory_hard_enabled",
                )
                st.session_state["memory_auto_hard_threshold"] = st.slider(
                    "Hard memory threshold",
                    min_value=0.90,
                    max_value=0.995,
                    value=float(st.session_state.get("memory_auto_hard_threshold", AppConfig.WORLD_MEMORY_HARD_CONFIDENCE)),
                    step=0.005,
                    key="workspace_memory_hard_threshold",
                )

        if st.button("Save Workspace Settings", type="primary", use_container_width=True, key="workspace_save"):
            data = load_app_config()
            data.update({
                "ui_theme": st.session_state.get("ui_theme", "Dark"),
                "auto_save": bool(st.session_state.get("auto_save", True)),
                "daily_word_goal": int(st.session_state.get("daily_word_goal", 1000)),
                "weekly_sessions_goal": int(st.session_state.get("weekly_sessions_goal", 4)),
                "focus_minutes": int(st.session_state.get("focus_minutes", 25)),
                "world_bible_confidence_threshold": float(st.session_state.get("world_bible_confidence_threshold", AppConfig.WORLD_BIBLE_CONFIDENCE)),
                "memory_auto_soft_enabled": bool(st.session_state.get("memory_auto_soft_enabled", AppConfig.WORLD_MEMORY_AUTO_SOFT)),
                "memory_auto_soft_threshold": float(st.session_state.get("memory_auto_soft_threshold", AppConfig.WORLD_MEMORY_SOFT_CONFIDENCE)),
                "memory_auto_hard_enabled": bool(st.session_state.get("memory_auto_hard_enabled", AppConfig.WORLD_MEMORY_AUTO_HARD)),
                "memory_auto_hard_threshold": float(st.session_state.get("memory_auto_hard_threshold", AppConfig.WORLD_MEMORY_HARD_CONFIDENCE)),
            })
            save_app_config(data)
            st.toast("Workspace settings saved.")

    # Render enhanced sidebar
    from app.layout.enhanced_sidebar import render_enhanced_sidebar
    
    def save_project_callback():
        """Callback for saving project from sidebar"""
        p = st.session_state.project
        if persist_project(p, action="save"):
            return True
        return False
    
    def close_project_callback():
        """Callback for closing project from sidebar"""
        p = st.session_state.project
        save_p()
        # Clear the last_project_path from config to prevent auto-reload
        config = load_app_config()
        config.pop("last_project_path", None)
        save_app_config(config)
        st.session_state.project = None
        st.session_state.page = "home"
        st.rerun()

    def on_sidebar_theme_change(theme_name: str):
        """Persist sidebar theme changes immediately."""
        resolved = "Light" if str(theme_name).strip().lower() == "light" else "Dark"
        st.session_state.ui_theme = resolved
        config = load_app_config()
        config["ui_theme"] = resolved
        save_app_config(config)

    if not st.session_state.get("sidebar_collapsed", False):
        render_enhanced_sidebar(
            version=AppConfig.VERSION,
            project=st.session_state.project,
            current_page=st.session_state.get("page", "home"),
            world_focus=st.session_state.get("world_focus_tab", ""),
            key_scope=key_scope,
            slugify=_slugify,
            save_project_callback=save_project_callback,
            close_project_callback=close_project_callback,
            on_theme_change=on_sidebar_theme_change,
        )
    else:
        st.html(
            """
            <style>
            section[data-testid="stSidebar"] { display: none !important; }
            </style>
            """
        )
        if st.button("Show Sidebar", key="sidebar_expand", type="secondary", use_container_width=False):
            st.session_state["sidebar_collapsed"] = False
            st.session_state["_scroll_top_nonce"] = int(st.session_state.get("_scroll_top_nonce", 0)) + 1
            st.rerun()

    def render_home():
        active_dir = get_active_projects_dir()
        recent_projects = _load_recent_projects(active_dir, st.session_state.projects_refresh_token)
        if recent_projects and st.session_state.get("first_run", True):
            st.session_state.first_run = False
        dashboard_banner_path = ASSETS_DIR / "branding" / "mantis_banner_dark.png"
        try:
            dashboard_visual = (
                dashboard_banner_path.read_bytes() if dashboard_banner_path.exists() else None
            )
        except Exception:
            dashboard_visual = None

        has_outline = any((p["meta"].get("outline") or "").strip() for p in recent_projects)
        has_chapter = any(
            (c.get("word_count") or 0) > 0
            for p in recent_projects
            for c in (p["meta"].get("chapters") or {}).values()
        )

        active_project = st.session_state.project
        recent_snapshot = _project_snapshot(recent_projects[0]["meta"]) if recent_projects else None

        project_title = (
            (active_project.title if active_project else None)
            or (recent_snapshot or {}).get("title")
            or "Your next story"
        )
        milestone_label = "Create/Advance Next Narrative Milestone"
        weekly_goal = max(1, int(st.session_state.weekly_sessions_goal))
        weekly_count = _weekly_activity_count()
        canon_icon, canon_label = get_canon_health()
        groq_key, _ = get_effective_key("groq", st.session_state.get("user_id"))
        openai_key, _ = get_effective_key("openai", st.session_state.get("user_id"))
        project_health_completed = sum(
            [bool(recent_projects), bool(has_outline), bool(has_chapter), canon_icon != "RISK"]
        )
        project_health_percent = int((project_health_completed / 4) * 100)
        latest_chapter_label = "You last worked on a chapter recently"
        latest_chapter_index = None
        latest_chapter_id = None

        if active_project and getattr(active_project, "chapters", None):
            ch = max(
                active_project.chapters.values(),
                key=lambda c: (c.modified_at or c.created_at or 0),
            )
            latest_chapter_index = ch.index
            latest_chapter_id = ch.id
            latest_chapter_label = f"Latest: Chapter {ch.index} - {ch.title}"

        primary_label = milestone_label
        primary_target = "projects"
        if canon_icon == "RISK":
            primary_label = "Fix story issues"
            primary_target = "world"
        elif has_chapter and latest_chapter_index:
            primary_label = f"Continue Chapter {latest_chapter_index}"
            primary_target = "chapters"
        elif has_outline:
            primary_label = "Build your outline"
            primary_target = "outline"

        def open_recent_project(target: str, focus_tab: Optional[str] = None) -> None:
            if not recent_projects and not st.session_state.project:
                st.session_state.page = "projects"
                st.toast("Create or import a project to unlock this module.")
                st.rerun()
            if recent_projects and not st.session_state.project:
                loaded = load_project_safe(recent_projects[0]["path"], context="recent project")
                if not loaded:
                    st.session_state.page = "projects"
                    st.toast("Select a project to continue.")
                    st.rerun()
                st.session_state.project = loaded
            if focus_tab:
                st.session_state.world_focus_tab = focus_tab
            st.session_state.page = target
            st.rerun()

        def open_export() -> None:
            export_path = None
            if st.session_state.project and st.session_state.project.filepath:
                export_path = st.session_state.project.filepath
            elif recent_projects:
                export_path = recent_projects[0]["path"]
            if export_path:
                st.session_state.export_project_path = export_path
                st.session_state.page = "projects"
                st.rerun()
            else:
                st.session_state.page = "projects"
                st.toast("Select a project to export.")
                st.rerun()

        def open_ai_settings() -> None:
            st.session_state.ai_settings__flash = True
            st.session_state.page = "ai"
            st.rerun()

        def open_primary_cta() -> None:
            if primary_target == "chapters" and latest_chapter_id:
                st.session_state.curr_chap_id = latest_chapter_id
                st.session_state.active_chapter_id = latest_chapter_id
            open_recent_project(primary_target)

        def open_new_project() -> None:
            st.session_state.page = "projects"
            st.rerun()

        today = datetime.date.today()
        last_action_ts = st.session_state.get("last_action_ts")
        last_action_day = (
            datetime.date.fromtimestamp(last_action_ts) if last_action_ts else None
        )
        last_action_label = (st.session_state.get("last_action") or "").lower()
        ai_action_tokens = ("analysis", "insights", "memory", "semantic", "entity", "ai")
        ai_ops_today = (
            1
            if last_action_day == today and any(token in last_action_label for token in ai_action_tokens)
            else 0
        )
        system_mode = "Cloud" if (groq_key or openai_key) else "Local"
        autosave_status = "Auto-saving" if st.session_state.get("auto_save", True) else "Saved"
        ai_connected = bool(groq_key or openai_key)
        project_entity_count = len(getattr(active_project, "world_db", {}) or {}) if active_project else 0
        world_recent_names: List[str] = []
        if active_project and getattr(active_project, "world_db", None):
            recent_entities = sorted(
                active_project.world_db.values(),
                key=lambda e: (getattr(e, "modified_at", None) or getattr(e, "created_at", 0) or 0),
                reverse=True,
            )
            world_recent_names = [e.name for e in recent_entities[:3] if (e.name or "").strip()]

        def _open_project(project_path: str, target: str = "chapters") -> None:
            loaded = load_project_safe(project_path, context="project")
            if loaded:
                st.session_state.project = loaded
                st.session_state.page = target
                st.rerun()

        # =====================================================================
        # DASHBOARD LAYOUT
        # =====================================================================
        from app.views.dashboard_workspace import render_dashboard_workspace

        st.markdown("## Dashboard")
        st.caption("Outcome-focused workspace for planning, drafting, canon, and export.")
        if dashboard_visual:
            st.image(dashboard_visual, use_container_width=True)
        st.caption(f"{autosave_status} - Canon {canon_label} - Mode: {system_mode}")
        st.html("<div style='height: 8px;'></div>")
        if not recent_projects:
            render_welcome_banner("home")
            st.html("<div style='height: 8px;'></div>")
        elif not ai_connected:
            st.warning("You can plan and write without AI. Connect a provider in AI Settings to unlock generation, summaries, and scans.")

        total_words = 0
        for project_entry in recent_projects:
            chapters = (project_entry.get("meta", {}).get("chapters") or {}).values()
            total_words += sum(int((ch or {}).get("word_count") or 0) for ch in chapters)

        render_dashboard_workspace(
            recent_projects=recent_projects,
            project_title=project_title,
            latest_chapter_label=latest_chapter_label,
            has_outline=has_outline,
            has_chapter=has_chapter,
            canon_icon=canon_icon,
            canon_label=canon_label,
            system_mode=system_mode,
            ai_ops_today=ai_ops_today,
            weekly_count=weekly_count,
            weekly_goal=weekly_goal,
            project_entity_count=project_entity_count,
            world_recent_names=world_recent_names,
            total_words=total_words,
            project_health_percent=project_health_percent,
            ai_connected=bool(groq_key or openai_key),
            on_resume_writing=lambda: open_recent_project("chapters"),
            on_open_outline=lambda: open_recent_project("outline"),
            on_new_project=open_new_project,
            on_open_world=lambda: open_recent_project("world"),
            on_open_memory=lambda: open_recent_project("memory"),
            on_open_insights=lambda: open_recent_project("insights"),
            on_open_project=_open_project,
        )

        with st.container(border=True):
            st.markdown("#### Data confidence")
            st.caption("Ownership and save behavior for long-form work.")
            st.markdown(f"**Autosave:** {autosave_status}")
            st.markdown(f"**Project storage:** {system_mode} workspace")
            st.caption("You keep project files and can export anytime from Projects.")
            action_cols = st.columns(2)
            with action_cols[0]:
                if st.button("AI Settings", use_container_width=True, key="home_data_ai_settings"):
                    st.session_state.page = "ai"
                    st.rerun()
            with action_cols[1]:
                if st.button("Open Projects", use_container_width=True, key="home_data_open_projects"):
                    st.session_state.page = "projects"
                    st.rerun()

        return

        feature_groups = [
            (
                "Intelligence",
                [
                    ("Narrative Analysis", "Analyze structure and milestones.", lambda: open_recent_project("outline")),
                    ("Semantic Tools", "Review memory and semantic coherence.", lambda: open_recent_project("memory")),
                    ("Entity Extraction", "Manage entities from the World Bible.", lambda: open_recent_project("world")),
                ],
            ),
            (
                "Writing",
                [
                    ("Editor", "Draft and revise scenes.", lambda: open_recent_project("chapters")),
                    ("Rewrite", "Open editor tools for rewrites.", lambda: open_recent_project("chapters")),
                    ("Tone Modulation", "Refine voice and tone from editor workflows.", lambda: open_recent_project("chapters")),
                ],
            ),
            (
                "Insights",
                [
                    ("Reports", "Prepare export-ready story reports.", open_export),
                    ("Analytics", "Open canon analytics and health.", lambda: open_recent_project("insights")),
                    ("Data Viewer", "Inspect project metadata and structure.", lambda: open_recent_project("projects")),
                ],
            ),
            (
                "System",
                [
                    ("Settings", "Manage AI providers and keys.", open_ai_settings),
                    ("Configuration", "Adjust themes and workspace preferences.", open_ai_settings),
                ],
            ),
        ]

        for group_idx, (group_name, features) in enumerate(feature_groups):
            render_feature_group(
                group_name,
                features,
                expanded=(group_idx == 0),
                key_prefix="dashboard"
            )

        # =====================================================================
        # AI PROVIDER CONNECTION STATUS
        # =====================================================================
        st.html("<div style='height: 24px;'></div>")
        
        if not groq_key or not openai_key:
            with card_block("Connect your AI providers", "Unlock generation, summaries, and entity tools with API access."):
                cta_left, cta_right = st.columns(2)
                with cta_left:
                    st.link_button("Create Groq Account", "https://console.groq.com/keys", use_container_width=True)
                with cta_right:
                    st.link_button(
                        "Create OpenAI Account",
                        "https://platform.openai.com/api-keys",
                        use_container_width=True,
                    )
        else:
            with card_block("AI providers connected", "Your AI providers are configured and ready to use."):
                cta_left, cta_right = st.columns(2)
                with cta_left:
                    if st.button("Manage AI Settings", use_container_width=True, key="dashboard__ai_connected_settings"):
                        open_ai_settings()
                with cta_right:
                    st.caption("Groq and OpenAI are active.")

        add_divider_with_spacing(top=2, bottom=2)
        events = []
        if st.session_state.get("last_action"):
            ts = st.session_state.get("last_action_ts")
            stamp = time.strftime("%H:%M", time.localtime(ts)) if ts else "now"
            events.append((stamp, "Info", st.session_state.get("last_action")))
        if recent_projects:
            events.append(("Today", "Project", f"{len(recent_projects)} workspaces available"))
        events.append(("Live", "System", f"{system_mode} mode - {autosave_status.lower()}"))
        render_activity_feed(events)


    def render_projects():
        render_welcome_banner("projects")
        active_dir = get_active_projects_dir()
        recent_projects = _load_recent_projects(active_dir, st.session_state.projects_refresh_token)

        def open_latest_project() -> None:
            if not recent_projects:
                st.toast("Complete the form below to create your first project.")
                return
            loaded = load_project_safe(recent_projects[0]["path"], context="recent project")
            if not loaded:
                return
            st.session_state.project = loaded
            st.session_state.page = "chapters"
            st.session_state.first_run = False
            st.rerun()

        render_page_header(
            "Projects",
            "Create, import, and manage your story worlds.",
            tag="Workspace",
            key_prefix="projects_header",
        )

        from app.utils.helpers import ai_connection_warning
        ai_connection_warning(st)

        with st.container(border=True):
            st.markdown("### Project manager")
            pm_cols = st.columns(3)
            with pm_cols[0]:
                st.metric("Projects", len(recent_projects))
            with pm_cols[1]:
                total_chapters = sum(len((p.get("meta", {}).get("chapters") or {})) for p in recent_projects)
                st.metric("Chapters", total_chapters)
            with pm_cols[2]:
                total_words = 0
                for entry in recent_projects:
                    chapters = (entry.get("meta", {}).get("chapters") or {}).values()
                    total_words += sum(int((c or {}).get("word_count") or 0) for c in chapters)
                st.metric("Total words", f"{total_words:,}")

        if recent_projects:
            with st.container(border=True):
                st.markdown("### Edit project metadata")
                options = []
                option_to_path = {}
                for entry in recent_projects:
                    meta = entry.get("meta", {})
                    label = meta.get("title") or os.path.basename(entry.get("path", "Untitled"))
                    options.append(label)
                    option_to_path[label] = entry.get("path")

                default_label = options[0]
                if st.session_state.project and st.session_state.project.filepath:
                    for label, pth in option_to_path.items():
                        if os.path.normcase(os.path.normpath(pth or "")) == os.path.normcase(os.path.normpath(st.session_state.project.filepath or "")):
                            default_label = label
                            break
                selected_label = st.selectbox(
                    "Project",
                    options,
                    index=options.index(default_label),
                    key="projects_meta_target_label",
                )
                target_path = option_to_path.get(selected_label)
                target_project = None
                if target_path:
                    try:
                        target_project = Project.load(target_path)
                    except Exception:
                        target_project = None

                if target_project:
                    if st.session_state.get("projects_meta_target_path") != target_path:
                        st.session_state["projects_meta_target_path"] = target_path
                        st.session_state["projects_meta_title_input"] = target_project.title or ""
                        st.session_state["projects_meta_genre_input"] = target_project.genre or ""
                        st.session_state["projects_meta_author_input"] = target_project.author or ""

                    c_meta1, c_meta2, c_meta3 = st.columns([2, 1, 1])
                    with c_meta1:
                        st.text_input("Title", key="projects_meta_title_input")
                    with c_meta2:
                        st.text_input("Genre", key="projects_meta_genre_input")
                    with c_meta3:
                        st.text_input("Author", key="projects_meta_author_input")

                    save_meta_cols = st.columns(2)
                    with save_meta_cols[0]:
                        if st.button("Save metadata", type="primary", use_container_width=True, key="projects_meta_save"):
                            target_project.title = (st.session_state.get("projects_meta_title_input") or "").strip() or target_project.title
                            target_project.genre = (st.session_state.get("projects_meta_genre_input") or "").strip()
                            target_project.author = (st.session_state.get("projects_meta_author_input") or "").strip()
                            if persist_project(target_project, action="save"):
                                if (
                                    st.session_state.project
                                    and os.path.normcase(os.path.normpath(st.session_state.project.filepath or ""))
                                    == os.path.normcase(os.path.normpath(target_path or ""))
                                ):
                                    st.session_state.project = target_project
                                _bump_projects_refresh()
                                st.toast("Project metadata saved.")
                                st.rerun()
                    with save_meta_cols[1]:
                        if st.button("Open this project", use_container_width=True, key="projects_meta_open"):
                            loaded = load_project_safe(target_path, context="project")
                            if loaded:
                                st.session_state.project = loaded
                                st.session_state.page = "chapters"
                                st.rerun()

        section_header(
            "Start a new project",
            "Set a title, genre, and author details to build your base.",
        )
        with st.container(border=True):
            genre_prefill = st.session_state.pop("new_project_genre_prefill", "")
            if "new_project_title_input" not in st.session_state:
                st.session_state["new_project_title_input"] = ""
            if "new_project_genre_input" not in st.session_state:
                st.session_state["new_project_genre_input"] = genre_prefill or ""
            elif genre_prefill and not (st.session_state.get("new_project_genre_input") or "").strip():
                st.session_state["new_project_genre_input"] = genre_prefill
            if "new_project_author_input" not in st.session_state:
                st.session_state["new_project_author_input"] = ""

            with st.form("new_project_form", clear_on_submit=False):
                c1, c2 = st.columns([2, 1])
                with c1:
                    st.text_input(
                        "Project Title",
                        key="new_project_title_input",
                        placeholder="e.g., The Chronicle of Ash",
                        help="Type your title and press Enter or Create Project. If title is empty, MANTIS auto-generates one."
                    )
                with c2:
                    st.text_input(
                        "Genre",
                        key="new_project_genre_input",
                        placeholder="e.g., Dark Fantasy, Sci-Fi Noir",
                        help="Type your genre(s). If genre is empty, MANTIS auto-generates it."
                    )
                st.text_input(
                    "Author Name (Optional)",
                    key="new_project_author_input",
                    placeholder="Your name",
                    help="Add your name or pen name for attribution. Leave blank if you prefer."
                )
                submitted = st.form_submit_button(
                    "Create Project",
                    type="primary",
                    use_container_width=True,
                    key="projects_create_submit",
                )
                if submitted:
                    t = (st.session_state.get("new_project_title_input") or "").strip()
                    g = (st.session_state.get("new_project_genre_input") or "").strip()
                    a = (st.session_state.get("new_project_author_input") or "").strip()
                    title, genre, genre_list, ai_used = _resolve_project_seed(t, g, a)
                    st.session_state["project_seed_output"] = {
                        "title": title,
                        "genres": genre_list,
                        "ai_used": ai_used,
                    }
                    p = Project.create(
                        title,
                        author=a,
                        genre=genre,
                        storage_dir=get_active_projects_dir(),
                    )
                    persist_project(p, action="create_project")
                    _bump_projects_refresh()
                    st.session_state.project = p
                    st.session_state.page = "outline"
                    st.session_state.first_run = False
                    st.rerun()

        section_header(
            "Import an existing draft",
            "Upload TXT/MD/DOCX/PDF. MANTIS routes content into outline and chapters intelligently.",
        )

        def _extract_text_from_upload(uploaded_file) -> str:
            name = (uploaded_file.name or "").lower()
            raw = uploaded_file.read()
            if name.endswith(".txt") or name.endswith(".md"):
                return raw.decode("utf-8", errors="replace")
            if name.endswith(".docx"):
                try:
                    import docx  # type: ignore
                    from io import BytesIO
                    doc = docx.Document(BytesIO(raw))
                    return "\n".join(p.text for p in doc.paragraphs if (p.text or "").strip())
                except Exception:
                    try:
                        import zipfile
                        import xml.etree.ElementTree as ET
                        from io import BytesIO
                        with zipfile.ZipFile(BytesIO(raw)) as zf:
                            xml_data = zf.read("word/document.xml")
                        root = ET.fromstring(xml_data)
                        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                        text_nodes = root.findall(".//w:t", ns)
                        return " ".join((n.text or "") for n in text_nodes if (n.text or "").strip())
                    except Exception:
                        return ""
            if name.endswith(".pdf"):
                try:
                    from io import BytesIO
                    import pypdf  # type: ignore
                    reader = pypdf.PdfReader(BytesIO(raw))
                    pages = [page.extract_text() or "" for page in reader.pages]
                    return "\n".join(pages).strip()
                except Exception:
                    return ""
            return ""

        def _route_import_into_project(project: "Project", text: str) -> None:
            clean_text = (text or "").strip()
            if not clean_text:
                return
            chapter_markers = re.findall(r"(?im)^\\s*(chapter|part)\\s+\\d+", clean_text)
            if chapter_markers:
                project.import_text_file(clean_text)
            else:
                active_provider = _normalize_provider(st.session_state.get("ai_provider", "groq"))
                active_key, _ = get_effective_key(active_provider, st.session_state.get("user_id"))
                if active_key and get_ai_model():
                    classify_prompt = (
                        "Classify this manuscript excerpt.\n"
                        "Return ONLY one label: OUTLINE or CHAPTER.\n\n"
                        f"TEXT:\n{clean_text[:5000]}"
                    )
                    guess = (AIEngine(provider=active_provider).generate(classify_prompt, get_ai_model()).get("text", "") or "").strip().upper()
                    if "OUTLINE" in guess:
                        project.outline = clean_text
                    else:
                        project.add_chapter("Imported Chapter 1", clean_text)
                else:
                    # Without AI, default to creating a starter chapter unless it looks list-like.
                    if re.search(r"(?im)^\\s*[-*]\\s+", clean_text) and len(clean_text.splitlines()) > 8:
                        project.outline = clean_text
                    else:
                        project.add_chapter("Imported Chapter 1", clean_text)

        with st.container(border=True):
            uf = st.file_uploader("Upload file", type=["txt", "md", "docx", "pdf"])
            if uf:
                max_bytes = AppConfig.MAX_UPLOAD_MB * 1024 * 1024
                uf_size = getattr(uf, "size", None)
                if uf_size and uf_size > max_bytes:
                    st.error(
                        f"File too large. Max size is {AppConfig.MAX_UPLOAD_MB} MB."
                    )
                else:
                    txt = _extract_text_from_upload(uf)
                    if not (txt or "").strip():
                        st.error("Could not extract readable text from this file. Try TXT/MD, or a text-based DOCX/PDF.")
                        return
                    if st.button("Import & Analyze", use_container_width=True, type="primary", key="projects_import_analyze"):
                        try:
                            # Show step indicator for import process
                            step_indicator(
                                steps=["Upload File", "Create Project", "Analyze Content", "Generate Outline"],
                                current_step=1
                            )
                            
                            p = Project.create("Imported Project", storage_dir=get_active_projects_dir())
                            _route_import_into_project(p, txt)
                            
                            # Update step indicator
                            step_indicator(
                                steps=["Upload File", "Create Project", "Analyze Content", "Generate Outline"],
                                current_step=2,
                                completed_steps=[0, 1]
                            )
                            
                            active_provider = _normalize_provider(st.session_state.get("ai_provider", "groq"))
                            active_key, _ = get_effective_key(active_provider, st.session_state.get("user_id"))
                            if active_key and get_ai_model():
                                # Enhanced loading state
                                loading_state(
                                    message="Analyzing your document...",
                                    subtext="This may take 10-30 seconds depending on length",
                                    show_spinner=True
                                )
                                p.outline = StoryEngine.reverse_engineer_outline(p, get_ai_model())
                                
                                # Show success feedback
                                feedback_message(
                                    "Document imported and outline generated successfully!",
                                    type="success"
                                )
                            else:
                                st.warning(
                                    f"Add a {_provider_label(active_provider)} API key and model to auto-generate an outline."
                                )
                            persist_project(p, action="create_project")
                        except Exception:
                            logger.warning("Import failed for uploaded draft", exc_info=True)
                            st.error("Import failed. Please check the file and try again.")
                            return
                        _bump_projects_refresh()
                        st.session_state.project = p
                        st.session_state.page = "outline"
                        st.session_state.first_run = False
                        st.rerun()

        section_header(
            "Your projects",
            "Open, export, or clean up older drafts.",
        )
        with st.container(border=True):
            if not recent_projects:
                st.info("No projects yet. Start a new project above to get going.")
            else:
                for project_entry in recent_projects[:30]:
                    full = project_entry["path"]
                    try:
                        meta = project_entry["meta"]
                        filename = os.path.basename(full)
                        title = meta.get("title") or filename
                        genre = meta.get("genre") or "-"
                        row_key = uuid.uuid5(uuid.NAMESPACE_URL, full).hex
                        row1, row2, row3, row4 = st.columns([5, 2, 1.5, 1])
                        with row1:
                            if st.button(f"{title}", key=f"proj_open_{row_key}", use_container_width=True):
                                loaded = load_project_safe(full, context="project")
                                if loaded:
                                    st.session_state.project = loaded
                                    st.session_state.page = "chapters"
                                    st.session_state.first_run = False
                                    st.rerun()
                        with row2:
                            st.caption(genre)
                        with row3:
                            if st.button("Export", key=f"proj_export_{row_key}", use_container_width=True):
                                st.session_state.export_project_path = full
                                st.rerun()
                        with row4:
                            if st.button("Delete", key=f"proj_delete_{row_key}", use_container_width=True):
                                st.session_state.delete_project_path = full
                                st.session_state.delete_project_title = title
                                st.rerun()
                    except Exception:
                        logger.warning("Failed to load project metadata: %s", full, exc_info=True)

        if st.session_state.delete_project_path:
            with st.container(border=True):
                title = st.session_state.delete_project_title or "this project"
                st.warning(f"Delete **{title}**? This cannot be undone.")
                d1, d2 = st.columns(2)
                with d1:
                    if st.button("Confirm delete", type="primary", use_container_width=True, key="projects_delete_confirm"):
                        target_path = st.session_state.delete_project_path
                        deleted, msg = Project.delete_file(target_path)
                        target_basename = os.path.basename(target_path or "")
                        target_id = target_basename.split("_", 1)[0] if "_" in target_basename else ""
                        if deleted:
                            if (
                                st.session_state.project
                                and (
                                    os.path.normcase(os.path.normpath(st.session_state.project.filepath or ""))
                                    == os.path.normcase(os.path.normpath(target_path or ""))
                                    or (target_id and st.session_state.project.id == target_id)
                                )
                            ):
                                st.session_state.project = None
                                st.session_state.page = "projects"
                            st.session_state.delete_project_path = None
                            st.session_state.delete_project_title = None
                            _bump_projects_refresh()
                            st.toast("Project deleted.")
                            st.rerun()
                        else:
                            st.error(f"Delete failed: {msg}")
                            st.session_state.delete_project_path = None
                            st.session_state.delete_project_title = None
                with d2:
                    if st.button("Cancel", use_container_width=True, key="projects_delete_cancel"):
                        st.session_state.delete_project_path = None
                        st.session_state.delete_project_title = None
                        st.rerun()

        export_path = st.session_state.get("export_project_path")
        if export_path:
            with st.container(border=True):
                try:
                    export_project = Project.load(export_path)
                except Exception:
                    st.error("Export failed. Unable to load project.")
                    st.session_state.export_project_path = None
                else:
                    st.markdown("### Export Project")
                    st.caption("Download as Markdown, Text, JSON, or DOCX.")
                    format_choice = st.selectbox(
                        "Format",
                        ["Markdown (.md)", "Text (.txt)", "JSON (.json)", "DOCX (.docx)"],
                        key="projects_inline_export_format",
                    )
                    if format_choice.startswith("Markdown"):
                        st.download_button(
                            "Download",
                            project_to_markdown(export_project),
                            file_name=f"{export_project.title}.md",
                            use_container_width=True,
                        )
                    elif format_choice.startswith("Text"):
                        st.download_button(
                            "Download",
                            project_to_text(export_project),
                            file_name=f"{export_project.title}.txt",
                            use_container_width=True,
                        )
                    elif format_choice.startswith("JSON"):
                        st.download_button(
                            "Download",
                            project_to_json(export_project),
                            file_name=f"{export_project.title}.json",
                            use_container_width=True,
                            mime="application/json",
                        )
                    else:
                        try:
                            docx_bytes = project_to_docx_bytes(export_project)
                        except Exception:
                            st.error("DOCX export requires python-docx support.")
                        else:
                            st.download_button(
                                "Download",
                                docx_bytes,
                                file_name=f"{export_project.title}.docx",
                                use_container_width=True,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                    if st.button("Close export", use_container_width=True):
                        st.session_state.export_project_path = None
                        st.rerun()


    def render_export():
        render_page_header(
            "Export",
            "Export your project as markdown for editors or collaborators.",
            tag="Export",
            key_prefix="export_header",
        )

        export_project = None
        export_path = st.session_state.get("export_project_path")
        if export_path:
            try:
                export_project = Project.load(export_path)
            except Exception:
                st.error("Export failed. Unable to load project.")
                st.session_state.export_project_path = None
                return
        elif st.session_state.project:
            export_project = st.session_state.project
            if not export_project.filepath:
                export_project.save()

        if not export_project:
            render_no_project_state("Export", "empty_export")
            return

        with st.container(border=True):
            st.markdown(f"### {export_project.title}")
            st.caption("Download your complete manuscript in your preferred format.")
            format_choice = st.selectbox(
                "Export format",
                ["Markdown (.md)", "Text (.txt)", "JSON (.json)", "DOCX (.docx)"],
                key="export_page_format",
            )
            if format_choice.startswith("Markdown"):
                st.download_button(
                    "Download",
                    project_to_markdown(export_project),
                    file_name=f"{export_project.title}.md",
                    use_container_width=True,
                    help="Download as Markdown.",
                )
            elif format_choice.startswith("Text"):
                st.download_button(
                    "Download",
                    project_to_text(export_project),
                    file_name=f"{export_project.title}.txt",
                    use_container_width=True,
                    help="Download as plain text.",
                )
            elif format_choice.startswith("JSON"):
                st.download_button(
                    "Download",
                    project_to_json(export_project),
                    file_name=f"{export_project.title}.json",
                    use_container_width=True,
                    mime="application/json",
                    help="Download raw project data.",
                )
            else:
                try:
                    docx_bytes = project_to_docx_bytes(export_project)
                except Exception:
                    st.error("DOCX export requires python-docx support.")
                else:
                    st.download_button(
                        "Download",
                        docx_bytes,
                        file_name=f"{export_project.title}.docx",
                        use_container_width=True,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        help="Download as DOCX.",
                    )

    def render_no_project_state(context: str, key_prefix: str) -> None:
        with st.container(border=True):
            st.markdown("### Start with a project")
            st.caption(f"No active project loaded for {context}. Open Projects to create or load one.")
            c1, c2 = st.columns(2)
            with c1:
                if st.button("Go to Projects", use_container_width=True, key=f"{key_prefix}_go_projects"):
                    st.session_state.page = "projects"
                    st.rerun()
            with c2:
                if st.button("Go to Dashboard", use_container_width=True, key=f"{key_prefix}_go_home"):
                    st.session_state.page = "home"
                    st.rerun()

    def render_outline():
        p = st.session_state.project
        if not p:
            render_no_project_state("Outline", "empty_outline")
            return
        provider, active_key, _ = get_active_key_status()
        render_page_header(
            "Outline",
            "Your blueprint. Generate structure, scan entities, and keep the story plan here.",
            tag="Project",
            key_prefix="outline_header",
        )
        if not active_key:
            st.info(
                f"Add a {_provider_label(provider)} API key in AI Settings to scan entities or generate outlines."
            )

        # Keep the outline editor widget in sync when switching projects/pages
        if st.session_state.get("out_txt_project_id") != p.id:
            st.session_state["out_txt_project_id"] = p.id
            st.session_state["_outline_sync"] = p.outline or ""  # apply on next rerun before widget renders

        # If AI updated the outline this run, apply it BEFORE the text_area is created.
        if st.session_state.get("_outline_sync") is not None:
            st.session_state["out_txt"] = st.session_state.pop("_outline_sync")

        # Keep outline metadata inputs synchronized with persisted project metadata.
        project_meta = (p.title or "", p.genre or "")
        if st.session_state.get("outline_meta_project_id") != p.id:
            st.session_state["outline_meta_project_id"] = p.id
            st.session_state["outline_project_title_input"] = project_meta[0]
            st.session_state["outline_project_genre_input"] = project_meta[1]
            st.session_state["outline_meta_snapshot"] = project_meta
        elif st.session_state.get("outline_meta_snapshot") != project_meta:
            st.session_state["outline_project_title_input"] = project_meta[0]
            st.session_state["outline_project_genre_input"] = project_meta[1]
            st.session_state["outline_meta_snapshot"] = project_meta

        with st.container(border=True):
            top1, top2, top3 = st.columns([2.2, 1.1, 1])
            with top1:
                new_title = st.text_input(
                    "Project Title",
                    key="outline_project_title_input",
                )
                if new_title != p.title:
                    p.title = new_title
                    save_p()
                    st.session_state["outline_meta_snapshot"] = (p.title or "", p.genre or "")
            with top2:
                new_genre = st.text_input(
                    "Genre",
                    key="outline_project_genre_input",
                    placeholder="e.g., Dark Fantasy",
                )
                if new_genre != p.genre:
                    p.genre = new_genre
                    save_p()
                    st.session_state["outline_meta_snapshot"] = (p.title or "", p.genre or "")
            with top3:
                if st.button("Save Project", key="outline_save_project", type="primary", use_container_width=True):
                    if persist_project(p, action="save"):
                        st.toast("Saved")
                        st.rerun()

        left, right = st.columns([2.1, 1])

        # Placeholder for streaming AI-generated outline; lives below the editor, like chapters.
        outline_stream_ph = st.empty()

        with left:
            with st.container(border=True):
                st.markdown("### Blueprint")
                val = st.text_area(
                    "Plot Outline",
                    height=560,
                    key="out_txt",
                    label_visibility="collapsed",
                    help="Write your story's plot outline here. Changes are auto-saved."
                )
                if val != p.outline:
                    p.outline = val
                    save_p()

                if st.button(
                    "Save Outline",
                    use_container_width=True,
                    help="Save and automatically scan for characters, locations, and other entities",
                    key="outline_save_outline_btn",
                ):
                    if persist_project(p, action="save"):
                        # Automatically scan entities on save so World Bible stays in sync.
                        scan_result = extract_entities_ui(
                            p.outline or "",
                            "Outline",
                            show_feedback=False,
                        )
                        if scan_result:
                            if scan_result.get("status") == "updated":
                                st.success(
                                    "Outline saved and scanned: "
                                    f"{scan_result.get('added', 0)} new, "
                                    f"{scan_result.get('matched', 0)} existing, "
                                    f"{scan_result.get('flagged', 0)} flagged."
                                )
                            elif scan_result.get("status") == "matched_only":
                                st.info("Outline saved. Entities were detected and matched existing entries.")
                            else:
                                st.info("Outline saved. No entities were detected in this outline text.")
                        else:
                            st.success("Outline saved.")

                with st.expander("Revision Tools"):
                    st.caption("Refine outline structure/tone with AI, then review before applying.")
                    if not active_key:
                        st.info(
                            f"Add a {_provider_label(provider)} API key in AI Settings to generate outline revisions."
                        )
                    outline_style = st.selectbox(
                        "Revision style",
                        list(REWRITE_PRESETS.keys()),
                        key=f"outline_revision_style_{p.id}",
                    )
                    outline_custom = ""
                    if outline_style == "Custom":
                        outline_custom = st.text_input(
                            "Custom instructions",
                            key=f"outline_revision_custom_{p.id}",
                        )
                    rev_cooldown = _cooldown_remaining(f"outline_revision_{p.id}", 12)
                    rev_label = (
                        f"Generate Revision ({rev_cooldown}s)"
                        if rev_cooldown
                        else "Generate Revision"
                    )
                    if st.button(
                        rev_label,
                        use_container_width=True,
                        key=f"outline_revision_generate_{p.id}",
                        disabled=bool(rev_cooldown) or not active_key,
                    ):
                        _mark_action(f"outline_revision_{p.id}")
                        prompt = rewrite_prompt(val or p.outline or "", outline_style, outline_custom)
                        full = ""
                        for chunk in AIEngine().generate_stream(prompt, get_ai_model()):
                            full += chunk
                            outline_stream_ph.markdown(f"**REVISING OUTLINE:**\n\n{full}")
                        if full.strip():
                            st.session_state["pending_outline_revision_text"] = full.strip()
                            st.session_state["pending_outline_revision_meta"] = {
                                "style": outline_style,
                                "custom": outline_custom,
                                "timestamp": time.time(),
                                "project_id": p.id,
                            }
                            st.rerun()

                pending_outline = st.session_state.get("pending_outline_revision_text") or ""
                pending_meta = st.session_state.get("pending_outline_revision_meta") or {}
                if pending_outline and pending_meta.get("project_id") == p.id:
                    with st.container(border=True):
                        st.markdown("#### ? Review Outline Revision")
                        st.code(pending_outline, language="markdown")
                        review_cols = st.columns(3)
                        with review_cols[0]:
                            if st.button(
                                "Apply Replace",
                                type="primary",
                                use_container_width=True,
                                key=f"outline_revision_apply_replace_{p.id}",
                            ):
                                st.session_state["outline_prev_text"] = p.outline or ""
                                p.outline = pending_outline
                                st.session_state["_outline_sync"] = p.outline
                                save_p()
                                st.session_state["pending_outline_revision_text"] = ""
                                st.session_state["pending_outline_revision_meta"] = {}
                                st.toast("Outline revision applied.")
                                st.rerun()
                        with review_cols[1]:
                            if st.button(
                                "Apply Append",
                                use_container_width=True,
                                key=f"outline_revision_apply_append_{p.id}",
                            ):
                                st.session_state["outline_prev_text"] = p.outline or ""
                                p.outline = (((p.outline or "").rstrip() + "\n\n" + pending_outline).strip())
                                st.session_state["_outline_sync"] = p.outline
                                save_p()
                                st.session_state["pending_outline_revision_text"] = ""
                                st.session_state["pending_outline_revision_meta"] = {}
                                st.toast("Outline revision appended.")
                                st.rerun()
                        with review_cols[2]:
                            if st.button(
                                "Discard",
                                use_container_width=True,
                                key=f"outline_revision_discard_{p.id}",
                            ):
                                st.session_state["pending_outline_revision_text"] = ""
                                st.session_state["pending_outline_revision_meta"] = {}
                                st.toast("Revision discarded.")
                                st.rerun()

                if st.button(
                    "Undo last outline apply",
                    use_container_width=True,
                    key=f"outline_revision_undo_{p.id}",
                    help="Restore the previous outline text from the last apply action.",
                ):
                    prev_outline = st.session_state.get("outline_prev_text")
                    if prev_outline is None:
                        st.info("No previous outline text available.")
                    else:
                        p.outline = prev_outline
                        st.session_state["_outline_sync"] = p.outline
                        save_p()
                        st.toast("Previous outline restored.")
                        st.rerun()

        with right:
            with st.container(border=True):
                st.markdown("### Architect (AI)")
                st.caption("Generate a chapter-by-chapter outline and append it to your blueprint.")

                init_state("outline_chapter_target", 12)
                chaps = st.number_input(
                    "Number of Chapters",
                    min_value=1,
                    max_value=50,
                    step=1,
                    key="outline_chapter_target",
                    help="Specify how many chapters to generate in the AI outline",
                )
                chaps = int(chaps)
                provider, active_key, _ = get_active_key_status()
                outline_cooldown = _cooldown_remaining("outline_generate", 12)
                outline_label = (
                    f"Generate Structure ({outline_cooldown}s)" if outline_cooldown else "Generate Structure"
                )
                
                # Determine help text based on disabled state
                button_help = get_ai_button_help(
                    outline_cooldown,
                    active_key,
                    "generate a detailed chapter-by-chapter outline for your story"
                )
                
                if not active_key:
                    show_api_key_notice("generate outlines")
                if st.button(
                    outline_label,
                    type="primary",
                    use_container_width=True,
                    disabled=bool(outline_cooldown),
                    key="outline_generate_structure_btn",
                    help=button_help,
                ):
                    if not active_key:
                        st.toast(
                            f"Add a {_provider_label(provider)} API key in AI Settings to generate outlines."
                        )
                        st.session_state.page = "ai"
                        st.rerun()
                    if not get_ai_model():
                        st.toast("Choose an AI model in AI Settings first.")
                        st.session_state.page = "ai"
                        st.rerun()
                    _mark_action("outline_generate")
                    # use outline_stream_ph defined above
                    full = ""
                    prompt = (
                        f"Write a detailed {chaps}-chapter outline for a {p.genre} novel: {p.title}. "
                        "Use structure: Chapter X: [Title] - [Summary]."
                    )
                    try:
                        # Show loading indicator before streaming starts
                        with outline_stream_ph.container():
                            st.markdown("**Generating chapter outline...**")

                        for chunk in AIEngine().generate_stream(prompt, get_ai_model()):
                            full += chunk
                            outline_stream_ph.markdown(full)

                        if full.strip():
                            new_outline = (((p.outline or "").rstrip() + "\n\n" + full.strip()).strip())
                            p.outline = new_outline
                            st.session_state["_outline_sync"] = new_outline  # apply on next rerun before widget renders
                            save_p()
                            st.rerun()
                        else:
                            st.warning("No outline text returned. Try again or switch model/provider in AI Settings.")
                    except Exception as exc:
                        logger.warning("Generate Structure failed", exc_info=True)
                        st.error(f"Generate Structure failed: {exc}")

    def build_expander_label(base: str, suffix: Optional[str] = None) -> str:
        if suffix:
            return f"{base} | {suffix}"
        return base

    def render_world():
        from app.views.world_relationships import render_world_relationships

        p = st.session_state.project
        if not p:
            render_no_project_state("World Bible", "empty_world")
            return
        def open_add_entity() -> None:
            tab_label = st.session_state.get("world_tabs") or "Characters"
            category_map = {
                "Characters": "Character",
                "Locations": "Location",
                "Factions": "Faction",
                "Lore": "Lore",
            }
            category = category_map.get(tab_label, "Character")
            st.session_state[f"add_open_{category}"] = True
            st.rerun()

        render_page_header(
            "World Bible",
            "Your story's encyclopedia: Track characters, locations, factions, and lore. Mantis uses this to keep your writing consistent.",
            tag="Canon",
            key_prefix="world_header",
        )
        if not get_active_key_status()[1]:
            show_api_key_notice("run scans")

        st.html(
            """
            <style>
            .world-overview-card {
                padding: 14px 18px;
                border-radius: 18px;
                background: var(--mantis-surface);
                border: 1px solid var(--mantis-card-border);
            }
            .world-pill {
                display: inline-flex;
                align-items: center;
                gap: 6px;
                padding: 6px 12px;
                border-radius: 999px;
                font-size: 0.85rem;
                font-weight: 600;
                border: 1px solid var(--mantis-primary-border);
                background: var(--mantis-accent-soft);
            }
            .world-pill.good {
                border-color: rgba(34,197,94,0.45);
                background: rgba(34,197,94,0.16);
            }
            .world-pill.warn {
                border-color: rgba(245,158,11,0.4);
                background: rgba(245,158,11,0.16);
            }
            .world-pill.risk {
                border-color: rgba(239,68,68,0.45);
                background: rgba(239,68,68,0.16);
            }
            .world-card {
                padding: 14px 18px;
                border-radius: 18px;
                background: var(--mantis-surface);
                border: 1px solid var(--mantis-card-border);
                box-shadow: 0 10px 22px rgba(15, 23, 42, 0.06);
            }
            .world-card.highlight {
                border-color: rgba(245,158,11,0.6);
                box-shadow: 0 12px 24px rgba(245,158,11,0.18);
            }
            .world-card-header {
                display: flex;
                align-items: center;
                justify-content: space-between;
                gap: 14px;
            }
            .world-card-title {
                display: flex;
                align-items: center;
                gap: 10px;
                font-size: 1.05rem;
                font-weight: 600;
            }
            .world-card-meta {
                display: flex;
                align-items: center;
                gap: 10px;
                font-size: 0.85rem;
                color: var(--mantis-muted);
            }
            .world-badge {
                padding: 4px 10px;
                border-radius: 999px;
                background: var(--mantis-accent-soft);
                border: 1px solid var(--mantis-primary-border);
                color: var(--mantis-text);
                font-weight: 600;
                text-transform: uppercase;
                letter-spacing: 0.06em;
                font-size: 0.65rem;
            }
            .world-card-actions {
                display: flex;
                align-items: center;
                gap: 10px;
            }
            .world-card-metric {
                font-weight: 600;
                color: var(--mantis-text);
            }
            </style>
            """,
        )

        entries = list(p.world_db.values())
        chapters = p.get_ordered_chapters()
        chapter_texts = [(c, (c.content or "")) for c in chapters]
        chapter_texts_lower = [(c, text.lower()) for c, text in chapter_texts]

        def _count_mentions(aliases: List[str], text: str) -> int:
            total = 0
            for alias in aliases:
                alias_clean = (alias or "").strip()
                if not alias_clean:
                    continue
                pattern = rf"\\b{re.escape(alias_clean.lower())}\\b"
                total += len(re.findall(pattern, text))
            return total

        mention_counts: Dict[str, int] = {}
        mention_refs: Dict[str, List[Chapter]] = {}
        for ent in entries:
            aliases = [ent.name] + (ent.aliases or [])
            total_hits = 0
            hit_chapters = []
            for chap, lower_text in chapter_texts_lower:
                hits = _count_mentions(aliases, lower_text)
                if hits:
                    hit_chapters.append(chap)
                total_hits += hits
            mention_counts[ent.id] = total_hits
            mention_refs[ent.id] = hit_chapters

        orphaned_ids = {eid for eid, count in mention_counts.items() if count == 0}
        under_described_ids = {
            e.id for e in entries if len((e.description or "").strip()) < 30
        }
        normalized_name_map: Dict[str, List[str]] = {}
        for ent in entries:
            normalized = Project._normalize_entity_name(ent.name)
            if not normalized:
                continue
            normalized_name_map.setdefault(normalized, []).append(ent.id)
        collision_ids = {
            ent_id
            for ent_ids in normalized_name_map.values()
            if len(ent_ids) > 1
            for ent_id in ent_ids
        }
        flagged_entity_ids = orphaned_ids | under_described_ids | collision_ids

        locked_entities = [
            e for e in entries if "locked" in (e.tags or "").lower()
        ]
        last_scan_ts = st.session_state.get("last_entity_scan") or p.last_modified
        canon_icon, canon_label = get_canon_health()
        canon_class = "good"
        if canon_icon == "RISK":
            canon_class = "warn"
        elif canon_icon == "RISK":
            canon_class = "risk"

        with card("World overview", "Live status of your canon database."):
            top_cols = st.columns([1, 1, 1, 1, 1.2])
            with top_cols[0]:
                stat_tile("Total entities", str(len(entries)), icon="\U0001F4DA")
            with top_cols[1]:
                stat_tile("Orphaned", str(len(orphaned_ids)), icon="\U0001F517")
            with top_cols[2]:
                stat_tile("Locked", str(len(locked_entities)), icon="\U0001F512")
            with top_cols[3]:
                stat_tile(
                    "Last scan",
                    time.strftime("%Y-%m-%d %H:%M", time.localtime(last_scan_ts)),
                    icon="\U0001F552",
                )
            with top_cols[4]:
                stat_tile("Canon health", f"{canon_icon} {canon_label}", icon="\u2705")

        with card("Search & filters", "Refine by status, recency, or canon risk."):
            f1, f2, f3, f4, f5 = st.columns([2.2, 1, 1, 1, 1.1])
            with f1:
                if st.session_state.get("world_search_pending"):
                    st.session_state["world_search"] = st.session_state.pop("world_search_pending")
                query = st.text_input(
                    "Search",
                    placeholder="Type a name or alias to filter...",
                    key="world_search",
                )
            with f2:
                show_orphaned = st.checkbox("Orphaned only", key="world_filter_orphaned")
            with f3:
                show_canon_risk = st.checkbox("Canon risk only", key="world_filter_risk")
            with f4:
                show_recent = st.checkbox("Recently added", key="world_filter_recent")
            with f5:
                sort_mode = st.selectbox(
                    "Sort by",
                    ["Name (A-Z)", "Most mentions", "Newest first"],
                    key="world_sort_mode",
                )

        render_world_relationships(
            project=p,
            entries=entries,
            chapters=chapters,
            mention_refs=mention_refs,
            mention_counts=mention_counts,
        )

        tab_options = ["Characters", "Locations", "Factions", "Lore"]
        default_tab = st.session_state.get("world_tabs", tab_options[0])
        if default_tab not in tab_options:
            default_tab = tab_options[0]
        selected_tab = st.radio(
            "World Bible sections",
            tab_options,
            index=tab_options.index(default_tab),
            horizontal=True,
            key="world_tabs",
        )

        review_queue = st.session_state.get("world_bible_review", [])
        if review_queue:
            with card("Review AI Suggestions", "AI suggestions are queued for review. Apply to update canon."):
                for idx, item in enumerate(list(review_queue)):
                    stype = item.get("type", "new")
                    label = f"{item.get('name', 'Unnamed')} | {item.get('category', 'Lore')}"
                    if stype == "update":
                        label = f"{label}"
                    elif stype == "alias_only":
                        label = f"{label}"
                    else:
                        label = f"{label}"
                    expander_label = build_expander_label(label, str(idx))
                    with st.expander(expander_label):
                        type_labels = {"update": "Update Existing", "alias_only": "Add Aliases", "new": "New Entry"}
                        st.markdown(f"**Action:** {type_labels.get(stype, stype.title())}")
                        if item.get("match_name"):
                            st.markdown(f"**Matches:** {item['match_name']}")
                        confidence = item.get("confidence")
                        if confidence is not None:
                            st.markdown(f"**Confidence:** {confidence:.2f}")
                        source = item.get("source")
                        if source:
                            st.markdown(f"**Source:** {source}")
                        if item.get("reason"):
                            st.info(item["reason"])
                        novel_bullets = item.get("novel_bullets") or []
                        if novel_bullets:
                            st.markdown("**New details to add:**")
                            for b in novel_bullets:
                                st.markdown(f"- {b}")
                        elif item.get("description"):
                            st.markdown("**Suggested Notes**")
                            st.write(item.get("description"))
                        novel_aliases = item.get("novel_aliases") or []
                        if novel_aliases:
                            st.markdown(f"**New aliases:** {', '.join(novel_aliases)}")
                        elif item.get("aliases"):
                            st.markdown(f"**Aliases:** {', '.join(item.get('aliases') or [])}")

                        c1, c2 = st.columns(2)
                        with c1:
                            if st.button("Apply", key=f"apply_suggestion_{idx}", use_container_width=True):
                                from app.services.world_bible_merge import apply_suggestion as _apply_suggestion
                                applied_ent, _action = _apply_suggestion(p, item)
                                # Clear cached widget values so the UI
                                # reflects the updated description/aliases
                                # on the next rerun.
                                if applied_ent is not None:
                                    st.session_state.pop(f"desc_{applied_ent.id}", None)
                                    st.session_state.pop(f"aliases_{applied_ent.id}", None)
                                review_queue.pop(idx)
                                st.session_state["world_bible_review"] = review_queue
                                persist_project(p)
                                if _action == "duplicate":
                                    st.toast("No new World Bible changes to apply.")
                                else:
                                    st.toast("World Bible updated.")
                                st.rerun()
                        with c2:
                            if st.button("Ignore", key=f"ignore_suggestion_{idx}", use_container_width=True):
                                review_queue.pop(idx)
                                st.session_state["world_bible_review"] = review_queue
                                st.toast("Suggestion removed.")
                                st.rerun()

        focus_entity = st.session_state.get("world_focus_entity")
        recent_cutoff = time.time() - (7 * 86400)

        def render_cat(category: str):
            with card():
                top = st.columns([1, 1.2, 1])
                with top[0]:
                    st.markdown(f"### {category}")
                with top[2]:
                    if st.button(
                        f"Add {category}",
                        use_container_width=True,
                        key=f"world_open_add_{category}",
                    ):
                        st.session_state[f"add_open_{category}"] = True
                        st.rerun()

                if st.session_state.get(f"add_open_{category}", False):
                    with st.form(f"add_{category}"):
                        cat_slug = re.sub(r"[^a-z0-9]+", "_", category.lower()).strip("_")
                        n = st.text_input("Name", key=f"add_{cat_slug}_name")
                        d = st.text_area("Description", key=f"add_{cat_slug}_desc")
                        a = st.text_input("Aliases (comma-separated)", key=f"add_{cat_slug}_aliases")
                        s1, s2 = st.columns(2)
                        with s1:
                            ok = st.form_submit_button("Save", type="primary", use_container_width=True)
                        with s2:
                            cancel = st.form_submit_button("Cancel", use_container_width=True)
                        if ok:
                            aliases = [alias.strip() for alias in (a or "").split(",") if alias.strip()]
                            p.upsert_entity(n, category, d, aliases=aliases, allow_merge=True, allow_alias=True)
                            persist_project(p)
                            st.session_state[f"add_open_{category}"] = False
                            st.rerun()
                        if cancel:
                            st.session_state[f"add_open_{category}"] = False
                            st.rerun()

                ents = [
                    e
                    for e in p.world_db.values()
                    if Project._normalize_category(e.category)
                    == Project._normalize_category(category)
                ]
                if query:
                    q = query.lower()
                    ents = [
                        e
                        for e in ents
                        if q in (e.name or "").lower()
                        or any(q in alias.lower() for alias in (e.aliases or []))
                    ]
                if show_orphaned:
                    ents = [e for e in ents if e.id in orphaned_ids]
                if show_canon_risk:
                    ents = [e for e in ents if e.id in flagged_entity_ids]
                if show_recent:
                    ents = [e for e in ents if e.created_at >= recent_cutoff]

                if not ents:
                    st.info(f"No {category} entries yet. Add one above or scan entities from your outline/chapters.")
                    return

                if sort_mode == "Most mentions":
                    ents = sorted(
                        ents,
                        key=lambda ent: (mention_counts.get(ent.id, 0), (ent.name or "").lower()),
                        reverse=True,
                    )
                elif sort_mode == "Newest first":
                    ents = sorted(ents, key=lambda ent: ent.created_at, reverse=True)
                else:
                    ents = sorted(ents, key=lambda ent: (ent.name or "").lower())

                st.caption(f"Showing {len(ents)} {category.lower()} entr{'y' if len(ents) == 1 else 'ies'}")

                for idx, e in enumerate(ents):
                    mention_count = mention_counts.get(e.id, 0)
                    is_orphaned = e.id in orphaned_ids
                    is_under_described = e.id in under_described_ids
                    is_collision = e.id in collision_ids
                    if is_orphaned:
                        status_icon = "OK"
                    elif is_collision:
                        status_icon = "OK"
                    elif is_under_described:
                        status_icon = "OK"
                    else:
                        status_icon = "OK"

                    highlight = "highlight" if e.id in flagged_entity_ids or e.id == focus_entity else ""

                    with st.container(border=True):
                        st.html(
                            f"""
                            <div class="world-card {highlight}">
                                <div class="world-card-header">
                                    <div class="world-card-title">{status_icon} {e.name}</div>
                                    <div class="world-card-meta">
                                        <span class="world-badge">{e.category}</span>
                                        <span class="world-card-metric">{mention_count} mentions</span>
                                    </div>
                                </div>
                            </div>
                            """,
                        )

                        issues = []
                        if is_orphaned:
                            issues.append("Orphaned")
                        if is_under_described:
                            issues.append("Needs detail")
                        if is_collision:
                            issues.append("Name collision")
                        if issues:
                            st.caption(f"{'  '.join(issues)}")

                        detail_suffix = e.name
                        detail_label = build_expander_label("Open details", detail_suffix)
                        with st.expander(detail_label, expanded=e.id == focus_entity):
                            new_desc = st.text_area("Notes", e.description, key=f"desc_{e.id}", height=140)
                            if new_desc != e.description:
                                e.description = new_desc
                                persist_project(p)

                            alias_text = st.text_input(
                                "Aliases (comma-separated)",
                                value=", ".join(e.aliases or []),
                                key=f"aliases_{e.id}",
                            )
                            if alias_text != ", ".join(e.aliases or []):
                                e.aliases = [a.strip() for a in alias_text.split(",") if a.strip()]
                                persist_project(p)

                            refs = mention_refs.get(e.id, [])
                            if refs:
                                options = {
                                    f"Chapter {chap.index}: {chap.title}": chap.id for chap in refs
                                }
                                sel = st.selectbox(
                                    "Jump to chapter",
                                    list(options.keys()),
                                    key=f"jump_select_{e.id}",
                                )
                                if st.button("Jump to Chapter", key=f"jump_{e.id}", use_container_width=True):
                                    st.session_state.page = "chapters"
                                    st.session_state.curr_chap_id = options[sel]
                                    st.session_state.active_chapter_id = options[sel]
                                    st.session_state._force_nav = True
                                    st.rerun()
                            else:
                                st.caption("No chapter references yet.")

                            d1, d2 = st.columns([1, 1])
                            with d1:
                                if st.session_state.delete_entity_id == e.id:
                                    st.warning(f"Delete **{st.session_state.delete_entity_name or e.name}**?")
                                    cdel1, cdel2 = st.columns(2)
                                    with cdel1:
                                        if st.button(
                                            "Confirm",
                                            type="primary",
                                            use_container_width=True,
                                            key=f"world_delete_confirm_{e.id}",
                                        ):
                                            p.delete_entity(e.id)
                                            persist_project(p)
                                            st.session_state.delete_entity_id = None
                                            st.session_state.delete_entity_name = None
                                            st.toast("Entity deleted.")
                                            st.rerun()
                                    with cdel2:
                                        if st.button(
                                            "Cancel",
                                            use_container_width=True,
                                            key=f"world_delete_cancel_{e.id}",
                                        ):
                                            st.session_state.delete_entity_id = None
                                            st.session_state.delete_entity_name = None
                                            st.rerun()
                                elif st.button("Delete", key=f"del_{e.id}", use_container_width=True):
                                    st.session_state.delete_entity_id = e.id
                                    st.session_state.delete_entity_name = e.name
                                    st.rerun()
                            with d2:
                                alias_total = len(e.aliases or [])
                                st.caption(
                                    f"Aliases: {alias_total} | Created: {time.strftime('%Y-%m-%d', time.localtime(e.created_at))}"
                                )

        if selected_tab == "Characters":
            render_cat("Character")
        elif selected_tab == "Locations":
            render_cat("Location")
        elif selected_tab == "Factions":
            render_cat("Faction")
        elif selected_tab == "Lore":
            render_cat("Lore")
        elif selected_tab == "Memory":
            st.markdown("### World Memory")
            st.caption("Keep canon notes, timelines, and facts the AI should always know.")
            with st.container(border=True):
                st.markdown("#### Auto-add settings")
                st.caption("Auto-add confidence controls are managed in Workspace Settings to avoid duplicate configuration.")
            st.markdown("#### Hard Canon Rules")
            hard_key = f"world_memory_hard_{p.id}"
            hard_default = p.memory_hard or p.memory
            hard_val = st.text_area("Hard Canon Rules", hard_default, height=160, key=hard_key)
            if hard_val != p.memory_hard:
                p.memory_hard = hard_val
                save_p()

            st.markdown("#### Soft Guidelines")
            soft_key = f"world_memory_soft_{p.id}"
            soft_val = st.text_area("Soft Guidelines", p.memory_soft, height=160, key=soft_key)
            if soft_val != p.memory_soft:
                p.memory_soft = soft_val
                save_p()

            st.markdown("#### Project Memory")
            memory_key = f"world_memory_{p.id}"
            memory_val = st.text_area("Memory", p.memory, height=320, key=memory_key)
            if memory_val != p.memory:
                p.memory = memory_val
                save_p()
            if st.button("Save Memory", use_container_width=True):
                if persist_project(p, action="save"):
                    st.toast("Memory saved")
                    st.rerun()

            st.divider()
            st.markdown("#### Coherence Check")
            scope_cols = st.columns(3)
            with scope_cols[0]:
                scope_outline = st.checkbox("Outline", value=True, key=f"coh_outline_{p.id}")
            with scope_cols[1]:
                scope_world = st.checkbox("World Bible", value=True, key=f"coh_world_{p.id}")
            with scope_cols[2]:
                scope_chapters = st.checkbox("Chapters", value=True, key=f"coh_chapters_{p.id}")

            provider, active_key, _ = get_active_key_status()
            coherence_cooldown = _cooldown_remaining(f"coherence_{p.id}", 15)
            coherence_label = (
                f"Run Coherence Check ({coherence_cooldown}s)"
                if coherence_cooldown
                else "Run Coherence Check"
            )
            if not active_key:
                st.info(
                    f"Add a {_provider_label(provider)} API key in AI Settings to run coherence checks."
                )
            if st.button(
                coherence_label,
                use_container_width=True,
                disabled=bool(coherence_cooldown) or not active_key,
                key=f"coh_run_{p.id}",
            ):
                if not (scope_outline or scope_world or scope_chapters):
                    st.warning("Choose at least one scope (Outline, World Bible, or Chapters).")
                    st.stop()
                model_name = get_ai_model()
                if not model_name:
                    st.warning("Choose an AI model in AI Settings first.")
                    st.session_state.page = "ai"
                    st.rerun()
                _mark_action(f"coherence_{p.id}")
                compiled_world_bible = "\n".join(
                    f"{e.name} ({e.category}): {e.description}"
                    for e in p.world_db.values()
                )
                chapter_payload = [
                    {
                        "chapter_index": c.index,
                        "summary": c.summary,
                        "excerpt": (c.content or "")[:800],
                    }
                    for c in p.get_ordered_chapters()
                ]
                outline_payload = p.outline if scope_outline else ""
                world_payload = compiled_world_bible if scope_world else ""
                chapters_payload = chapter_payload if scope_chapters else []
                with st.spinner("Running coherence check..."):
                    results = AnalysisEngine.coherence_check(
                        memory=p.memory,
                        author_note=p.author_note,
                        style_guide=p.style_guide,
                        outline=outline_payload,
                        world_bible=world_payload,
                        chapters=chapters_payload,
                        model=model_name,
                    )
                st.session_state["coherence_results"] = results or []
                canon_icon, canon_label = get_canon_health()
                st.session_state.setdefault("canon_health_log", [])
                st.session_state["canon_health_log"].append(
                    {
                        "timestamp": time.time(),
                        "status": canon_icon,
                        "issue_count": len(st.session_state.get("coherence_results", [])),
                    }
                )
                st.session_state["canon_health_log"] = st.session_state["canon_health_log"][-30:]
                if results:
                    st.toast("Coherence issues found.")
                else:
                    st.toast("No coherence issues detected.")
                st.rerun()

            results = st.session_state.get("coherence_results", [])
            if results:
                st.markdown("#### Coherence Issues")
                for idx, issue in enumerate(list(results)):
                    with st.container(border=True):
                        chapter_idx = issue.get("chapter_index", "?")
                        st.markdown(f"**Chapter #{chapter_idx}**")
                        st.markdown(f"**Issue:** {issue.get('issue', 'Unspecified issue')}")
                        st.markdown(f"**Confidence:** {issue.get('confidence', 'unknown')}")
                        st.markdown("**Suggested Rewrite:**")
                        st.write(issue.get("suggested_rewrite", ""))

                        a1, a2 = st.columns(2)
                        with a1:
                            if st.button("Apply Fix", key=f"coh_apply_{idx}", use_container_width=True):
                                target_excerpt = issue.get("target_excerpt", "")
                                try:
                                    chapter_num = int(chapter_idx)
                                except (TypeError, ValueError):
                                    chapter_num = None
                                target_chapter = None
                                if chapter_num is not None:
                                    for c in p.get_ordered_chapters():
                                        if c.index == chapter_num:
                                            target_chapter = c
                                            break
                                if not target_chapter:
                                    st.warning("Chapter not found for this issue.")
                                elif target_excerpt and target_excerpt in (target_chapter.content or ""):
                                    updated = (target_chapter.content or "").replace(
                                        target_excerpt,
                                        issue.get("suggested_rewrite", ""),
                                        1,
                                    )
                                    target_chapter.update_content(updated, "Coherence Fix")
                                    persist_project(p)
                                    results.pop(idx)
                                    st.session_state["coherence_results"] = results
                                    update_locked_chapters()
                                    st.toast("Applied fix.")
                                    st.rerun()
                                elif issue.get("suggested_rewrite"):
                                    insertion = issue.get("suggested_rewrite", "").strip()
                                    if insertion:
                                        spacer = "\n\n" if (target_chapter.content or "").strip() else ""
                                        updated = f"{(target_chapter.content or '').rstrip()}{spacer}{insertion}"
                                        target_chapter.update_content(updated, "Coherence Fix (Appended)")
                                        persist_project(p)
                                        results.pop(idx)
                                        st.session_state["coherence_results"] = results
                                        update_locked_chapters()
                                        st.toast("Applied fix (appended).")
                                        st.rerun()
                                else:
                                    st.warning("Target excerpt not found in chapter content.")
                        with a2:
                            if st.button("Ignore", key=f"coh_ignore_{idx}", use_container_width=True):
                                results.pop(idx)
                                st.session_state["coherence_results"] = results
                                update_locked_chapters()
                                st.toast("Issue ignored.")
                                st.rerun()
        elif selected_tab == "Insights":
            st.markdown("### World Bible Insights")
            st.caption("Quick stats on your current canon database.")
            entries = list(p.world_db.values())
            total_entries = len(entries)
            counts = {"Character": 0, "Location": 0, "Faction": 0, "Lore": 0}
            for ent in entries:
                category = Project._normalize_category(ent.category)
                counts[category] = counts.get(category, 0) + 1

            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Total Entries", total_entries)
            c2.metric("Characters", counts.get("Character", 0))
            c3.metric("Locations", counts.get("Location", 0))
            c4.metric("Factions", counts.get("Faction", 0))

            c5, c6 = st.columns(2)
            c5.metric("Lore", counts.get("Lore", 0))
            c6.metric(
                "Last Updated",
                time.strftime("%Y-%m-%d", time.localtime(p.last_modified)),
            )

            st.divider()
            if not st.session_state.get("is_premium", False):
                st.info("Canon history is a Premium feature.")
            else:
                st.markdown("### Canon Health History")
                for entry in st.session_state.get("canon_health_log", []):
                    st.caption(
                        f"{time.strftime('%Y-%m-%d %H:%M', time.localtime(entry['timestamp']))} "
                        f"{entry['status']} ({entry['issue_count']} issues)"
                    )

            st.divider()
            st.markdown("### ? Timeline Heatmap")
            for chap in p.get_ordered_chapters():
                intensity = min(1.0, chap.word_count / 2000)
                st.progress(intensity, text=f"Chapter {chap.index}: {chap.word_count} words")

            st.divider()
            st.markdown("#### Entity Utilization")
            utilization_rows = []
            for ent in entries:
                aliases = [ent.name] + (ent.aliases or [])
                total_hits = mention_counts.get(ent.id, 0)
                utilization_rows.append(
                    {
                        "Name": ent.name,
                        "Category": ent.category,
                        "Appearances": total_hits,
                        "Orphaned": ent.id in orphaned_ids,
                        "Under-described": ent.id in under_described_ids,
                    }
                )
            if utilization_rows:
                st.dataframe(utilization_rows, use_container_width=True, hide_index=True)
            else:
                st.info("No entities yet to analyze.")

            st.divider()
            st.markdown("#### Flagged Entities")
            flagged_entities = [ent for ent in entries if ent.id in flagged_entity_ids]
            if flagged_entities:
                for ent in flagged_entities:
                    reasons = []
                    if ent.id in orphaned_ids:
                        reasons.append("Orphaned")
                    if ent.id in under_described_ids:
                        reasons.append("Needs detail")
                    if ent.id in collision_ids:
                        reasons.append("Name collision")
                    with st.container(border=True):
                        r1, r2 = st.columns([3, 1])
                        with r1:
                            st.markdown(f"**{ent.name}**  {ent.category}")
                            st.caption("  ".join(reasons))
                        with r2:
                            if st.button("Jump to Entity", key=f"jump_entity_{ent.id}", use_container_width=True):
                                st.session_state["world_focus_entity"] = ent.id
                                st.session_state["world_search_pending"] = ent.name
                                st.toast("Entity highlighted in World Bible.")
                                st.rerun()
            else:
                st.success("No flagged entities right now.")

            st.divider()
            st.markdown("#### Canon Risk Flags")
            coherence_results = st.session_state.get("coherence_results", [])
            high_canon_drift = len(coherence_results) > 0
            alias_collision = any(len(names) > 1 for names in normalized_name_map.values())

            timeline_dense = False
            dense_chapters = []
            for chap in chapters:
                summary_text = (chap.summary or "").strip()
                base_text = summary_text if summary_text else (chap.content or "")[:800]
                sentence_count = len([s for s in re.split(r"[.!?]+", base_text) if s.strip()])
                if sentence_count > 3:
                    timeline_dense = True
                    dense_chapters.append(chap.index)

            if high_canon_drift:
                st.warning("High Canon Drift: coherence issues detected.")
            if alias_collision:
                st.warning("Alias Collision: multiple entities share normalized names.")
            if timeline_dense:
                chap_list = ", ".join(str(idx) for idx in dense_chapters)
                st.warning(f"Timeline Density: >3 major events in chapters {chap_list}.")
            if not any([high_canon_drift, alias_collision, timeline_dense]):
                st.success("No canon risk flags detected.")

            st.divider()
            st.markdown("#### ? AI Readiness Score")
            readiness = 0
            if (p.memory or "").strip():
                readiness += 20
            if len(entries) > 10:
                readiness += 20
            if (p.outline or "").strip():
                readiness += 20
            if chapters:
                readiness += 20
            if not coherence_results:
                readiness += 20
            st.metric("AI Readiness", f"{readiness}%")

    def render_memory():
        p = st.session_state.project
        if not p:
            render_no_project_state("Memory", "empty_memory")
            return

        render_page_header(
            "Memory",
            "Hard canon rules, soft guidance, and project memory used by AI systems.",
            tag="Canon",
            key_prefix="memory_header",
        )

        st.markdown("### World Memory")
        st.caption("Keep canon notes, timelines, and facts the AI should always know.")
        with st.container(border=True):
            st.markdown("#### Auto-add settings")
            st.caption("Auto-add confidence controls are managed in Workspace Settings to avoid duplicate configuration.")
        st.markdown("#### Hard Canon Rules")
        hard_key = f"world_memory_hard_{p.id}"
        hard_default = p.memory_hard or p.memory
        hard_val = st.text_area("Hard Canon Rules", hard_default, height=160, key=hard_key)
        if hard_val != p.memory_hard:
            p.memory_hard = hard_val
            save_p()

        st.markdown("#### Soft Guidelines")
        soft_key = f"world_memory_soft_{p.id}"
        soft_val = st.text_area("Soft Guidelines", p.memory_soft, height=160, key=soft_key)
        if soft_val != p.memory_soft:
            p.memory_soft = soft_val
            save_p()

        st.markdown("#### Project Memory")
        memory_key = f"world_memory_{p.id}"
        memory_val = st.text_area("Memory", p.memory, height=320, key=memory_key)
        if memory_val != p.memory:
            p.memory = memory_val
            save_p()
        if st.button("Save Memory", use_container_width=True):
            if persist_project(p, action="save"):
                st.toast("Memory saved")
                st.rerun()

        st.divider()
        st.markdown("#### Coherence Check")
        scope_cols = st.columns(3)
        with scope_cols[0]:
            scope_outline = st.checkbox("Outline", value=True, key=f"coh_outline_{p.id}")
        with scope_cols[1]:
            scope_world = st.checkbox("World Bible", value=True, key=f"coh_world_{p.id}")
        with scope_cols[2]:
            scope_chapters = st.checkbox("Chapters", value=True, key=f"coh_chapters_{p.id}")

        provider, active_key, _ = get_active_key_status()
        coherence_cooldown = _cooldown_remaining(f"coherence_{p.id}", 15)
        coherence_label = (
            f"Run Coherence Check ({coherence_cooldown}s)"
            if coherence_cooldown
            else "Run Coherence Check"
        )
        if not active_key:
            st.info(
                f"Add a {_provider_label(provider)} API key in AI Settings to run coherence checks."
            )
        if st.button(
            coherence_label,
            use_container_width=True,
            disabled=bool(coherence_cooldown) or not active_key,
            key=f"coh_run_memory_{p.id}",
        ):
            if not (scope_outline or scope_world or scope_chapters):
                st.warning("Choose at least one scope (Outline, World Bible, or Chapters).")
                st.stop()
            model_name = get_ai_model()
            if not model_name:
                st.warning("Choose an AI model in AI Settings first.")
                st.session_state.page = "ai"
                st.rerun()
            _mark_action(f"coherence_{p.id}")
            compiled_world_bible = "\n".join(
                f"{e.name} ({e.category}): {e.description}"
                for e in p.world_db.values()
            )
            chapter_payload = [
                {
                    "chapter_index": c.index,
                    "summary": c.summary,
                    "excerpt": (c.content or "")[:800],
                }
                for c in p.get_ordered_chapters()
            ]
            outline_payload = p.outline if scope_outline else ""
            world_payload = compiled_world_bible if scope_world else ""
            chapters_payload = chapter_payload if scope_chapters else []
            with st.spinner("Running coherence check..."):
                results = AnalysisEngine.coherence_check(
                    memory=p.memory,
                    author_note=p.author_note,
                    style_guide=p.style_guide,
                    outline=outline_payload,
                    world_bible=world_payload,
                    chapters=chapters_payload,
                    model=model_name,
                )
            st.session_state["coherence_results"] = results or []
            canon_icon, _canon_label = get_canon_health()
            st.session_state.setdefault("canon_health_log", [])
            st.session_state["canon_health_log"].append(
                {
                    "timestamp": time.time(),
                    "status": canon_icon,
                    "issue_count": len(st.session_state.get("coherence_results", [])),
                }
            )
            st.session_state["canon_health_log"] = st.session_state["canon_health_log"][-30:]
            if results:
                st.toast("Coherence issues found.")
            else:
                st.toast("No coherence issues detected.")
            st.rerun()

        results = st.session_state.get("coherence_results", [])
        if results:
            st.markdown("#### Coherence Issues")
            for idx, issue in enumerate(list(results)):
                with st.container(border=True):
                    chapter_idx = issue.get("chapter_index", "?")
                    st.markdown(f"**Chapter #{chapter_idx}**")
                    st.markdown(f"**Issue:** {issue.get('issue', 'Unspecified issue')}")
                    st.markdown(f"**Confidence:** {issue.get('confidence', 'unknown')}")
                    st.markdown("**Suggested Rewrite:**")
                    st.write(issue.get("suggested_rewrite", ""))

                    a1, a2 = st.columns(2)
                    with a1:
                        if st.button("Apply Fix", key=f"coh_apply_{idx}", use_container_width=True):
                            target_excerpt = issue.get("target_excerpt", "")
                            try:
                                chapter_num = int(chapter_idx)
                            except (TypeError, ValueError):
                                chapter_num = None
                            target_chapter = None
                            if chapter_num is not None:
                                for c in p.get_ordered_chapters():
                                    if c.index == chapter_num:
                                        target_chapter = c
                                        break
                            if not target_chapter:
                                st.warning("Chapter not found for this issue.")
                            elif target_excerpt and target_excerpt in (target_chapter.content or ""):
                                updated = (target_chapter.content or "").replace(
                                    target_excerpt,
                                    issue.get("suggested_rewrite", ""),
                                    1,
                                )
                                target_chapter.update_content(updated, "Coherence Fix")
                                persist_project(p)
                                results.pop(idx)
                                st.session_state["coherence_results"] = results
                                update_locked_chapters()
                                st.toast("Applied fix.")
                                st.rerun()
                            elif issue.get("suggested_rewrite"):
                                insertion = issue.get("suggested_rewrite", "").strip()
                                if insertion:
                                    spacer = "\n\n" if (target_chapter.content or "").strip() else ""
                                    updated = f"{(target_chapter.content or '').rstrip()}{spacer}{insertion}"
                                    target_chapter.update_content(updated, "Coherence Fix (Appended)")
                                    persist_project(p)
                                    results.pop(idx)
                                    st.session_state["coherence_results"] = results
                                    update_locked_chapters()
                                    st.toast("Applied fix (appended).")
                                    st.rerun()
                            else:
                                st.warning("Target excerpt not found in chapter content.")
                    with a2:
                        if st.button("Ignore", key=f"coh_ignore_{idx}", use_container_width=True):
                            results.pop(idx)
                            st.session_state["coherence_results"] = results
                            update_locked_chapters()
                            st.toast("Issue ignored.")
                            st.rerun()

    def render_insights():
        p = st.session_state.project
        if not p:
            render_no_project_state("Insights", "empty_insights")
            return

        render_page_header(
            "Insights",
            "Decision-grade signals: continuity risk, writing momentum, and next actions.",
            tag="Analytics",
            key_prefix="insights_header",
        )

        entries = list(p.world_db.values())
        chapters = p.get_ordered_chapters()
        chapter_texts = [(c, (c.content or "")) for c in chapters]
        chapter_texts_lower = [(c, text.lower()) for c, text in chapter_texts]

        def _count_mentions(aliases: List[str], text: str) -> int:
            total = 0
            for alias in aliases:
                alias_clean = (alias or "").strip()
                if not alias_clean:
                    continue
                pattern = rf"\\b{re.escape(alias_clean.lower())}\\b"
                total += len(re.findall(pattern, text))
            return total

        mention_counts: Dict[str, int] = {}
        for ent in entries:
            aliases = [ent.name] + (ent.aliases or [])
            total_hits = 0
            for _chap, lower_text in chapter_texts_lower:
                total_hits += _count_mentions(aliases, lower_text)
            mention_counts[ent.id] = total_hits

        orphaned_ids = {eid for eid, count in mention_counts.items() if count == 0}
        under_described_ids = {e.id for e in entries if len((e.description or "").strip()) < 30}
        normalized_name_map: Dict[str, List[str]] = {}
        for ent in entries:
            normalized = Project._normalize_entity_name(ent.name)
            if normalized:
                normalized_name_map.setdefault(normalized, []).append(ent.id)
        collision_ids = {
            ent_id
            for ent_ids in normalized_name_map.values()
            if len(ent_ids) > 1
            for ent_id in ent_ids
        }
        flagged_entity_ids = orphaned_ids | under_described_ids | collision_ids
        coherence_results = st.session_state.get("coherence_results", [])
        review_queue = st.session_state.get("world_bible_review", [])

        total_words = p.get_total_word_count()
        chapter_count = len(chapters)
        active_chapters = len([c for c in chapters if (c.word_count or 0) > 0])
        avg_words = int(total_words / chapter_count) if chapter_count else 0
        weekly_sessions = _weekly_activity_count()
        streak = _activity_streak()
        memory_size = len((p.memory_hard or "").strip()) + len((p.memory_soft or "").strip()) + len((p.memory or "").strip())

        metric_cols = st.columns(5)
        metric_cols[0].metric("Canon issues", len(coherence_results))
        metric_cols[1].metric("Flagged entities", len(flagged_entity_ids))
        metric_cols[2].metric("Drafted chapters", f"{active_chapters}/{chapter_count}")
        metric_cols[3].metric("Avg words/chapter", avg_words)
        metric_cols[4].metric("Review queue", len(review_queue))

        left, right = st.columns([1.45, 1])
        with left:
            with st.container(border=True):
                st.markdown("### Risk Radar")
                risk_items = []
                if coherence_results:
                    risk_items.append(f"Coherence issues open: {len(coherence_results)}")
                if len(flagged_entity_ids) > 0:
                    risk_items.append(f"Entities need review: {len(flagged_entity_ids)}")
                if len(review_queue) > 0:
                    risk_items.append(f"World Bible review queue pending: {len(review_queue)}")
                if not (p.outline or "").strip():
                    risk_items.append("Outline missing")
                if chapter_count == 0:
                    risk_items.append("No chapters drafted yet")
                if not risk_items:
                    st.success("No immediate risks detected.")
                else:
                    for item in risk_items:
                        st.warning(item)
                action_cols = st.columns(3)
                with action_cols[0]:
                    if st.button("Open World Bible", key=f"ins_world_{p.id}", use_container_width=True):
                        st.session_state.page = "world"
                        st.rerun()
                with action_cols[1]:
                    if st.button("Open Memory", key=f"ins_memory_{p.id}", use_container_width=True):
                        st.session_state.page = "memory"
                        st.rerun()
                with action_cols[2]:
                    if st.button("Open Chapters", key=f"ins_chapters_{p.id}", use_container_width=True):
                        st.session_state.page = "chapters"
                        st.rerun()

            with st.container(border=True):
                st.markdown("### Next Best Action")
                if not (p.outline or "").strip():
                    st.info("Create an outline so chapter generation has structure.")
                elif chapter_count == 0:
                    st.info("Create Chapter 1 from your outline.")
                elif coherence_results:
                    st.info("Resolve coherence issues before major rewrites.")
                elif len(flagged_entity_ids) > 0:
                    st.info("Review flagged entities to stabilize canon.")
                else:
                    st.success("Keep drafting. Your project health is strong.")

        with right:
            with st.container(border=True):
                st.markdown("### Momentum")
                st.metric("Sessions (7d)", f"{weekly_sessions}")
                st.metric("Streak", f"{streak} days")
                st.metric("Total words", total_words)
                latest = max(chapters, key=lambda c: (c.modified_at or c.created_at or 0), default=None)
                if latest:
                    st.caption(f"Last touchpoint: Chapter {latest.index} - {latest.title}")
                else:
                    st.caption("Last touchpoint: none yet")

            with st.container(border=True):
                st.markdown("### AI Readiness")
                checks = [
                    ("Memory foundation", memory_size > 120),
                    ("World Bible coverage", len(entries) >= 8),
                    ("Outline complete", bool((p.outline or "").strip())),
                    ("Draft baseline", active_chapters >= 1),
                    ("No open coherence issues", len(coherence_results) == 0),
                ]
                passed = sum(1 for _, ok in checks if ok)
                readiness = int((passed / len(checks)) * 100) if checks else 0
                st.metric("Readiness score", f"{readiness}%")
                for label, ok in checks:
                    st.caption(f"{'?' if ok else '?'} {label}")

        st.divider()
        st.markdown("### Entity Coverage Focus")
        if not entries:
            st.info("No World Bible entries yet.")
        else:
            hotlist = []
            for ent in entries:
                hits = mention_counts.get(ent.id, 0)
                hotlist.append(
                    {
                        "Entity": ent.name,
                        "Category": ent.category,
                        "Mentions": hits,
                        "Priority": "High" if ent.id in flagged_entity_ids else ("Medium" if hits <= 1 else "Low"),
                    }
                )
            hotlist = sorted(hotlist, key=lambda row: (row["Priority"], row["Mentions"]))
            st.dataframe(hotlist[:20], use_container_width=True, hide_index=True)

    def render_chapters():
        from app.views.editor_workspace import render_chapter_sidebar, render_editor_utility_bar

        # Get project first
        p = st.session_state.project

        # Single welcome banner for the editor page (avoid duplicate widgets / duplicate keys).
        render_welcome_banner("editor")
        
        if not p:
            render_no_project_state("Chapters", "empty_chapters")
            return

        def _persist_chapter_update() -> None:
            st.session_state.project = p
            persist_project(p)
            _bump_projects_refresh()
            st.rerun()

        def _serialize_chapter(chapter: Chapter) -> Dict[str, Any]:
            return {
                "id": chapter.id,
                "index": chapter.index,
                "title": chapter.title,
                "content": chapter.content,
                "summary": chapter.summary,
                "word_count": chapter.word_count,
                "target_words": chapter.target_words,
                "created_at": chapter.created_at,
                "modified_at": chapter.modified_at,
            }

        def _set_active_chapter(chapter_id: Optional[str]) -> None:
            st.session_state.active_chapter_id = chapter_id
            st.session_state.curr_chap_id = chapter_id

        def _sync_session_chapters(force: bool = False) -> None:
            if (
                force
                or st.session_state.get("chapters_project_id") != p.id
                or not isinstance(st.session_state.get("chapters"), list)
            ):
                st.session_state.chapters = [_serialize_chapter(c) for c in p.get_ordered_chapters()]
                st.session_state.chapters_project_id = p.id
            chapter_ids = [c.get("id") for c in st.session_state.get("chapters", []) if c.get("id")]
            if st.session_state.get("active_chapter_id") not in chapter_ids:
                _set_active_chapter(chapter_ids[0] if chapter_ids else None)
            elif st.session_state.get("active_chapter_id"):
                st.session_state.curr_chap_id = st.session_state.active_chapter_id

        def _update_session_chapter(chapter: Chapter) -> None:
            chapters = st.session_state.get("chapters", [])
            for idx, item in enumerate(chapters):
                if item.get("id") == chapter.id:
                    chapters[idx] = {**item, **_serialize_chapter(chapter)}
                    st.session_state.chapters = chapters
                    return

        def _append_session_chapter(chapter: Chapter) -> None:
            chapters = st.session_state.get("chapters", [])
            chapters.append(_serialize_chapter(chapter))
            st.session_state.chapters = chapters

        _sync_session_chapters()
        chapters = st.session_state.get("chapters", [])

        def _outline_chapter_title(chapter_index: int) -> str:
            outline_text = p.outline or ""
            if not outline_text.strip():
                return ""
            patterns = [
                rf"Chapter\s*{chapter_index}\s*[:\-]\s*(.+?)(?=\n|$)",
                rf"^\s*{chapter_index}\.\s*\*{0,2}\s*(.+?)\s*\*{0,2}\s*(?:[-:]\s*.+)?$",
                rf"^\s*{chapter_index}\)\s*\*{0,2}\s*(.+?)\s*\*{0,2}\s*(?:[-:]\s*.+)?$",
            ]
            for pattern in patterns:
                match = re.search(pattern, outline_text, re.IGNORECASE | re.MULTILINE)
                if not match:
                    continue
                raw = (match.group(1) or "").strip()
                raw = re.sub(r"\*\*", "", raw).strip()
                raw = re.split(r"\s[-:]\s", raw, 1)[0].strip()
                cleaned = sanitize_chapter_title(raw)
                if cleaned:
                    return cleaned
            return ""

        def create_next_chapter() -> None:
            try:
                ordered = p.get_ordered_chapters()
                next_idx = (ordered[-1].index + 1) if ordered else 1
                title = f"Chapter {next_idx}"
                outline_title = _outline_chapter_title(next_idx)
                if outline_title:
                    title = outline_title
                new_chapter = p.add_chapter(title)
                _append_session_chapter(new_chapter)
                _set_active_chapter(new_chapter.id)
                _persist_chapter_update()
            except Exception as exc:
                logger.warning("Failed to create chapter", exc_info=True)
                st.error(f"Could not create chapter: {exc}")

        render_page_header(
            "Editor",
            "Write chapters, update summaries, and apply AI improvements.",
            tag="Drafting",
            key_prefix="editor_header",
        )

        if not chapters:
            with st.container(border=True):
                st.markdown("### Start your draft")
                st.caption("No chapters yet. Create Chapter 1 to begin writing, or open Outline to plan beats first.")
                cta_cols = st.columns(2)
                with cta_cols[0]:
                    if st.button(
                        "Create Chapter 1",
                        use_container_width=True,
                        help="Create your first chapter and start drafting",
                        key="editor_create_chapter_1",
                    ):
                        create_next_chapter()
                with cta_cols[1]:
                    if st.button(
                        "Open Outline",
                        use_container_width=True,
                        help="Plan your chapter flow before drafting",
                        key="editor_empty_open_outline",
                    ):
                        st.session_state.page = "outline"
                        st.rerun()

                if (p.outline or "").strip():
                    st.success("Outline detected. Chapter title will auto-align with your outline when available.")
            return

        chapter_map = {c["id"]: c for c in chapters if c.get("id")}
        active_id = st.session_state.get("active_chapter_id")
        if active_id not in chapter_map:
            _set_active_chapter(chapters[0]["id"] if chapters else None)
            active_id = st.session_state.get("active_chapter_id")

        curr = p.chapters.get(active_id) if active_id else None
        if curr is None and active_id:
            _sync_session_chapters(force=True)
            chapters = st.session_state.get("chapters", [])
            chapter_map = {c["id"]: c for c in chapters if c.get("id")}
            active_id = st.session_state.get("active_chapter_id")
            curr = p.chapters.get(active_id) if active_id else None
        if curr is None:
            st.warning("Select a chapter to begin editing.")
            return
        curr_entry = chapter_map.get(curr.id, _serialize_chapter(curr))
        locked_chapters = st.session_state.get("locked_chapters", set())

        render_editor_utility_bar(
            project=p,
            current_chapter=curr,
            on_open_outline=lambda: (
                st.session_state.update({"page": "outline"}),
                st.rerun(),
            ),
            on_open_world=lambda: (
                st.session_state.update({"page": "world"}),
                st.rerun(),
            ),
            on_word_target_change=lambda value: (
                setattr(p, "default_word_count", int(value)),
                save_p(),
            ),
        )
        # --- SAFELY sync programmatic chapter updates into the editor widget (before widget exists)
        ed_key = f"ed_{curr.id}"
        if (
            st.session_state.get("_chapter_sync_id") == curr.id
            and ed_key in st.session_state
        ):
            st.session_state[ed_key] = st.session_state.get("_chapter_sync_text", "") or ""
            st.session_state._chapter_sync_id = None
            st.session_state._chapter_sync_text = None

        outline_title = _outline_chapter_title(curr.index)
        if outline_title and (
            "Untitled" in (curr.title or "")
            or (curr.title or "").strip().lower() in {f"chapter {curr.index}", f"chapter {curr.index} (0)"}
        ):
            curr.title = outline_title
            _update_session_chapter(curr)
            persist_project(p)

        col_nav, col_editor, col_ai = st.columns([1.05, 3.2, 1.25])

        with col_nav:
            def _delete_current_chapter() -> None:
                p.delete_chapter(curr.id)
                _sync_session_chapters(force=True)
                refreshed = st.session_state.get("chapters", [])
                _set_active_chapter(refreshed[0]["id"] if refreshed else None)
                st.session_state.delete_chapter_id = None
                st.session_state.delete_chapter_title = None
                persist_project(p)
                st.toast("Chapter deleted.")
                st.rerun()

            render_chapter_sidebar(
                chapters=chapters,
                current_chapter_id=curr.id,
                on_select=lambda chapter_id: (_set_active_chapter(chapter_id), st.rerun()),
                on_create_next=create_next_chapter,
                on_delete_current=_delete_current_chapter,
                delete_pending=st.session_state.get("delete_chapter_id") == curr.id,
                current_title=st.session_state.get("delete_chapter_title") or curr.title,
                on_cancel_delete=lambda: (
                    st.session_state.update(
                        {"delete_chapter_id": None, "delete_chapter_title": None}
                    ),
                    st.rerun(),
                ),
            )

        stream_ph = st.empty()
        provider, active_key, _ = get_active_key_status()

        with col_editor:
            with st.container(border=True):
                title_key = f"editor_title_{curr.id}"
                title_val = st.text_input(
                    "Chapter Title",
                    curr_entry.get("title", curr.title),
                    label_visibility="collapsed",
                    help="Edit the title of this chapter",
                    key=title_key,
                )
                if title_val != curr.title:
                    curr.title = title_val
                    _update_session_chapter(curr)
                    save_p()

                val = st.text_area(
                    "Chapter Content",
                    curr_entry.get("content", curr.content),
                    height=680,
                    label_visibility="collapsed",
                    key=f"ed_{curr.id}",
                    help="Write your chapter content here. Changes are auto-saved."
                )
                if val != curr.content:
                    curr.update_content(val, "manual")
                    _update_session_chapter(curr)
                    save_p()

                st.caption(f"Chapter: {curr.word_count} words  Total: {p.get_total_word_count()} words")

                c1, c2 = st.columns([1, 1])
                with c1:
                    if st.button(
                        "Save Chapter",
                        type="primary",
                        use_container_width=True,
                        key=f"editor_save_{curr.id}",
                    ):
                        curr.update_content(val, "manual")
                        _update_session_chapter(curr)
                        if persist_project(p, action="save"):
                            # Automatically scan entities from this chapter when the user explicitly saves it.
                            scan_result = extract_entities_ui(
                                curr.content or "",
                                f"Ch {curr.index}",
                                show_feedback=False,
                            )
                            if scan_result:
                                if scan_result.get("status") == "updated":
                                    st.success(
                                        "Chapter saved and scanned: "
                                        f"{scan_result.get('added', 0)} new, "
                                        f"{scan_result.get('matched', 0)} existing, "
                                        f"{scan_result.get('flagged', 0)} flagged."
                                    )
                                elif scan_result.get("status") == "matched_only":
                                    st.info("Chapter saved. Entities were detected and matched existing entries.")
                                else:
                                    st.info("Chapter saved. No entities were detected in this chapter.")
                            else:
                                st.success("Chapter saved.")
                with c2:
                    summary_cooldown = _cooldown_remaining(f"summary_{curr.id}", 10)
                    summary_label = (
                        f"Update Summary ({summary_cooldown}s)"
                        if summary_cooldown
                        else "Update Summary"
                    )
                    if not active_key:
                        st.info(
                            f"Add a {_provider_label(provider)} API key in AI Settings to update summaries."
                        )
                    if st.button(
                        summary_label,
                        use_container_width=True,
                        disabled=bool(summary_cooldown) or not active_key,
                        key=f"editor_summary_{curr.id}",
                    ):
                        _mark_action(f"summary_{curr.id}")
                        summary = StoryEngine.summarize(curr.content or "", get_ai_model())
                        if summary.strip().startswith("ERROR: Canon violation detected."):
                            st.error("Canon violation detected. Resolve issues before generating AI content.")
                        else:
                            curr.summary = summary
                            _update_session_chapter(curr)
                            persist_project(p)
                            st.rerun()

                def generate_improvement(style_choice, custom_instructions):
                    prompt = rewrite_prompt(curr.content or "", style_choice, custom_instructions)
                    full = ""
                    for chunk in AIEngine().generate_stream(prompt, get_ai_model()):
                        full += chunk
                        stream_ph.markdown(f"**IMPROVING:**\n\n{full}")
                    if full.strip().startswith("ERROR: Canon violation detected."):
                        st.error("Canon violation detected. Resolve issues before generating AI content.")
                        return ""
                    violations = detect_hard_canon_violation(p, curr.index, full)
                    if violations:
                        results = st.session_state.get("coherence_results", [])
                        results.extend(violations)
                        st.session_state["coherence_results"] = results
                        update_locked_chapters()
                        st.warning("Hard Canon conflict detected. AI output was not applied.")
                        return ""
                    return full

                with st.expander("Modify / Improve Text"):
                    rewrite_locked = curr.index in locked_chapters
                    style = st.selectbox(
                        "How to improve?",
                        list(REWRITE_PRESETS.keys()),
                        disabled=rewrite_locked,
                        key=f"editor_improve__style_{curr.id}",
                    )
                    cust = ""
                    if style == "Custom":
                        cust = st.text_input(
                            "Instructions",
                            disabled=rewrite_locked,
                            key=f"editor_improve__instructions_{curr.id}",
                        )
                    if rewrite_locked:
                        st.caption("Rewrite tools are disabled for locked chapters.")
                    if not active_key:
                        st.info(
                            f"Add a {_provider_label(provider)} API key in AI Settings to generate improvements."
                        )
                    improve_cooldown = _cooldown_remaining(f"improve_{curr.id}", 12)
                    improve_label = (
                        f"Generate Improvement ({improve_cooldown}s)"
                        if improve_cooldown
                        else "Generate Improvement"
                    )
                    if st.button(
                        improve_label,
                        use_container_width=True,
                        disabled=rewrite_locked or bool(improve_cooldown) or not active_key,
                        key=f"editor_improve__generate_{curr.id}",
                    ):
                        _mark_action(f"improve_{curr.id}")
                        full = generate_improvement(style, cust)
                        if full:
                            st.session_state.pending_improvement_text = full
                            st.session_state.pending_improvement_meta = {
                                "mode": style,
                                "custom": cust,
                                "timestamp": time.time(),
                                "chapter_id": curr.id,
                            }
                            st.rerun()

                if st.button(
                    "Undo last apply",
                    use_container_width=True,
                    key=f"editor_improve__undo_{curr.id}",
                    help="Restore the previous chapter text, if available.",
                ):
                    prev = st.session_state.get("chapter_text_prev", {})
                    if prev.get("chapter_id") == curr.id and prev.get("text") is not None:
                        curr.update_content(prev.get("text") or "", "Undo Apply")
                        _update_session_chapter(curr)
                        persist_project(p)
                        st.session_state._chapter_sync_id = curr.id
                        st.session_state._chapter_sync_text = prev.get("text") or ""
                        st.toast("Previous chapter text restored.")
                        st.rerun()
                    else:
                        st.info("No previous chapter text available to restore.")

                chapter_drafts = [
                    draft for draft in st.session_state.get("chapter_drafts", [])
                    if draft.get("chapter_id") == curr.id
                ]
                if chapter_drafts:
                    st.markdown("#### Draft Versions")
                    for idx, draft in enumerate(chapter_drafts, start=1):
                        label = f"Draft {idx}  {draft.get('mode', 'Improved')}  {time.strftime('%Y-%m-%d %H:%M', time.localtime(draft.get('timestamp', 0)))}"
                        with st.expander(label):
                            st.text_area(
                                "Draft Text",
                                draft.get("text", ""),
                                height=200,
                                disabled=True,
                                key=f"editor_improve__draft_text_{curr.id}_{idx}",
                            )
                            if st.button(
                                "Set as active",
                                type="primary",
                                use_container_width=True,
                                key=f"editor_improve__set_active_{curr.id}_{idx}",
                            ):
                                st.session_state.chapter_text_prev = {
                                    "chapter_id": curr.id,
                                    "text": curr.content or "",
                                }
                                curr.update_content(draft.get("text", ""), "Draft Applied")
                                _update_session_chapter(curr)
                                persist_project(p)
                                st.session_state._chapter_sync_id = curr.id
                                st.session_state._chapter_sync_text = draft.get("text", "")
                                st.toast("Draft set as active.")
                                st.rerun()

        with col_ai:
            with st.container(border=True):
                st.markdown("### Assistant")
                st.caption("Generate new prose from your outline + previous context.")

                canon_icon, canon_label = get_canon_health()
                canon_blocked = canon_icon == "RISK"
                if canon_blocked:
                    st.button(
                        "Auto-Write Disabled (Canon Risk)",
                        disabled=True,
                        use_container_width=True,
                        help="Resolve canon issues in World Bible ? Memory before generating.",
                        key=f"editor_autowrite_blocked_{curr.id}",
                    )
                else:
                    if not active_key:
                        st.info(
                            f"Add a {_provider_label(provider)} API key in AI Settings to auto-write chapters."
                        )
                    auto_cooldown = _cooldown_remaining(f"auto_write_{curr.id}", 12)
                    auto_label = (
                        f"Auto-Write Chapter ({auto_cooldown}s)"
                        if auto_cooldown
                        else "Auto-Write Chapter"
                    )
                    if st.button(
                        auto_label,
                        type="primary",
                        use_container_width=True,
                        disabled=bool(auto_cooldown) or not active_key,
                        key=f"editor_autowrite_{curr.id}",
                    ):
                        _mark_action(f"auto_write_{curr.id}")
                        chapter_lengths = [
                            int(ch.word_count or 0)
                            for ch in p.get_ordered_chapters()
                            if int(ch.word_count or 0) > 0
                        ]
                        avg_length = int(sum(chapter_lengths) / len(chapter_lengths)) if chapter_lengths else 0
                        ai_target_words = int(avg_length or p.default_word_count or 1000)
                        prompt = StoryEngine.generate_chapter_prompt(p, curr.index, ai_target_words)
                        full = ""
                        for chunk in AIEngine().generate_stream(prompt, get_ai_model()):
                            full += chunk
                            stream_ph.markdown(f"**GENERATING:**\n\n{full}")

                        if full.strip():
                            if full.strip().startswith("ERROR: Canon violation detected."):
                                st.error("Canon violation detected. Resolve issues before generating AI content.")
                                return
                            violations = detect_hard_canon_violation(p, curr.index, full)
                            if violations:
                                results = st.session_state.get("coherence_results", [])
                                results.extend(violations)
                                st.session_state["coherence_results"] = results
                                update_locked_chapters()
                                st.warning("Hard Canon conflict detected. AI output was not applied.")
                                return
                            new_text = ((curr.content or "") + "\n" + full.strip()).strip()
                            curr.update_content(new_text, "AI Auto-Write")
                            _update_session_chapter(curr)
                            persist_project(p)

                            # Queue sync into editor widget on next run (do NOT set ed_key here!)
                            st.session_state._chapter_sync_id = curr.id
                            st.session_state._chapter_sync_text = new_text

                            st.toast("Chapter Updated")
                            time.sleep(0.3)
                            st.rerun()

                st.divider()
                st.markdown("#### Summary")
                if curr.summary:
                    st.info(curr.summary)
                else:
                    st.info("No summary yet. Click **Update Summary** to generate one.")

        pending_text = st.session_state.get("pending_improvement_text") or ""
        pending_meta = st.session_state.get("pending_improvement_meta") or {}
        pending_for_chapter = pending_text and pending_meta.get("chapter_id") == curr.id
        if pending_for_chapter:
            with st.container(border=True):
                st.markdown("### ? Review Changes")
                diff_toggle = st.toggle(
                    "Diff view",
                    value=False,
                    key=f"editor_improve__diff_toggle_{curr.id}",
                    help="Show a unified diff instead of the full texts.",
                )
                col_left, col_right = st.columns(2)
                if diff_toggle:
                    diff_lines = difflib.unified_diff(
                        (curr.content or "").splitlines(),
                        (pending_text or "").splitlines(),
                        fromfile="Original",
                        tofile="Improved",
                        lineterm="",
                    )
                    diff_text = "\n".join(diff_lines).strip() or "No differences found."
                    with col_left:
                        st.text_area(
                            "Original",
                            curr.content or "",
                            height=260,
                            disabled=True,
                            key=f"editor_improve__original_{curr.id}",
                        )
                    with col_right:
                        st.code(diff_text, language="diff")
                else:
                    with col_left:
                        st.text_area(
                            "Original",
                            curr.content or "",
                            height=260,
                            disabled=True,
                            key=f"editor_improve__original_{curr.id}",
                        )
                    with col_right:
                        st.text_area(
                            "Improved",
                            pending_text or "",
                            height=260,
                            disabled=True,
                            key=f"editor_improve__improved_{curr.id}",
                        )

                apply_mode = st.selectbox(
                    "Apply result as",
                    [
                        "Replace chapter",
                        "Save as new draft/version",
                        "Append to end (advanced)",
                    ],
                    key=f"editor_improve__apply_mode_{curr.id}",
                )

                b1, b2, b3, b4 = st.columns([1, 1, 1, 1])
                with b1:
                    if st.button(
                        "Apply Changes",
                        type="primary",
                        use_container_width=True,
                        key=f"editor_improve__apply_{curr.id}",
                    ):
                        if apply_mode == "Replace chapter":
                            st.session_state.chapter_text_prev = {
                                "chapter_id": curr.id,
                                "text": curr.content or "",
                            }
                            new_text = pending_text or ""
                            curr.update_content(new_text, "AI Improve")
                            _update_session_chapter(curr)
                            persist_project(p)
                            st.session_state._chapter_sync_id = curr.id
                            st.session_state._chapter_sync_text = new_text
                            st.toast("Chapter replaced with improved text.")
                        elif apply_mode == "Save as new draft/version":
                            draft_entry = {
                                "chapter_id": curr.id,
                                "text": pending_text or "",
                                "mode": pending_meta.get("mode", "Improve"),
                                "timestamp": time.time(),
                            }
                            st.session_state.chapter_drafts = (
                                st.session_state.get("chapter_drafts", []) + [draft_entry]
                            )
                            st.toast("Saved as a new draft version.")
                        else:
                            st.session_state.chapter_text_prev = {
                                "chapter_id": curr.id,
                                "text": curr.content or "",
                            }
                            new_text = ((curr.content or "") + "\n" + (pending_text or "")).strip()
                            curr.update_content(new_text, "AI Improve Append")
                            _update_session_chapter(curr)
                            persist_project(p)
                            st.session_state._chapter_sync_id = curr.id
                            st.session_state._chapter_sync_text = new_text
                            st.toast("Improved text appended to chapter.")

                        st.session_state.pending_improvement_text = ""
                        st.session_state.pending_improvement_meta = {}
                        st.rerun()
                with b2:
                    if st.button(
                        "Copy Improved",
                        use_container_width=True,
                        key=f"editor_improve__copy_{curr.id}",
                    ):
                        st.session_state.editor_improve__copy_buffer = pending_text or ""
                        st.toast("Improved text ready to copy (Ctrl/Cmd+C).")
                with b3:
                    regen_cooldown = _cooldown_remaining(f"improve_regen_{curr.id}", 12)
                    regen_label = (
                        f"Regenerate ({regen_cooldown}s)" if regen_cooldown else "Regenerate"
                    )
                    if st.button(
                        regen_label,
                        use_container_width=True,
                        disabled=bool(regen_cooldown) or not active_key,
                        key=f"editor_improve__regenerate_{curr.id}",
                    ):
                        _mark_action(f"improve_regen_{curr.id}")
                        regen_style = pending_meta.get("mode", "Improve Flow")
                        regen_custom = pending_meta.get("custom", "")
                        full = generate_improvement(regen_style, regen_custom)
                        if full:
                            st.session_state.pending_improvement_text = full
                            st.session_state.pending_improvement_meta = {
                                "mode": regen_style,
                                "custom": regen_custom,
                                "timestamp": time.time(),
                                "chapter_id": curr.id,
                            }
                            st.rerun()
                with b4:
                    if st.button(
                        "Discard",
                        use_container_width=True,
                        key=f"editor_improve__discard_{curr.id}",
                    ):
                        st.session_state.pending_improvement_text = ""
                        st.session_state.pending_improvement_meta = {}
                        st.rerun()

    def _render_current_page() -> None:
        rendered_page = False
        current_page = st.session_state.get("page")
        logger.info(f"Rendering page: {current_page}")
        last_page = st.session_state.get("_last_rendered_page")
        scroll_nonce = int(st.session_state.get("_scroll_top_nonce", 0))
        last_scroll_nonce = int(st.session_state.get("_last_scroll_top_nonce", -1))
        should_scroll_top = (current_page != last_page) or (scroll_nonce != last_scroll_nonce)
        if should_scroll_top:
            st.session_state["_last_rendered_page"] = current_page
            st.session_state["_last_scroll_top_nonce"] = scroll_nonce
            components.html(
                """
                <script>
                  (function () {
                    const topWin = window.parent || window;
                    const topDoc = topWin.document || document;
                    function doScroll() {
                      const targets = [
                        topDoc.querySelector('section.main'),
                        topDoc.querySelector('[data-testid="stAppViewContainer"]'),
                        topDoc.querySelector('[data-testid="stMain"]'),
                        topDoc.querySelector('main'),
                        topDoc.scrollingElement,
                        topDoc.documentElement,
                        topDoc.body,
                      ].filter(Boolean);
                      targets.forEach((el) => {
                        try {
                          if (typeof el.scrollTo === "function") {
                            el.scrollTo({ top: 0, behavior: "auto" });
                          }
                          el.scrollTop = 0;
                        } catch (_) {}
                      });
                      try { topWin.scrollTo({ top: 0, behavior: "auto" }); } catch (_) {}
                      try { topWin.scrollTo(0, 0); } catch (_) {}
                      try {
                        if (topWin.location && topWin.location.hash) {
                          topWin.location.hash = "";
                        }
                      } catch (_) {}
                    }
                    doScroll();
                    try { topWin.requestAnimationFrame(doScroll); } catch (_) {}
                    setTimeout(doScroll, 40);
                    setTimeout(doScroll, 140);
                    setTimeout(doScroll, 320);
                    setTimeout(doScroll, 700);
                  })();
                </script>
                """,
                height=0,
            )
        
        if st.session_state.page == "home":
            logger.debug("Rendering home/dashboard page")
            with key_scope("dashboard"):
                render_home()
            rendered_page = True
        elif st.session_state.page == "projects":
            logger.debug("Rendering projects page")
            with key_scope("projects"):
                render_projects()
            rendered_page = True
        elif st.session_state.page == "ai":
            logger.debug("Rendering AI settings page")
            with key_scope("settings"):
                render_ai_settings()
            rendered_page = True
        elif st.session_state.page == "workspace":
            logger.debug("Rendering workspace settings page")
            with key_scope("workspace"):
                render_workspace_settings()
            rendered_page = True
        elif st.session_state.page == "legal":
            logger.debug("Rendering legal page")
            with key_scope("legal"):
                render_legal_redirect()
            rendered_page = True
        elif st.session_state.page == "privacy":
            logger.debug("Rendering privacy page")
            with key_scope("privacy"):
                render_privacy()
            rendered_page = True
        elif st.session_state.page == "terms":
            logger.debug("Rendering terms page")
            with key_scope("terms"):
                render_terms()
            rendered_page = True
        elif st.session_state.page == "copyright":
            logger.debug("Rendering copyright page")
            with key_scope("copyright"):
                render_copyright()
            rendered_page = True
        elif st.session_state.page == "cookie":
            logger.debug("Rendering cookie page")
            with key_scope("cookie"):
                render_cookie()
            rendered_page = True
        elif st.session_state.page == "help":
            logger.debug("Rendering help page")
            with key_scope("help"):
                render_help()
            rendered_page = True
        elif st.session_state.page == "export":
            logger.debug("Legacy export route hit; redirecting to projects")
            st.session_state.page = "projects"
            st.rerun()
        else:
            pg = st.session_state.page
            if pg == "outline":
                logger.debug("Rendering outline page")
                with key_scope("outline"):
                    render_outline()
                rendered_page = True
            elif pg == "world":
                logger.debug("Rendering world bible page")
                with key_scope("world"):
                    render_world()
                rendered_page = True
            elif pg == "memory":
                logger.debug("Rendering memory page")
                with key_scope("memory"):
                    render_memory()
                rendered_page = True
            elif pg == "insights":
                logger.debug("Rendering insights page")
                with key_scope("insights"):
                    render_insights()
                rendered_page = True
            elif pg == "chapters":
                logger.debug("Rendering chapters/editor page")
                with key_scope("editor"):
                    render_chapters()
                rendered_page = True
            else:
                logger.warning(f"Unknown page '{pg}', redirecting to home")
                st.session_state.page = "home"
                st.rerun()
        if rendered_page:
            logger.debug("Rendering footer")
            render_app_footer()
            logger.info(f"Page '{current_page}' rendered successfully")

    # Render with comprehensive error handling and fallback UI
    try:
        logger.info("=" * 60)
        logger.info("Starting page render cycle")
        logger.info("=" * 60)
        st.markdown('<div id="mantis-top"></div>', unsafe_allow_html=True)
        _render_current_page()
        logger.info("Page render cycle completed successfully")
    except Exception as exc:
        # Comprehensive error handling with fallback UI to ensure app never shows blank page
        error_msg = f"{type(exc).__name__}: {exc}"
        st.session_state["last_exception"] = error_msg
        logger.error("=" * 60)
        logger.error("UNHANDLED UI EXCEPTION")
        logger.error("=" * 60)
        logger.error(f"Page: {st.session_state.get('page', 'unknown')}")
        logger.error(f"Error: {error_msg}")
        logger.exception("Exception details:", exc_info=True)
        logger.error("=" * 60)
        
        # Render fallback error UI
        try:
            st.error("**Something went wrong while rendering this page.**")
            st.markdown("""
            ### Troubleshooting Steps:
            1. Try reloading the app (F5 or Ctrl+R)
            2. Return to the dashboard using the sidebar or button below
            3. Check the terminal/logs for detailed error messages
            4. If the issue persists, please report it on GitHub with the error details below
            """)
            
            with st.expander("Error Details (Click to expand)", expanded=False):
                st.code(error_msg, language="text")
                if debug_enabled():
                    st.write("**Debug Mode Active - Full Stack Trace:**")
                    st.exception(exc)
            
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Return to Dashboard", use_container_width=True, type="primary", key="main_error_fallback_home"):
                    st.session_state.page = "home"
                    st.rerun()
            with col2:
                if st.button("Reload App", use_container_width=True, key="main_error_fallback_reload"):
                    st.rerun()
        except Exception as fallback_exc:
            # Last resort: even the error UI failed
            logger.critical(f"Error fallback UI failed: {fallback_exc}", exc_info=True)
            try:
                st.text("Critical error - please refresh the page")
                if st.button("Refresh", key="main_error_critical_refresh"):
                    st.rerun()
            except Exception:
                # Nothing we can do at this point
                pass


# Execute the UI when running under Streamlit (not when imported by tests/other modules)
# This is the main entry point when `streamlit run app/main.py` is executed
try:
    import streamlit as st
    from streamlit import runtime
    
    # Only run UI if we're already inside a Streamlit runtime
    # (which happens when `streamlit run app/main.py` is executed)
    if runtime.exists():
        _run_ui()
except ImportError:
    # Streamlit not available - don't run UI
    # This allows imports in tests and other contexts where Streamlit is not installed
    pass


def run_selftest() -> int:
    """
    Quick non-UI integrity test. Intended to be called by the launcher.
    Creates a temporary project, exercises save/load, chapters/entities/export, then cleans up.
    """
    print("[MANTIS SELFTEST]")
    try:
        os.makedirs(AppConfig.PROJECTS_DIR, exist_ok=True)
        # Local-first architecture: all projects stored in default directory
        user_projects_dir = AppConfig.PROJECTS_DIR

        p = Project.create("SELFTEST_PROJECT", author="MANTIS", genre="Test", storage_dir=user_projects_dir)
        p.outline = "Chapter 1: Test - This is a test outline."
        path = p.save()
        if not os.path.exists(path):
            raise RuntimeError("Save failed: project file not found on disk.")

        p2 = Project.load(path)
        if p2.title != "SELFTEST_PROJECT":
            raise RuntimeError("Load failed: project title mismatch.")

        ch = p2.add_chapter("Test Chapter", "Hello world")
        if ch.word_count != 2:
            raise RuntimeError(f"Chapter word_count incorrect: expected 2, got {ch.word_count}")

        ent = p2.add_entity("Tester", "Character", "A test entity.")
        if ent is None:
            raise RuntimeError("Entity add failed: returned None.")

        md = project_to_markdown(p2)
        if "# SELFTEST_PROJECT" not in md:
            raise RuntimeError("Export markdown failed: missing title header.")

        # Cleanup
        try:
            os.remove(path)
        except OSError:
            logger.warning("Selftest cleanup failed for %s", path, exc_info=True)
        try:
            if os.path.isdir(user_projects_dir):
                shutil.rmtree(user_projects_dir)
        except OSError:
            logger.warning("Selftest cleanup failed for %s", user_projects_dir, exc_info=True)
        print("SELFTEST RESULT: PASS")
        return 0
    except Exception as e:
        logger.error("Selftest failed", exc_info=True)
        print("SELFTEST RESULT: FAIL")
        print(str(e))
        return 1



def run_repair() -> int:
    """
    Repairs/normalizes project JSON files in the projects folder.
    - Creates a timestamped backup copy into projects/.backups
    - Attempts to load via Project.load (validates schema)
    - Rewrites JSON in a normalized format (indent=2, ensure_ascii=False)
    Returns 0 on success, 1 if any file failed.
    """
    projects_dir = AppConfig.PROJECTS_DIR
    backups_dir = AppConfig.BACKUPS_DIR
    os.makedirs(projects_dir, exist_ok=True)
    os.makedirs(backups_dir, exist_ok=True)

    files = [f for f in os.listdir(projects_dir) if f.lower().endswith(".json")]
    if not files:
        print("[REPAIR] No project files found.")
        return 0

    stamp = time.strftime("%Y%m%d_%H%M%S")
    failures = 0

    def _atomic_write(path: str, data: dict) -> None:
        tmp = path + ".tmp"
        with open(tmp, "w", encoding="utf-8") as fh:
            json.dump(data, fh, indent=2, ensure_ascii=False)
        os.replace(tmp, path)

    print(f"[REPAIR] Found {len(files)} project file(s). Backups -> {backups_dir}")
    for fn in files:
        full = os.path.join(projects_dir, fn)
        try:
            # Backup first
            backup_name = f"{stamp}__{fn}"
            backup_path = os.path.join(backups_dir, backup_name)
            try:
                with open(full, "rb") as src, open(backup_path, "wb") as dst:
                    dst.write(src.read())
            except Exception:
                # Backup failure shouldn't block repair attempt, but count it
                print(f"[REPAIR][WARN] Could not backup: {fn}")
                logger.warning("Repair backup failed for %s", full, exc_info=True)

            # Validate/normalize via Project model
            proj = Project.load(full)
            data = proj.to_dict()
            _atomic_write(full, data)
            print(f"[REPAIR][OK] {fn}")
        except Exception as e:
            failures += 1
            print(f"[REPAIR][FAIL] {fn} :: {e}")
            logger.error("Repair failed for %s", full, exc_info=True)

    if failures:
        print(f"[REPAIR] Completed with failures: {failures}")
        return 1

    print("[REPAIR] Completed successfully.")
    return 0


# If launched directly by Python (e.g., from the .bat launcher), support utility flags.
# Only execute when run as main script, not when imported by tests or other modules.
def _is_running_under_streamlit() -> bool:
    """Best-effort detection for scripts executed by Streamlit itself."""
    if os.getenv("STREAMLIT_SERVER_PORT"):
        return True
    try:
        from streamlit import runtime

        return bool(runtime.exists())
    except Exception:
        return False


if __name__ == "__main__":
    if SELFTEST_MODE:
        raise SystemExit(run_selftest())

    if REPAIR_MODE:
        raise SystemExit(run_repair())

    # Default: run the Streamlit UI.
    # Important for cloud deploys: avoid nested `streamlit run` launches when
    # Streamlit is already the process owner (can manifest as redirect loops).
    if not _is_running_under_streamlit():
        import streamlit.web.cli as stcli
        sys.argv = ["streamlit", "run", __file__]
        sys.exit(stcli.main())









